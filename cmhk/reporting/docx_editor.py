"""Round-trip CMHK Word reports through the in-browser report editor.

The browser edits a constrained Tiptap JSON document.  This module converts
the existing DOCX body to that JSON and writes edited JSON back into a copy of
the original package.  Reusing the source package keeps its page setup,
headers, footers, styles, theme, and embedded brand assets intact.
"""

from __future__ import annotations

import base64
import binascii
import hashlib
import re
import shutil
import uuid
from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


MAX_EDITOR_NODES = 20_000
MAX_EDITOR_TEXT_CHARS = 1_200_000
MAX_EDITOR_IMAGE_BYTES = 8 * 1024 * 1024

_ALIGN_TO_WEB = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
}
_WEB_TO_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_HIGHLIGHT_TO_HEX = {
    WD_COLOR_INDEX.YELLOW: "#ffff00",
    WD_COLOR_INDEX.BRIGHT_GREEN: "#00ff00",
    WD_COLOR_INDEX.TURQUOISE: "#00ffff",
    WD_COLOR_INDEX.PINK: "#ff00ff",
    WD_COLOR_INDEX.BLUE: "#0000ff",
    WD_COLOR_INDEX.RED: "#ff0000",
    WD_COLOR_INDEX.DARK_BLUE: "#000080",
    WD_COLOR_INDEX.TEAL: "#008080",
    WD_COLOR_INDEX.GREEN: "#008000",
    WD_COLOR_INDEX.VIOLET: "#800080",
    WD_COLOR_INDEX.DARK_RED: "#800000",
    WD_COLOR_INDEX.DARK_YELLOW: "#808000",
    WD_COLOR_INDEX.GRAY_50: "#808080",
    WD_COLOR_INDEX.GRAY_25: "#c0c0c0",
    WD_COLOR_INDEX.BLACK: "#000000",
}
_HEX_TO_HIGHLIGHT = {value: key for key, value in _HIGHLIGHT_TO_HEX.items()}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _inches(value: Any) -> float:
    try:
        return round(float(value) / 914400.0, 4)
    except (TypeError, ValueError):
        return 0.0


def _points(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except (AttributeError, TypeError, ValueError):
        return None


def _style_chain(style: Any) -> Iterable[Any]:
    """Yield one Word style and its bases without looping on malformed files."""
    seen: set[str] = set()
    current = style
    while current is not None:
        style_id = str(getattr(current, "style_id", "") or id(current))
        if style_id in seen:
            break
        seen.add(style_id)
        yield current
        current = getattr(current, "base_style", None)


def _normal_style(paragraph: Paragraph) -> Any | None:
    try:
        return paragraph.part.document.styles["Normal"]
    except (AttributeError, KeyError):
        return None


def _paragraph_format_sources(paragraph: Paragraph) -> Iterable[Any]:
    yield paragraph.paragraph_format
    seen: set[str] = set()
    for style in _style_chain(paragraph.style):
        seen.add(str(getattr(style, "style_id", "") or ""))
        yield style.paragraph_format
    normal = _normal_style(paragraph)
    if normal is not None and str(getattr(normal, "style_id", "") or "") not in seen:
        for style in _style_chain(normal):
            yield style.paragraph_format


def _effective_paragraph_value(paragraph: Paragraph, attribute: str) -> Any:
    for formatting in _paragraph_format_sources(paragraph):
        value = getattr(formatting, attribute, None)
        if value is not None:
            return value
    return None


def _font_name(font: Any) -> str:
    direct = str(getattr(font, "name", "") or "").strip()
    if direct:
        return direct
    element = getattr(font, "_element", None)
    properties = getattr(element, "rPr", None)
    fonts = getattr(properties, "rFonts", None)
    if fonts is None:
        return ""
    for attribute in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
        value = str(fonts.get(qn(attribute)) or "").strip()
        if value and not value.startswith("+"):
            return value
    return ""


def _run_font_sources(run: Run) -> Iterable[Any]:
    yield run.font
    try:
        for style in _style_chain(run.style):
            yield style.font
    except (AttributeError, KeyError):
        pass
    paragraph = run._parent
    if isinstance(paragraph, Paragraph):
        seen: set[str] = set()
        for style in _style_chain(paragraph.style):
            seen.add(str(getattr(style, "style_id", "") or ""))
            yield style.font
        normal = _normal_style(paragraph)
        if normal is not None and str(getattr(normal, "style_id", "") or "") not in seen:
            for style in _style_chain(normal):
                yield style.font


def _effective_font_value(run: Run, attribute: str) -> Any:
    for font in _run_font_sources(run):
        value = _font_name(font) if attribute == "name" else getattr(font, attribute, None)
        if attribute == "color" and value is not None and getattr(value, "rgb", None) is None:
            continue
        if value is not None and value != "":
            return value
    return None


def _iter_blocks(parent: _Document | _Cell) -> Iterable[Paragraph | Table]:
    element = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _paragraph_style(paragraph: Paragraph) -> str:
    try:
        return str(paragraph.style.name or "")
    except (AttributeError, KeyError):
        return ""


def _heading_level(paragraph: Paragraph, style_name: str) -> int | None:
    style_id = ""
    try:
        style_id = str(paragraph.style.style_id or "")
    except (AttributeError, KeyError):
        pass
    for value in (style_id, style_name):
        match = re.search(r"(?:heading|title|\u6807\u9898)\s*([1-6])", value, flags=re.I)
        if match:
            return int(match.group(1))
    return None


def _paragraph_attrs(paragraph: Paragraph, style_name: str) -> dict[str, Any]:
    alignment = _effective_paragraph_value(paragraph, "alignment")
    attrs: dict[str, Any] = {
        "docxStyle": style_name or None,
        "textAlign": _ALIGN_TO_WEB.get(alignment),
        "lineHeight": None,
        "spaceBefore": _points(_effective_paragraph_value(paragraph, "space_before")),
        "spaceAfter": _points(_effective_paragraph_value(paragraph, "space_after")),
        "firstLineIndent": _points(_effective_paragraph_value(paragraph, "first_line_indent")),
        "leftIndent": _points(_effective_paragraph_value(paragraph, "left_indent")),
        "rightIndent": _points(_effective_paragraph_value(paragraph, "right_indent")),
        "keepWithNext": _effective_paragraph_value(paragraph, "keep_with_next"),
        "keepTogether": _effective_paragraph_value(paragraph, "keep_together"),
        "pageBreakBefore": _effective_paragraph_value(paragraph, "page_break_before"),
        "widowControl": _effective_paragraph_value(paragraph, "widow_control"),
        "paragraphMark": _paragraph_mark_attrs(paragraph),
        "docxDrawingIds": [_drawing_id(element) for element in paragraph._p.xpath(".//w:drawing")],
        "tabStops": _paragraph_tab_stops(paragraph),
    }
    spacing = _effective_paragraph_value(paragraph, "line_spacing")
    if isinstance(spacing, (int, float)):
        attrs["lineHeight"] = str(round(float(spacing), 2))
    else:
        points = _points(spacing)
        if points:
            attrs["lineHeight"] = f"{points:g}pt"
    return attrs


def _paragraph_mark_attrs(paragraph: Paragraph) -> dict[str, Any] | None:
    properties = paragraph._p.pPr
    run_properties = properties.find(qn("w:rPr")) if properties is not None else None
    if run_properties is None:
        return None
    attrs: dict[str, Any] = {}
    fonts = run_properties.find(qn("w:rFonts"))
    if fonts is not None:
        for attribute in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
            value = str(fonts.get(qn(attribute)) or "").strip()
            if value and not value.startswith("+"):
                attrs["fontFamily"] = value
                break
        hint = str(fonts.get(qn("w:hint")) or "").strip()
        if hint:
            attrs["fontHint"] = hint
        cs_font = str(fonts.get(qn("w:cs")) or "").strip()
        if cs_font:
            attrs["csFontFamily"] = cs_font
    size = run_properties.find(qn("w:sz"))
    if size is not None:
        try:
            attrs["fontSize"] = round(float(size.get(qn("w:val"))) / 2, 2)
        except (TypeError, ValueError):
            pass
    color = run_properties.find(qn("w:color"))
    normalized_color = _normalize_hex(color.get(qn("w:val")) if color is not None else "")
    if normalized_color:
        attrs["color"] = f"#{normalized_color.lower()}"
    for key, tag in (("bold", "w:b"), ("italic", "w:i"), ("strike", "w:strike")):
        element = run_properties.find(qn(tag))
        if element is not None:
            value = str(element.get(qn("w:val")) or "true").lower()
            attrs[key] = value not in {"0", "false", "off", "none"}
    underline = run_properties.find(qn("w:u"))
    if underline is not None:
        attrs["underline"] = str(underline.get(qn("w:val")) or "single").lower() not in {"0", "false", "off", "none"}
    for key, tag in (("characterSpacing", "w:spacing"), ("position", "w:position")):
        element = run_properties.find(qn(tag))
        if element is not None:
            try:
                attrs[key] = int(element.get(qn("w:val")))
            except (TypeError, ValueError):
                pass
    return attrs or None


def _paragraph_tab_stops(paragraph: Paragraph) -> list[dict[str, Any]] | None:
    properties = paragraph._p.pPr
    tabs = properties.find(qn("w:tabs")) if properties is not None else None
    result: list[dict[str, Any]] = []
    if tabs is not None:
        for tab in tabs.findall(qn("w:tab")):
            try:
                position = int(tab.get(qn("w:pos")))
            except (TypeError, ValueError):
                continue
            result.append({
                "positionTwips": position,
                "alignment": str(tab.get(qn("w:val")) or "left"),
                "leader": str(tab.get(qn("w:leader")) or "none"),
            })
    return result or None


def _run_marks(run: Run, hyperlink: str = "") -> list[dict[str, Any]]:
    marks: list[dict[str, Any]] = []
    if _effective_font_value(run, "bold"):
        marks.append({"type": "bold"})
    if _effective_font_value(run, "italic"):
        marks.append({"type": "italic"})
    if _effective_font_value(run, "underline"):
        marks.append({"type": "underline"})
    if _effective_font_value(run, "strike"):
        marks.append({"type": "strike"})
    if _effective_font_value(run, "superscript"):
        marks.append({"type": "superscript"})
    if _effective_font_value(run, "subscript"):
        marks.append({"type": "subscript"})
    text_style: dict[str, Any] = {}
    font_name = _effective_font_value(run, "name")
    font_size = _effective_font_value(run, "size")
    font_color = _effective_font_value(run, "color")
    if font_name:
        text_style["fontFamily"] = str(font_name)
    if font_size:
        text_style["fontSize"] = f"{font_size.pt:g}pt"
    if font_color and font_color.rgb:
        text_style["color"] = f"#{str(font_color.rgb).lower()}"
    properties = run._r.rPr
    fonts = properties.find(qn("w:rFonts")) if properties is not None else None
    font_hint = str(fonts.get(qn("w:hint")) or "").strip() if fonts is not None else ""
    if font_hint:
        text_style["docxFontHint"] = font_hint
    for key, tag in (("docxCharacterSpacing", "w:spacing"), ("docxPosition", "w:position")):
        element = properties.find(qn(tag)) if properties is not None else None
        if element is not None:
            try:
                text_style[key] = int(element.get(qn("w:val")))
            except (TypeError, ValueError):
                pass
    if text_style:
        marks.append({"type": "textStyle", "attrs": text_style})
    highlight = _HIGHLIGHT_TO_HEX.get(run.font.highlight_color)
    if highlight:
        marks.append({"type": "highlight", "attrs": {"color": highlight}})
    if hyperlink:
        marks.append({"type": "link", "attrs": {"href": hyperlink, "target": "_blank"}})
    return marks


def _image_node(blip: Any, drawing: Any, paragraph: Paragraph) -> dict[str, Any] | None:
    relation_id = blip.get(qn("r:embed"))
    if not relation_id:
        return None
    part = paragraph.part.related_parts.get(relation_id)
    if not part or not hasattr(part, "blob"):
        return None
    blob = bytes(part.blob)
    if not blob or len(blob) > MAX_EDITOR_IMAGE_BYTES:
        return None
    mime = str(getattr(part, "content_type", "") or "image/png")
    extent = next(iter(drawing.xpath(".//wp:extent")), None)
    width = height = None
    if extent is not None:
        try:
            width = round(int(extent.get("cx")) / 914400 * 96)
            height = round(int(extent.get("cy")) / 914400 * 96)
        except (TypeError, ValueError):
            width = height = None
    return {
        "type": "image",
        "attrs": {
            "src": f"data:{mime};base64,{base64.b64encode(blob).decode('ascii')}",
            "alt": "\u62a5\u544a\u56fe\u7247",
            "title": None,
            "width": width,
            "height": height,
            "docxDrawingId": _drawing_id(drawing),
        },
    }


def _drawing_id(element: Any) -> str:
    return hashlib.sha256(element.xml.encode("utf-8")).hexdigest()


def _drawing_parent_run(element: Any) -> Any | None:
    parent = element.getparent()
    while parent is not None and parent.tag != qn("w:p"):
        if parent.tag == qn("w:r"):
            return parent
        parent = parent.getparent()
    return None


def _run_nodes(run_element: Any, paragraph: Paragraph, hyperlink: str = "") -> list[dict[str, Any]]:
    run = Run(run_element, paragraph)
    marks = _run_marks(run, hyperlink)
    nodes: list[dict[str, Any]] = []
    for child in run_element.iterchildren():
        if child.tag in {qn("w:t"), qn("w:instrText")}:
            text = child.text or ""
            if text:
                node: dict[str, Any] = {"type": "text", "text": text}
                if marks:
                    node["marks"] = marks
                nodes.append(node)
        elif child.tag == qn("w:tab"):
            node = {"type": "text", "text": "\t"}
            if marks:
                node["marks"] = marks
            nodes.append(node)
        elif child.tag in {qn("w:br"), qn("w:cr")}:
            if child.get(qn("w:type")) == "page":
                nodes.append({"type": "pageBreak"})
            else:
                nodes.append({"type": "hardBreak"})
        elif child.tag in {qn("w:drawing"), qn("w:pict")}:
            for blip in child.xpath(".//a:blip"):
                image = _image_node(blip, child, paragraph)
                if image:
                    nodes.append(image)
    return nodes


def _paragraph_inline_nodes(paragraph: Paragraph) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:r"):
            nodes.extend(_run_nodes(child, paragraph))
        elif child.tag == qn("w:hyperlink"):
            relation_id = child.get(qn("r:id"))
            href = ""
            if relation_id:
                relation = paragraph.part.rels.get(relation_id)
                href = str(getattr(relation, "target_ref", "") or "")
            for run_element in child.iterchildren(qn("w:r")):
                nodes.extend(_run_nodes(run_element, paragraph, href))
        elif child.tag in {qn("w:smartTag"), qn("w:ins"), qn("w:sdt"), qn("w:fldSimple")}:
            for run_element in child.iter(qn("w:r")):
                nodes.extend(_run_nodes(run_element, paragraph))
    return nodes


def _paragraph_node(paragraph: Paragraph) -> dict[str, Any]:
    style_name = _paragraph_style(paragraph)
    heading_level = _heading_level(paragraph, style_name)
    node: dict[str, Any] = {
        "type": "heading" if heading_level else "paragraph",
        "attrs": _paragraph_attrs(paragraph, style_name),
    }
    if heading_level:
        node["attrs"]["level"] = heading_level
    content = _paragraph_inline_nodes(paragraph)
    if content:
        node["content"] = content
    return node


def _cell_background(cell: _Cell) -> str | None:
    shading = cell._tc.tcPr.find(qn("w:shd")) if cell._tc.tcPr is not None else None
    fill = str(shading.get(qn("w:fill")) or "") if shading is not None else ""
    if re.fullmatch(r"[0-9A-Fa-f]{6}", fill):
        return f"#{fill.lower()}"
    return None


def _table_node(table: Table) -> dict[str, Any]:
    rows: list[dict[str, Any]] = []
    for row_index, row in enumerate(table.rows):
        cells: list[dict[str, Any]] = []
        for cell in row.cells:
            content = [_block_node(block) for block in _iter_blocks(cell)]
            content = [item for item in content if item]
            if not content:
                content = [{"type": "paragraph", "attrs": _paragraph_attrs(cell.paragraphs[0], "")}]
            cells.append({
                "type": "tableHeader" if row_index == 0 else "tableCell",
                "attrs": {
                    "colspan": 1,
                    "rowspan": 1,
                    "colwidth": None,
                    "backgroundColor": _cell_background(cell),
                },
                "content": content,
            })
        rows.append({"type": "tableRow", "content": cells})
    style_name = ""
    try:
        style_name = str(table.style.name or "")
    except (AttributeError, KeyError):
        pass
    return {"type": "table", "attrs": {"docxStyle": style_name or None}, "content": rows}


def _block_node(block: Paragraph | Table) -> dict[str, Any]:
    return _paragraph_node(block) if isinstance(block, Paragraph) else _table_node(block)


def load_docx_for_editor(path: Path) -> dict[str, Any]:
    path = Path(path)
    document = Document(path)
    content = [_block_node(block) for block in _iter_blocks(document)]
    if not content:
        content = [{"type": "paragraph", "attrs": {"docxStyle": "Normal"}}]
    section = document.sections[0]
    page = {
        "widthIn": _inches(section.page_width) or 8.27,
        "heightIn": _inches(section.page_height) or 11.69,
        "topMarginIn": _inches(section.top_margin) or 0.8,
        "rightMarginIn": _inches(section.right_margin) or 0.8,
        "bottomMarginIn": _inches(section.bottom_margin) or 0.8,
        "leftMarginIn": _inches(section.left_margin) or 0.8,
    }
    stat = path.stat()
    return {
        "name": path.name,
        "sourceSha256": sha256_file(path),
        "sourceMtimeNs": stat.st_mtime_ns,
        "page": page,
        "document": {"type": "doc", "content": content},
    }


def _validate_editor_document(document: Any) -> dict[str, Any]:
    if not isinstance(document, dict) or document.get("type") != "doc":
        raise ValueError("\u7f16\u8f91\u5668\u5185\u5bb9\u4e0d\u662f\u6709\u6548\u6587\u6863")
    node_count = 0
    text_count = 0
    stack: list[Any] = [document]
    while stack:
        node = stack.pop()
        if not isinstance(node, dict):
            raise ValueError("\u7f16\u8f91\u5668\u5185\u5bb9\u7ed3\u6784\u65e0\u6548")
        node_count += 1
        if node_count > MAX_EDITOR_NODES:
            raise ValueError("\u7f16\u8f91\u5185\u5bb9\u8fc7\u4e8e\u590d\u6742")
        text = node.get("text")
        if text is not None:
            if not isinstance(text, str):
                raise ValueError("\u7f16\u8f91\u5668\u6587\u672c\u7ed3\u6784\u65e0\u6548")
            text_count += len(text)
            if text_count > MAX_EDITOR_TEXT_CHARS:
                raise ValueError("\u7f16\u8f91\u5185\u5bb9\u8d85\u8fc7\u53ef\u4fdd\u5b58\u7684\u6587\u5b57\u4e0a\u9650")
        content = node.get("content", [])
        if content is not None:
            if not isinstance(content, list):
                raise ValueError("\u7f16\u8f91\u5668\u5185\u5bb9\u5c42\u7ea7\u65e0\u6548")
            stack.extend(reversed(content))
        if node.get("type") == "image":
            src = str((node.get("attrs") or {}).get("src") or "")
            if len(src) > MAX_EDITOR_IMAGE_BYTES * 2:
                raise ValueError("\u63d2\u5165\u7684\u56fe\u7247\u8fc7\u5927")
    return document


def _style_exists(document: _Document, style_name: str) -> bool:
    if not style_name:
        return False
    try:
        document.styles[style_name]
        return True
    except KeyError:
        return False


def _length_points(value: Any, minimum: float = -360, maximum: float = 720) -> Pt | None:
    if value in {None, ""}:
        return None
    try:
        numeric = max(minimum, min(maximum, float(value)))
    except (TypeError, ValueError):
        return None
    return Pt(numeric)


def _apply_paragraph_attrs(paragraph: Paragraph, attrs: dict[str, Any]) -> None:
    alignment = _WEB_TO_ALIGN.get(str(attrs.get("textAlign") or ""))
    if alignment is not None:
        paragraph.alignment = alignment
    formatting = paragraph.paragraph_format
    line_height = str(attrs.get("lineHeight") or "").strip().lower()
    if line_height.endswith("pt"):
        try:
            formatting.line_spacing = Pt(max(6, min(144, float(line_height[:-2]))))
        except ValueError:
            pass
    elif line_height:
        try:
            formatting.line_spacing = max(0.8, min(4, float(line_height)))
        except ValueError:
            pass
    for key, field in (
        ("spaceBefore", "space_before"),
        ("spaceAfter", "space_after"),
        ("firstLineIndent", "first_line_indent"),
        ("leftIndent", "left_indent"),
        ("rightIndent", "right_indent"),
    ):
        length = _length_points(attrs.get(key))
        if length is not None:
            setattr(formatting, field, length)
    for key, field in (
        ("keepWithNext", "keep_with_next"),
        ("keepTogether", "keep_together"),
        ("pageBreakBefore", "page_break_before"),
        ("widowControl", "widow_control"),
    ):
        if key in attrs and attrs[key] is not None:
            setattr(formatting, field, bool(attrs[key]))
    tab_stops = attrs.get("tabStops")
    if isinstance(tab_stops, list) and tab_stops:
        properties = paragraph._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        for item in tab_stops[:32]:
            if not isinstance(item, dict):
                continue
            try:
                position = max(0, min(31680, int(item.get("positionTwips"))))
            except (TypeError, ValueError):
                continue
            tab = OxmlElement("w:tab")
            tab.set(qn("w:pos"), str(position))
            alignment = str(item.get("alignment") or "left")
            if alignment not in {"left", "center", "right", "decimal", "bar", "clear", "num"}:
                alignment = "left"
            tab.set(qn("w:val"), alignment)
            leader = str(item.get("leader") or "none")
            if leader not in {"none", "dot", "hyphen", "underscore", "heavy", "middleDot"}:
                leader = "none"
            if leader != "none":
                tab.set(qn("w:leader"), leader)
            tabs.append(tab)
        if len(tabs):
            run_properties = properties.find(qn("w:rPr"))
            properties.insert(properties.index(run_properties) if run_properties is not None else len(properties), tabs)
    mark = attrs.get("paragraphMark")
    if isinstance(mark, dict):
        properties = paragraph._p.get_or_add_pPr()
        run_properties = properties.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            properties.append(run_properties)
        font_family = re.sub(r"[\r\n\x00]", "", str(mark.get("fontFamily") or "")).strip()[:80]
        if font_family:
            fonts = OxmlElement("w:rFonts")
            for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                fonts.set(qn(attribute), font_family)
            if mark.get("csFontFamily"):
                fonts.set(qn("w:cs"), str(mark["csFontFamily"])[:80])
            if mark.get("fontHint"):
                fonts.set(qn("w:hint"), str(mark["fontHint"])[:20])
            run_properties.append(fonts)
        try:
            font_size = max(6, min(96, float(mark.get("fontSize"))))
        except (TypeError, ValueError):
            font_size = 0
        if font_size:
            for tag in ("w:sz", "w:szCs"):
                size = OxmlElement(tag)
                size.set(qn("w:val"), str(round(font_size * 2)))
                run_properties.append(size)
        color = _normalize_hex(mark.get("color"))
        if color:
            element = OxmlElement("w:color")
            element.set(qn("w:val"), color)
            run_properties.append(element)
        for key, tag in (("characterSpacing", "w:spacing"), ("position", "w:position")):
            if key not in mark:
                continue
            try:
                value = max(-31680, min(31680, int(mark[key])))
            except (TypeError, ValueError):
                continue
            element = OxmlElement(tag)
            element.set(qn("w:val"), str(value))
            run_properties.append(element)
        for key, tag in (("bold", "w:b"), ("italic", "w:i"), ("strike", "w:strike"), ("underline", "w:u")):
            if key not in mark:
                continue
            element = OxmlElement(tag)
            if not bool(mark[key]):
                element.set(qn("w:val"), "0")
            elif key == "underline":
                element.set(qn("w:val"), "single")
            run_properties.append(element)


def _normalize_hex(value: Any) -> str:
    match = re.fullmatch(r"#?([0-9a-fA-F]{6})", str(value or "").strip())
    return match.group(1).upper() if match else ""


def _marks_by_type(node: dict[str, Any]) -> dict[str, dict[str, Any]]:
    result: dict[str, dict[str, Any]] = {}
    for mark in node.get("marks") or []:
        if isinstance(mark, dict) and mark.get("type"):
            result[str(mark["type"])] = mark.get("attrs") if isinstance(mark.get("attrs"), dict) else {}
    return result


def _apply_run_marks(run: Run, marks: dict[str, dict[str, Any]]) -> None:
    run.bold = "bold" in marks
    run.italic = "italic" in marks
    run.underline = "underline" in marks
    run.font.strike = "strike" in marks
    run.font.superscript = "superscript" in marks
    run.font.subscript = "subscript" in marks
    text_style = marks.get("textStyle", {})
    font_family = re.sub(r"[\r\n\x00]", "", str(text_style.get("fontFamily") or "")).strip()[:80]
    if font_family:
        run.font.name = font_family
        fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            fonts.set(qn(attribute), font_family)
    else:
        fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    font_hint = str(text_style.get("docxFontHint") or "")
    if font_hint in {"default", "eastAsia", "cs"}:
        fonts.set(qn("w:hint"), font_hint)
    font_size = str(text_style.get("fontSize") or "").strip().lower()
    size_match = re.fullmatch(r"(\d+(?:\.\d+)?)(pt|px)?", font_size)
    if size_match:
        size = float(size_match.group(1)) * (0.75 if size_match.group(2) == "px" else 1)
        run.font.size = Pt(max(6, min(96, size)))
    color = _normalize_hex(text_style.get("color"))
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    properties = run._r.get_or_add_rPr()
    for key, tag in (("docxCharacterSpacing", "w:spacing"), ("docxPosition", "w:position")):
        if key not in text_style:
            continue
        try:
            value = max(-31680, min(31680, int(text_style[key])))
        except (TypeError, ValueError):
            continue
        element = OxmlElement(tag)
        element.set(qn("w:val"), str(value))
        properties.append(element)
    highlight = _normalize_hex((marks.get("highlight") or {}).get("color"))
    background = _normalize_hex(text_style.get("backgroundColor"))
    selected = f"#{(highlight or background).lower()}" if highlight or background else ""
    if selected in _HEX_TO_HIGHLIGHT:
        run.font.highlight_color = _HEX_TO_HIGHLIGHT[selected]


def _add_hyperlink_run(paragraph: Paragraph, text: str, href: str) -> Run:
    relationship_id = paragraph.part.relate_to(href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    return Run(run_element, paragraph)


def _decode_image(src: str) -> BytesIO:
    match = re.fullmatch(r"data:image/[a-zA-Z0-9.+-]+;base64,([A-Za-z0-9+/=\r\n]+)", src)
    if not match:
        raise ValueError("\u53ea\u5141\u8bb8\u4fdd\u5b58\u7f16\u8f91\u5668\u4e2d\u7684\u672c\u5730\u56fe\u7247")
    try:
        blob = base64.b64decode(match.group(1), validate=True)
    except (ValueError, binascii.Error) as exc:
        raise ValueError("\u56fe\u7247\u5185\u5bb9\u65e0\u6cd5\u89e3\u6790") from exc
    if not blob or len(blob) > MAX_EDITOR_IMAGE_BYTES:
        raise ValueError("\u63d2\u5165\u7684\u56fe\u7247\u8fc7\u5927")
    return BytesIO(blob)


def _render_inline(paragraph: Paragraph, content: list[dict[str, Any]], source_drawings: dict[str, Any]) -> set[str]:
    restored_drawings: set[str] = set()
    for node in content:
        node_type = str(node.get("type") or "")
        if node_type == "text":
            text = str(node.get("text") or "")
            if not text:
                continue
            marks = _marks_by_type(node)
            link = str((marks.get("link") or {}).get("href") or "").strip()
            for index, segment in enumerate(text.split("\n")):
                if index:
                    paragraph.add_run().add_break()
                if not segment:
                    continue
                run = _add_hyperlink_run(paragraph, segment, link) if re.match(r"^https?://", link, re.I) else paragraph.add_run(segment)
                _apply_run_marks(run, marks)
        elif node_type == "hardBreak":
            paragraph.add_run().add_break()
        elif node_type == "pageBreak":
            paragraph.add_run().add_break(WD_BREAK.PAGE)
        elif node_type == "image":
            attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
            drawing_id = str(attrs.get("docxDrawingId") or "")
            if drawing_id in source_drawings:
                paragraph._p.append(deepcopy(source_drawings[drawing_id]))
                restored_drawings.add(drawing_id)
                continue
            width_value = attrs.get("width")
            height_value = attrs.get("height")
            width = height = None
            try:
                if width_value:
                    width = Inches(max(0.2, min(7.5, float(width_value) / 96)))
                if height_value:
                    height = Inches(max(0.2, min(10.0, float(height_value) / 96)))
            except (TypeError, ValueError):
                width = height = None
            paragraph.add_run().add_picture(_decode_image(str(attrs.get("src") or "")), width=width, height=height)
    return restored_drawings


def _new_paragraph(
    container: _Document | _Cell,
    node: dict[str, Any],
    source_drawings: dict[str, Any],
    *,
    list_style: str = "",
) -> Paragraph:
    paragraph = container.add_paragraph()
    attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
    document = container if isinstance(container, _Document) else container.part.document
    style_name = str(attrs.get("docxStyle") or "")
    node_type = str(node.get("type") or "")
    if list_style and _style_exists(document, list_style):
        paragraph.style = list_style
    elif style_name and _style_exists(document, style_name):
        paragraph.style = style_name
    elif node_type == "heading":
        level = max(1, min(6, int(attrs.get("level") or 1)))
        heading_style = f"Heading {level}"
        if _style_exists(document, heading_style):
            paragraph.style = heading_style
    elif node_type == "blockquote" and _style_exists(document, "Quote"):
        paragraph.style = "Quote"
    elif node_type == "codeBlock" and _style_exists(document, "No Spacing"):
        paragraph.style = "No Spacing"
    _apply_paragraph_attrs(paragraph, attrs)
    restored = _render_inline(paragraph, [item for item in node.get("content") or [] if isinstance(item, dict)], source_drawings)
    for drawing_id in attrs.get("docxDrawingIds") or []:
        drawing_id = str(drawing_id or "")
        if drawing_id in source_drawings and drawing_id not in restored:
            paragraph._p.append(deepcopy(source_drawings[drawing_id]))
    if node_type == "codeBlock":
        for run in paragraph.runs:
            run.font.name = "Menlo"
            run.font.size = run.font.size or Pt(10)
    return paragraph


def _render_list(container: _Document | _Cell, node: dict[str, Any], source_drawings: dict[str, Any], depth: int = 0) -> None:
    ordered = node.get("type") == "orderedList"
    style = "List Number" if ordered else "List Bullet"
    nested_style = f"{style} {min(depth + 2, 3)}" if depth else style
    for item in node.get("content") or []:
        if not isinstance(item, dict):
            continue
        blocks = [block for block in item.get("content") or [] if isinstance(block, dict)]
        wrote = False
        for block in blocks:
            if block.get("type") in {"bulletList", "orderedList"}:
                _render_list(container, block, source_drawings, depth + 1)
            elif block.get("type") in {"paragraph", "heading", "blockquote", "codeBlock"}:
                _new_paragraph(container, block, source_drawings, list_style=nested_style if not wrote else "")
                wrote = True


def _shade_cell(cell: _Cell, color: str) -> None:
    normalized = _normalize_hex(color)
    if not normalized:
        return
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), normalized)


def _render_table(document: _Document, node: dict[str, Any], source_drawings: dict[str, Any]) -> None:
    rows = [item for item in node.get("content") or [] if isinstance(item, dict)]
    columns = max((len(row.get("content") or []) for row in rows), default=1)
    table = document.add_table(rows=max(1, len(rows)), cols=max(1, columns))
    attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
    style_name = str(attrs.get("docxStyle") or "")
    if style_name and _style_exists(document, style_name):
        table.style = style_name
    elif _style_exists(document, "Table Grid"):
        table.style = "Table Grid"
    for row_index, row_node in enumerate(rows):
        for column_index, cell_node in enumerate(row_node.get("content") or []):
            if column_index >= len(table.rows[row_index].cells) or not isinstance(cell_node, dict):
                continue
            cell = table.rows[row_index].cells[column_index]
            cell._tc.clear_content()
            cell_attrs = cell_node.get("attrs") if isinstance(cell_node.get("attrs"), dict) else {}
            _shade_cell(cell, str(cell_attrs.get("backgroundColor") or ("d9eaf7" if row_index == 0 else "")))
            _render_blocks(cell, [item for item in cell_node.get("content") or [] if isinstance(item, dict)], source_drawings)
            if not cell.paragraphs:
                cell.add_paragraph()
            if cell_node.get("type") == "tableHeader":
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _add_horizontal_rule(container: _Document | _Cell) -> None:
    paragraph = container.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7C9D6")
    borders.append(bottom)
    properties.append(borders)


def _render_blocks(container: _Document | _Cell, content: list[dict[str, Any]], source_drawings: dict[str, Any]) -> None:
    for node in content:
        node_type = str(node.get("type") or "")
        if node_type in {"paragraph", "heading", "blockquote", "codeBlock"}:
            _new_paragraph(container, node, source_drawings)
        elif node_type in {"bulletList", "orderedList"}:
            _render_list(container, node, source_drawings)
        elif node_type == "horizontalRule":
            _add_horizontal_rule(container)
        elif node_type == "table" and isinstance(container, _Document):
            _render_table(container, node, source_drawings)
        elif node_type == "image":
            paragraph = container.add_paragraph()
            _render_inline(paragraph, [node], source_drawings)


def save_editor_document(source_path: Path, target_path: Path, document_payload: Any) -> dict[str, Any]:
    source_path = Path(source_path)
    target_path = Path(target_path)
    document_json = _validate_editor_document(document_payload)
    pending = target_path.with_name(f".{target_path.name}.{uuid.uuid4().hex}.tmp")
    shutil.copy2(source_path, pending)
    try:
        document = Document(pending)
        source_drawings: dict[str, Any] = {}
        for drawing in document._element.body.iter(qn("w:drawing")):
            parent_run = _drawing_parent_run(drawing)
            if parent_run is not None:
                source_drawings[_drawing_id(drawing)] = deepcopy(parent_run)
        body = document._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        _render_blocks(document, [item for item in document_json.get("content") or [] if isinstance(item, dict)], source_drawings)
        if not document.paragraphs and not document.tables:
            document.add_paragraph()
        document.core_properties.modified = datetime.now()
        document.save(pending)
        pending.replace(target_path)
    finally:
        pending.unlink(missing_ok=True)
    return {
        "path": target_path,
        "sha256": sha256_file(target_path),
        "size": target_path.stat().st_size,
        "mtimeNs": target_path.stat().st_mtime_ns,
    }
