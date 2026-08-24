from __future__ import annotations

import json
import unittest
from tempfile import TemporaryDirectory
from pathlib import Path
from unittest import mock

from docx import Document

import generate_carrier_performance_report as report


def sample_sections() -> list[dict]:
    return [
        {
            "company": "测试运营商",
            "title": "测试运营商关键摘要",
            "items": [
                "派息：全年每股派息0.50港元，同比增长5%。",
                "资本开支：2025年资本开支100亿元。",
                "战略升级：围绕AI和云服务升级，收入达到200亿元；产品目录和重复信息较多，需要压缩。",
                "券商观点：机构维持买入评级，目标价10港元。",
                "市场反应：股价由8港元升至9港元，上涨12.5%。",
            ],
        }
    ]


class CarrierPerformanceAiEditorTests(unittest.TestCase):
    def test_editor_requests_are_split_into_small_parallel_batches(self) -> None:
        packs = [
            {
                "company": f"测试运营商{index}",
                "title": "测试",
                "evidence": {},
                "web_research": {},
            }
            for index in range(5)
        ]
        batch_sizes = []

        def ai_client(batch):
            batch_sizes.append(len(batch))
            return {
                "companies": [
                    {"company": pack["company"], "fields": {}} for pack in batch
                ]
            }, "test-model"

        returned, model, failed = report.call_performance_editor_batches(
            packs,
            ai_client=ai_client,
            progress=lambda _message: None,
        )

        self.assertEqual(sorted(batch_sizes), [1, 2, 2])
        self.assertEqual(len(returned), 5)
        self.assertEqual(model, "test-model")
        self.assertEqual(failed, set())

    def test_rewrite_keeps_structure_and_accepts_grounded_fields(self) -> None:
        response = {
            "companies": [
                {
                    "company": "测试运营商",
                    "fields": {
                        "dividend": "全年每股派息0.50港元，同比增长5%。",
                        "capex": "2025年资本开支100亿元。",
                        "strategy": "公司围绕AI和云服务推进升级，相关收入达到200亿元。",
                        "broker": "机构维持买入评级，目标价10港元。",
                        "market": "股价由8港元升至9港元，上涨12.5%。",
                    },
                }
            ]
        }
        with TemporaryDirectory() as temp_dir, mock.patch.object(
            report, "PERFORMANCE_AI_AUDIT_PATH", Path(temp_dir) / "audit.json"
        ):
            rewritten = report.rewrite_performance_sections_with_ai(
                sample_sections(), ai_client=lambda _packs: (response, "test-model"), progress=lambda _message: None
            )

        self.assertEqual(len(rewritten), 1)
        self.assertEqual(len(rewritten[0]["items"]), 5)
        self.assertIn("公司围绕AI和云服务推进升级", rewritten[0]["items"][2])

    def test_invented_number_falls_back_to_locked_evidence(self) -> None:
        response = {
            "companies": [
                {
                    "company": "测试运营商",
                    "fields": {
                        "dividend": "全年每股派息0.80港元，同比增长8%。",
                    },
                }
            ]
        }
        with TemporaryDirectory() as temp_dir, mock.patch.object(
            report, "PERFORMANCE_AI_AUDIT_PATH", Path(temp_dir) / "audit.json"
        ):
            rewritten = report.rewrite_performance_sections_with_ai(
                sample_sections(), ai_client=lambda _packs: (response, "test-model"), progress=lambda _message: None
            )

        self.assertEqual(rewritten[0]["items"][0], sample_sections()[0]["items"][0])

    def test_model_failure_preserves_existing_summary(self) -> None:
        def fail(_packs):
            raise RuntimeError("offline")

        limitations = []
        messages = []
        with TemporaryDirectory() as temp_dir, mock.patch.object(
            report, "PERFORMANCE_AI_AUDIT_PATH", Path(temp_dir) / "audit.json"
        ):
            rewritten = report.rewrite_performance_sections_with_ai(
                sample_sections(),
                ai_client=fail,
                progress=messages.append,
                limitations=limitations,
            )

        self.assertEqual(rewritten, sample_sections())
        self.assertTrue(any(item["stage"] == "ai_batch" for item in limitations))
        self.assertTrue(any("[业绩摘要局限][ai_batch]" in message for message in messages))

    def test_online_research_is_passed_to_ai_and_can_support_a_new_number(self) -> None:
        response = {
            "companies": [
                {
                    "company": "测试运营商",
                    "fields": {
                        "capex": "2026年资本开支计划为88亿元。",
                    },
                }
            ]
        }
        captured_packs = []

        def ai_client(packs):
            captured_packs.extend(packs)
            return response, "test-model"

        web_research = {
            "测试运营商": {
                "query": "测试运营商 2026 资本开支",
                "provider": "unit",
                "results": [
                    {
                        "title": "测试运营商公布2026年资本开支计划",
                        "url": "https://example.com/capex",
                        "snippet": "公司公布2026年资本开支计划为88亿元。",
                    }
                ],
                "error": "",
            }
        }
        with TemporaryDirectory() as temp_dir, mock.patch.object(
            report, "PERFORMANCE_AI_AUDIT_PATH", Path(temp_dir) / "audit.json"
        ):
            rewritten = report.rewrite_performance_sections_with_ai(
                sample_sections(),
                ai_client=ai_client,
                progress=lambda _message: None,
                web_research=web_research,
            )
            audit = (Path(temp_dir) / "audit.json").read_text(encoding="utf-8")

        self.assertEqual(captured_packs[0]["web_research"], web_research["测试运营商"])
        self.assertIn("88亿元", rewritten[0]["items"][1])
        self.assertIn("https://example.com/capex", audit)

    def test_company_research_records_limitations_and_continues_when_searches_are_empty(self) -> None:
        def empty_search(query, _limit):
            return {"query": query, "provider": "", "results": [], "error": "offline"}

        limitations = []
        messages = []
        researched = report.research_performance_companies_online(
            sample_sections(),
            search_client=empty_search,
            progress=messages.append,
            limitations=limitations,
        )

        self.assertFalse(researched["测试运营商"]["results"])
        self.assertEqual(limitations[0]["stage"], "web_research")
        self.assertIn("两轮公开网页搜索", limitations[0]["reason"])
        self.assertTrue(any("[业绩摘要局限][web_research]" in message for message in messages))

    def test_company_research_repairs_a_first_round_search_gap(self) -> None:
        calls = []

        def second_round_search(query, _limit):
            calls.append(query)
            if len(calls) == 1:
                return {"query": query, "provider": "", "results": [], "error": "no result"}
            return {
                "query": query,
                "provider": "unit",
                "results": [
                    {
                        "title": "测试运营商官方年报",
                        "url": "https://example.com/annual-report",
                        "snippet": "测试运营商公布最新年度业绩。",
                    }
                ],
                "error": "",
            }

        researched = report.research_performance_companies_online(
            sample_sections(),
            search_client=second_round_search,
            progress=lambda _message: None,
        )

        self.assertEqual(len(calls), 2)
        self.assertEqual(researched["测试运营商"]["provider"], "unit")

    def test_template_failure_still_creates_business_only_report_and_audit(self) -> None:
        model = report.fallback_performance_model([])
        model["generationLimitations"] = [
            report.performance_limitation_entry(
                "web_research",
                "offline",
                impact="没有新网页证据",
                action="保留确定性字段",
            )
        ]
        model["generationMode"] = "limited"
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            output_path = temp_path / "业绩摘要.docx"
            with (
                mock.patch.object(report, "ROOT", temp_path),
                mock.patch.object(report, "TEMPLATE_PATH", temp_path / "missing-template.docx"),
                mock.patch.object(report, "build_dynamic_model", return_value=model),
                mock.patch.object(report, "dated_output_path", return_value=output_path),
            ):
                rendered_path = report.render_report()

            self.assertEqual(rendered_path, output_path)
            self.assertTrue(output_path.exists())
            sidecar = report.performance_quality_sidecar_path(output_path)
            self.assertTrue(sidecar.exists())
            audit = json.loads(sidecar.read_text(encoding="utf-8"))
            document = Document(output_path)
            report_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            report_text += "\n" + "\n".join(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )

        self.assertEqual(audit["generationMode"], "limited")
        self.assertTrue(any(item["stage"] == "template_render" for item in audit["limitations"]))
        for forbidden in report.PERFORMANCE_FORBIDDEN_REPORT_PHRASES:
            self.assertNotIn(forbidden, report_text)

    def test_source_config_failure_returns_fallback_model_instead_of_raising(self) -> None:
        messages = []
        with (
            mock.patch.object(report, "refresh_feishu_mirror", return_value=None),
            mock.patch.object(report, "load_source_config", side_effect=ValueError("broken config")),
        ):
            model = report.build_dynamic_model(progress=messages.append)

        self.assertEqual(model["generationMode"], "limited")
        self.assertEqual(len(model["sections"]), len(report.DEFAULT_PERFORMANCE_COMPANIES))
        self.assertEqual(model["generationLimitations"][0]["stage"], "source_config")
        self.assertTrue(any("[业绩摘要局限][source_config]" in message for message in messages))


if __name__ == "__main__":
    unittest.main()
