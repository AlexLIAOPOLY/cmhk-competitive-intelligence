import base64
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest import mock

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from cmhk.reporting.docx_editor import load_docx_for_editor, save_editor_document, sha256_file
import web_app


ROOT = Path(__file__).resolve().parents[1]


class ReportEditorRoundTripTests(unittest.TestCase):
    def test_docx_round_trip_preserves_page_header_table_and_writes_edits(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "8月30日周报.docx"
            target = root / "8月30日周报（编辑稿）.docx"
            document = Document()
            section = document.sections[0]
            section.header.paragraphs[0].text = "CMHK 战略竞对中心"
            title = document.add_heading("战略双周报", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph = document.add_paragraph()
            original_run = paragraph.add_run("原始内容")
            original_run.bold = True
            original_run.font.name = "FangSong"
            original_fonts = original_run._r.get_or_add_rPr().get_or_add_rFonts()
            original_fonts.set(qn("w:hint"), "eastAsia")
            character_spacing = OxmlElement("w:spacing")
            character_spacing.set(qn("w:val"), "-20")
            original_run._r.get_or_add_rPr().append(character_spacing)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            table = document.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "指标"
            table.cell(0, 1).text = "结果"
            table.cell(1, 0).text = "收入"
            table.cell(1, 1).text = "100"
            picture = document.add_paragraph()
            picture.add_run().add_picture(BytesIO(base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M/wHwAEAQH/6p0qWQAAAABJRU5ErkJggg=="
            )))
            document.save(source)

            payload = load_docx_for_editor(source)
            payload["document"]["content"][1]["content"][0]["text"] = "页面编辑后的真实内容"
            result = save_editor_document(source, target, payload["document"])
            reopened = Document(target)

            self.assertEqual(reopened.sections[0].header.paragraphs[0].text, "CMHK 战略竞对中心")
            self.assertEqual(reopened.sections[0].page_width, section.page_width)
            self.assertIn("页面编辑后的真实内容", "\n".join(item.text for item in reopened.paragraphs))
            self.assertEqual(len(reopened.tables), 1)
            self.assertEqual(reopened.tables[0].cell(1, 1).text, "100")
            source_drawings = [item.xml for item in Document(source)._element.body.iter(qn("w:drawing"))]
            saved_drawings = [item.xml for item in reopened._element.body.iter(qn("w:drawing"))]
            self.assertEqual(saved_drawings, source_drawings)
            saved_text_run = reopened.paragraphs[1].runs[0]
            saved_fonts = saved_text_run._r.rPr.find(qn("w:rFonts"))
            self.assertEqual(saved_fonts.get(qn("w:hint")), "eastAsia")
            self.assertEqual(saved_text_run._r.rPr.find(qn("w:spacing")).get(qn("w:val")), "-20")
            saved_tab = reopened.paragraphs[1]._p.pPr.find(qn("w:tabs")).find(qn("w:tab"))
            self.assertEqual(saved_tab.get(qn("w:val")), "right")
            self.assertEqual(result["sha256"], sha256_file(target))
            self.assertNotEqual(sha256_file(source), sha256_file(target))

    def test_editor_payload_exposes_page_and_revision_hash(self):
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "业绩摘要.docx"
            document = Document()
            document.add_paragraph("摘要正文")
            document.save(source)

            payload = load_docx_for_editor(source)

            self.assertEqual(payload["name"], source.name)
            self.assertEqual(payload["sourceSha256"], sha256_file(source))
            self.assertGreater(payload["page"]["widthIn"], 7)
            self.assertEqual(payload["document"]["type"], "doc")

    def test_editor_resolves_inherited_word_font_size_and_spacing(self):
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "继承样式.docx"
            document = Document()
            body = document.styles["Body Text"]
            body.font.name = "FangSong"
            body.font.size = Pt(15.5)
            body.paragraph_format.line_spacing = 1.05
            body.paragraph_format.space_after = Pt(1)
            body.paragraph_format.keep_with_next = True
            body.paragraph_format.widow_control = True
            paragraph = document.add_paragraph("继承仿宋正文", style="Body Text")
            self.assertIsNone(paragraph.runs[0].font.name)
            blank = document.add_paragraph()
            mark = OxmlElement("w:rPr")
            mark_size = OxmlElement("w:sz")
            mark_size.set(qn("w:val"), "28")
            mark.append(mark_size)
            blank._p.get_or_add_pPr().append(mark)
            document.save(source)

            payload = load_docx_for_editor(source)
            block = payload["document"]["content"][0]
            marks = {item["type"]: item.get("attrs", {}) for item in block["content"][0]["marks"]}

            self.assertEqual(marks["textStyle"]["fontFamily"], "FangSong")
            self.assertEqual(marks["textStyle"]["fontSize"], "15.5pt")
            self.assertEqual(block["attrs"]["lineHeight"], "1.05")
            self.assertEqual(block["attrs"]["spaceAfter"], 1.0)
            self.assertTrue(block["attrs"]["keepWithNext"])
            self.assertTrue(block["attrs"]["widowControl"])
            self.assertEqual(payload["document"]["content"][1]["attrs"]["paragraphMark"]["fontSize"], 14.0)

            target = Path(folder) / "继承样式（编辑稿）.docx"
            save_editor_document(source, target, payload["document"])
            saved_run = Document(target).paragraphs[0].runs[0]
            saved_fonts = saved_run._r.get_or_add_rPr().get_or_add_rFonts()
            self.assertEqual(saved_fonts.get(qn("w:eastAsia")), "FangSong")
            self.assertTrue(Document(target).paragraphs[0].paragraph_format.keep_with_next)
            self.assertTrue(Document(target).paragraphs[0].paragraph_format.widow_control)
            saved_mark = Document(target).paragraphs[1]._p.pPr.find(qn("w:rPr")).find(qn("w:sz"))
            self.assertEqual(saved_mark.get(qn("w:val")), "28")


class ReportEditorServiceTests(unittest.TestCase):
    def test_generated_report_saves_as_edit_copy_then_updates_with_archive(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            (root / "data" / "reporting").mkdir(parents=True)
            source = root / "8月30日周报.docx"
            document = Document()
            document.add_paragraph("正式生成内容")
            document.save(source)

            def fake_preview(report_path):
                from cmhk.reporting.pdf_preview import pdf_preview_path

                preview = pdf_preview_path(report_path, root / "web" / "static" / "report-previews")
                preview.parent.mkdir(parents=True, exist_ok=True)
                preview.write_bytes(b"%PDF-test")
                return preview

            with (
                mock.patch.object(web_app, "ROOT", root),
                mock.patch.object(web_app, "REPORT_METADATA_PATH", root / "data" / "reporting" / "report_file_metadata.json"),
                mock.patch.object(web_app, "build_status", return_value={"outputs": []}),
                mock.patch.object(web_app, "delete_audio_for_report"),
                mock.patch.object(web_app, "report_audio_metadata", return_value={"exists": False}),
                mock.patch("cmhk.reporting.pdf_preview.convert_docx_to_pdf_preview", side_effect=fake_preview),
            ):
                opened = web_app.load_report_editor_payload(source.name)
                self.assertEqual(opened["previewUrl"], "")
                opened["document"]["content"][0]["content"][0]["text"] = "第一次页面编辑"
                first = web_app.save_report_editor_payload({
                    "path": source.name,
                    "sourceSha256": opened["sourceSha256"],
                    "saveMode": "update",
                    "document": opened["document"],
                }, actor={"display_name": "测试编辑者"})
                reopened = web_app.load_report_editor_payload(first["path"])
                self.assertIn("/static/report-previews/", reopened["previewUrl"])
                reopened["document"]["content"][0]["content"][0]["text"] = "第二次页面编辑"
                second = web_app.save_report_editor_payload({
                    "path": first["path"],
                    "sourceSha256": reopened["sourceSha256"],
                    "saveMode": "update",
                    "document": reopened["document"],
                }, actor={"display_name": "测试编辑者"})

            self.assertEqual(Document(source).paragraphs[0].text, "正式生成内容")
            self.assertEqual(first["path"], "8月30日周报（编辑稿）.docx")
            self.assertEqual(second["path"], first["path"])
            self.assertEqual(Document(root / second["path"]).paragraphs[0].text, "第二次页面编辑")
            metadata = web_app.json.loads((root / "data" / "reporting" / "report_file_metadata.json").read_text(encoding="utf-8"))
            self.assertTrue(metadata[first["path"]]["isEdited"])
            self.assertEqual(metadata[first["path"]]["editorRevision"], 2)
            self.assertEqual(metadata[first["path"]]["sourcePath"], source.name)
            self.assertEqual(len(list((root / "archives" / "report_edits").rglob("*.docx"))), 1)


class ReportEditorUiTests(unittest.TestCase):
    def test_word_like_editor_is_local_full_screen_and_available_from_both_report_surfaces(self):
        index = (ROOT / "web/static/index.html").read_text(encoding="utf-8")
        app = (ROOT / "web/static/app.js").read_text(encoding="utf-8")
        workspace = (ROOT / "web/static/workspace-tabs.js").read_text(encoding="utf-8")
        source = (ROOT / "web/static/report-editor-source.js").read_text(encoding="utf-8")
        style = (ROOT / "web/static/report-editor.css").read_text(encoding="utf-8")
        bundle = ROOT / "web/static/vendor/tiptap-report-editor-3.30.5.min.js"
        bundle_source = bundle.read_text(encoding="utf-8")

        self.assertIn('id="reportEditorModal"', index)
        self.assertIn('id="reportEditorRibbon"', index)
        self.assertIn('id="reportEditorProofPane"', index)
        self.assertIn('id="reportEditorProofToggle"', index)
        self.assertIn('class="report-editor-modal"', index)
        self.assertIn("position: fixed; inset: 0; z-index: 4000", style)
        self.assertIn("EditorTable", source)
        self.assertIn("replaceAll", source)
        self.assertIn("cmhk-report-editor-draft", source)
        self.assertEqual(source.count('data-custom-select="native"'), 4)
        self.assertEqual(bundle_source.count('data-custom-select="native"'), 4)
        self.assertNotIn("underline: false", source)
        self.assertIn("preferredZoom", source)
        self.assertIn("renderProof", source)
        self.assertIn("右侧为最近保存版", source)
        self.assertIn("report-editor-proof-pane", style)
        self.assertIn('id="reportEditorZoom" type="range" min="35"', index)
        self.assertIn('class="row-icon-button edit-report-button"', app)
        self.assertIn("data-report-editor-path", workspace)
        self.assertIn("/api/report-editor", source)
        self.assertTrue(bundle.is_file())
        self.assertGreater(bundle.stat().st_size, 300_000)
        self.assertIn("tiptap-report-editor-3.30.5.min.js?v=4", index)
        self.assertNotIn("cdn.jsdelivr", index)
        self.assertNotIn("unpkg.com", index)

    def test_subscription_page_selects_weekly_version_and_shows_day_hour_countdown(self):
        script = (ROOT / "web/static/subscription-admin.js").read_text(encoding="utf-8")
        style = (ROOT / "web/static/subscription-admin.css").read_text(encoding="utf-8")

        self.assertIn("手动推送周报版本", script)
        self.assertIn("自动选择最新正式版", script)
        self.assertIn("weeklyReportPath: state.manualWeeklyPath", script)
        self.assertIn("data-weekly-picker-search", script)
        self.assertIn('action: "setWeeklyReportPreference"', script)
        self.assertIn("weekly-picker-options", style)
        self.assertIn("max-height: 224px", style)
        self.assertIn("距离下一次自动发送还有 ${days} 天 ${hours} 小时", script)
        self.assertIn("data-report-schedule-countdown", script)
        self.assertIn(".report-schedule-countdown", style)
        self.assertIn("color: #7c8d94", style)


if __name__ == "__main__":
    unittest.main()
