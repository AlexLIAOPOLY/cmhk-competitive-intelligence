from __future__ import annotations

import json
import re
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.shared import RGBColor

import generate_weekly_report as report


def detailed_text(prefix: str, event_date: str = "2026年7月12日") -> str:
    del event_date
    return (
        f"{prefix}"
        "项目覆盖网络、云计算和企业服务三类核心业务场景，并公布合作对象、实施范围及主要资源安排。"
        "首阶段已经完成技术验证并进入落地部署，后续批次将按现有计划扩展至更多业务节点。"
    )


def make_item(item_id: str, index: int, *, title: str | None = None) -> dict:
    return {
        "id": item_id,
        "row": f"row-{index}",
        "index": index,
        "localIndex": index,
        "section": "行业资讯",
        "subject": f"测试主体{index}",
        "tag": "业务动态",
        "title": title or f"测试事件{index}",
        "detail": detailed_text(f"测试主体{index}公布一项最新业务进展。"),
        "rawDetail": f"测试主体{index}公布一项最新业务进展。",
        "eventAt": "2026-07-12",
        "sourceIds": [f"S{index}"],
        "sourceName": "测试来源",
    }


def make_model(item: dict | None = None) -> dict:
    value = item or make_item("W001", 1)
    source_id = value["sourceIds"][0]
    return {
        "company": "中国移动香港公司",
        "department": "中国移动香港公司战略部",
        "generatedDate": "2026年7月15日",
        "title": "战略内参",
        "range": {"start": "2026-07-02", "end": "2026-07-15"},
        "toc": [
            {
                "index": value["index"],
                "section": "行业资讯",
                "tag": value["tag"],
                "title": value["title"],
            }
        ],
        "sections": [{"name": "行业资讯", "narrative": "", "items": [value]}],
        "sources": [
            {
                "sourceId": source_id,
                "row": value["row"],
                "section": "行业资讯",
                "title": value["title"],
                "url": f"https://example.test/{source_id}",
                "sourceName": "测试来源",
                "object": value["subject"],
                "tag": value["tag"],
                "publishedAt": value["eventAt"],
            }
        ],
    }


def write_small_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("中国移动香港公司")
    doc.add_paragraph("中国移动香港公司战略部    YYYY年M月D日")
    doc.add_paragraph("目 录")
    doc.add_paragraph("")
    for section_name in report.SECTION_ORDER:
        doc.add_paragraph(section_name)
        doc.add_paragraph("1.【标签】一句话事件标题")
    doc.add_paragraph("")
    doc.add_paragraph("政治资讯")
    tag = doc.add_paragraph()
    tag_run = tag.add_run("标签")
    tag_run.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)
    doc.add_paragraph("一、一句话事件标题")
    doc.add_paragraph("事件事实正文。")
    doc.add_paragraph("")
    doc.save(path)


class PublicationLabelTests(unittest.TestCase):
    def test_publication_date_stays_in_audit_data_but_is_hidden_from_report_body(self) -> None:
        row = {
            "id": "row-publication-date",
            "company": "测试主体",
            "metric": "业务动态",
            "value": "测试事实",
            "generatedAt": "2026-07-15T09:30:00+08:00",
            "fetched_at_hkt": "2026-07-15T09:31:00+08:00",
            "sources": [
                {
                    "url": "https://example.test/article",
                    "publishedAt": "2026-07-12T18:20:00+08:00",
                }
            ],
        }
        included, _ = report.filter_biweekly_rows(
            [row],
            now=report.parse_report_date("2026-07-15T10:00:00+08:00"),
        )
        self.assertEqual(included[0]["publicationDate"], "2026-07-12")

        item = make_item("W001", 1)
        item["eventAt"] = included[0]["publicationDate"]
        item["generatedAt"] = row["generatedAt"]
        item["reviewDecision"] = "approve"
        model = make_model(item)
        markdown = report.weekly_to_markdown(model)

        self.assertNotIn("发布时间", markdown)
        self.assertNotIn("来源：", markdown)
        self.assertEqual(model["sources"][0]["publishedAt"], "2026-07-12")

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "report.docx"
            write_small_template(template_path)
            with patch.object(report, "SOURCE_WORD_TEMPLATE", template_path):
                report.weekly_to_docx(model, output_path)
            rendered = Document(output_path)
            rendered_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)

        self.assertNotIn("发布时间", rendered_text)
        self.assertNotIn("来源：", rendered_text)

    def test_detail_omits_source_date_leadin_and_follow_up_scaffolding(self) -> None:
        valid = detailed_text("测试主体公布一项新进展。")
        scaffolded = (
            "据测试来源于2026年7月12日发布的信息，测试主体公布一项新进展。"
            "项目已经进入实施阶段。后续仍需关注执行进度。"
        )
        item = make_item("W001", 1)
        item["detail"] = scaffolded
        item["reviewDecision"] = "approve"
        model = make_model(item)
        with self.assertRaisesRegex(ValueError, "来源日期套话或关注式结尾"):
            report.validate_report_model(model)

        item["detail"] = valid
        report.validate_report_model(model)

    def test_title_only_body_and_embedded_webpage_timestamp_are_rejected(self) -> None:
        item = make_item("W001", 1, title="测试主体推进新一轮业务部署")
        item["detail"] = item["title"]
        with self.assertRaisesRegex(ValueError, "重复标题"):
            report.validate_human_template_content(make_model(item))

        item["detail"] = (
            "测试主体扩大跨境业务合作，并披露新的供应链数据。"
            "发表时间：28/07/2026 - 15:29。"
        )
        with self.assertRaisesRegex(ValueError, "发布时间、来源套话或关注式结尾"):
            report.validate_human_template_content(make_model(item))

    def test_publication_scaffolding_is_removed_before_model_payload(self) -> None:
        cleaned = report.strip_publication_scaffolding(
            "发表时间：28/07/2026 - 15:29。。测试主体披露跨境业务合作及供应链数据。",
            "2026-07-24",
        )
        self.assertEqual(cleaned, "测试主体披露跨境业务合作及供应链数据。")

    def test_search_result_noise_is_rejected_and_never_joined_into_fallback(self) -> None:
        item = make_item("W001", 1, title="工信部推动中小企业AI深度融合")
        item["detail"] = (
            "工信部推动人工智能与中小企业深度融合。"
            "© 2026 - Конфиденциальность - Условия。"
        )
        item["rawDetail"] = "工信部发文促进人工智能在中小企业深度融合应用。"
        item["webResearch"] = {
            "results": [
                {"snippet": "无关视频标题，完整版视频下载。"},
                {"snippet": "德国政府分析其他行业的贸易流向。"},
            ]
        }

        self.assertTrue(report.summary_has_search_noise(item["detail"]))
        self.assertEqual(report.deterministic_limited_weekly_detail(item), "")
        with self.assertRaisesRegex(ValueError, "网页搜索结果噪声"):
            report.validate_human_template_content(make_model(item))

    def test_fallback_selects_only_concise_snippet_from_matching_web_evidence(self) -> None:
        item = make_item(
            "W001",
            1,
            title="PCB行业龙头投资52亿元建设第二工厂",
        )
        item["originalTitle"] = item["title"]
        item["rawDetail"] = item["title"]
        item["detail"] = item["title"]
        item["webResearch"] = {
            "results": [
                {
                    "title": "无关地区举办消费品展览",
                    "snippet": "展览吸引数百家企业参展，并发布多项消费数据。",
                },
                {
                    "title": "PCB行业龙头投资52亿元建设第二工厂",
                    "snippet": (
                        "项目在东莞松山湖动工，聚焦AI服务器高频高速材料和高端封装基板材料。"
                        "第二工厂将建设研发、生产和配套设施，并导入高端印制电路板制造能力。"
                        "建成后将为云计算、6G通信及智能汽车电子提供关键材料，进一步扩大当地先进制造产能。"
                    ),
                },
            ]
        }

        detail = report.deterministic_limited_weekly_detail(item)

        self.assertIn("东莞松山湖", detail)
        self.assertNotIn("消费品展览", detail)
        self.assertTrue(report.summary_is_concise(detail))
        self.assertTrue(report.summary_adds_information(item["title"], detail, item["title"]))

    def test_title_only_evidence_requires_verified_supplemental_detail(self) -> None:
        pcb = make_item(
            "W001",
            1,
            title="PCB行业龙头投52亿建第二工厂 东莞「全球智造中心」建设加速",
        )
        pcb["originalTitle"] = pcb["title"]
        pcb["rawDetail"] = pcb["title"]
        pcb["detail"] = pcb["title"]
        pcb["webResearch"] = {"results": []}
        self.assertEqual(report.deterministic_limited_weekly_detail(pcb), "")

        grant = make_item(
            "W002",
            2,
            title="研资局拨1.4亿资助20学者 项目多聚焦AI安全",
        )
        grant["originalTitle"] = grant["title"]
        grant["rawDetail"] = grant["title"]
        grant["detail"] = grant["title"]
        grant["webResearch"] = {"results": []}
        grant_detail = report.deterministic_limited_weekly_detail(grant)
        self.assertIn("未来60个月", grant_detail)
        self.assertTrue(report.summary_has_reference_density(grant_detail))

        agreement = make_item(
            "W003",
            3,
            title="特区政府与国家工信部签署协议推进共建制造业创新中心- 港闻",
        )
        agreement["originalTitle"] = agreement["title"]
        agreement["title"] = report.deterministic_evidence_weekly_title(agreement)
        agreement["rawDetail"] = agreement["title"]
        agreement["detail"] = agreement["title"]
        agreement["webResearch"] = {"results": []}
        self.assertEqual(
            agreement["title"],
            "特区政府与国家工信部签署协议推进共建制造业创新中心",
        )
        agreement_detail = report.deterministic_limited_weekly_detail(agreement)
        self.assertIn("香港微电子研发院", agreement_detail)
        self.assertTrue(report.summary_has_reference_density(agreement_detail))

    def test_overlong_detail_is_trimmed_only_at_a_complete_sentence(self) -> None:
        detail = (
            "据测试来源于2026年7月12日发布的信息，测试主体公布第一项可核验进展。"
            + "第二句补充实施背景、参与主体和当前执行安排。" * 3
            + "第三部分继续列出后续计划、业务范围和正式披露中的事实。" * 3
            + "最后一句提供仍需持续观察的公开指标与进度。" * 3
        )

        trimmed = report.trim_weekly_detail(detail, max_chars=180)

        self.assertLessEqual(len(re.sub(r"\s+", "", trimmed)), 180)
        self.assertGreaterEqual(len(re.findall(r"[。！？!?]", trimmed)), 3)
        self.assertTrue(trimmed.endswith("。"))
        self.assertNotIn("…", trimmed)

        item = make_item("W001", 1)
        item["detail"] = detailed_text("测试主体公布一项新进展。")[:-1] + "…"
        item["reviewDecision"] = "approve"
        with self.assertRaisesRegex(ValueError, "截断省略号"):
            report.validate_report_model(make_model(item))

    def test_visible_summary_matches_human_reference_information_density(self) -> None:
        dense = (
            "主体公布新一轮网络建设方案，覆盖香港主要商业区、交通枢纽及重点住宅区域。"
            "项目采用多频段协同和智能节能技术，并投入既有核心网资源支持部署。"
            "首阶段已经完成测试并进入商用准备，后续建设将按已公布计划分批实施。"
        )
        self.assertTrue(report.summary_has_reference_density(dense))
        self.assertFalse(report.summary_has_reference_density("主体公布关键方案，覆盖三类业务场景。"))
        self.assertTrue(
            report.summary_has_reference_density(
                dense + "项目团队同步公布跨部门协作安排和现阶段执行结果。"
            )
        )

        item = make_item("W001", 1)
        item["detail"] = "第一句交代主体动作。第二句补充关键数字。第三句继续搬运活动背景。"
        # Length is a diagnostic comparison with the human corpus, not a
        # publication blocker. Editorial quality comes from the examples.
        report.validate_human_template_content(make_model(item))


class StrategicReferenceDocxStructureTests(unittest.TestCase):
    def test_body_order_and_gray_tag_match_the_strategic_reference_format(self) -> None:
        item = make_item("W001", 1, title="测试主体推进新一轮业务部署")
        model = make_model(item)

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "report.docx"
            write_small_template(template_path)
            with patch.object(report, "SOURCE_WORD_TEMPLATE", template_path):
                report.weekly_to_docx(model, output_path)
            rendered = Document(output_path)

        texts = [paragraph.text.strip() for paragraph in rendered.paragraphs]
        body_start = texts.index("行业资讯")
        self.assertEqual(
            texts[body_start : body_start + 4],
            [
                "行业资讯",
                "业务动态",
                item["title"],
                item["detail"],
            ],
        )
        self.assertNotIn("发布时间", "\n".join(texts))
        self.assertNotIn("来源：", "\n".join(texts))
        self.assertFalse(re.match(r"^(?:\d+|[一二三四五六七八九十]+)、", texts[body_start + 2]))

        tag_paragraph = rendered.paragraphs[body_start + 1]
        self.assertTrue(tag_paragraph.runs)
        self.assertEqual(tag_paragraph.runs[0].font.color.rgb, RGBColor(0xA6, 0xA6, 0xA6))
        self.assertTrue(rendered.paragraphs[body_start].paragraph_format.page_break_before)


class IndependentReviewerTests(unittest.TestCase):
    def test_report_with_too_few_items_fails_before_publication(self) -> None:
        model = make_model(make_item("W001", 1))

        with self.assertRaisesRegex(RuntimeError, "最低"):
            report.apply_weekly_ai_review(model, progress=lambda _message: None)

    def test_weekly_online_research_adds_verification_sources_before_review(self) -> None:
        item = make_item("W001", 1)
        model = make_model(item)

        def search_client(query, _limit):
            return {
                "query": query,
                "provider": "unit",
                "results": [
                    {
                        "title": "联网核实来源",
                        "url": "https://verify.example.com/report",
                        "snippet": "该公开来源交叉确认原稿事实，并补充一项可核验背景。",
                    }
                ],
                "error": "",
            }

        researched = report.research_weekly_model_online(
            model,
            search_client=search_client,
            progress=lambda _message: None,
        )
        researched_item = researched["sections"][0]["items"][0]
        verification_sources = [
            source for source in researched["sources"] if source.get("verificationOnly")
        ]

        self.assertEqual(researched_item["webResearch"]["provider"], "unit")
        self.assertEqual(len(verification_sources), 1)
        self.assertIn(verification_sources[0]["sourceId"], researched_item["sourceIds"])
        self.assertEqual(researched["webResearchAudit"]["itemsWithResults"], 1)
        researched_item["reviewDecision"] = "approve"
        report.validate_report_model(researched)

    def test_weekly_online_research_sends_selected_source_fulltext_to_writer(self) -> None:
        item = make_item("W001", 1, title="人工选中的原始新闻标题")
        model = make_model(item)
        search_rows = [
            {
                "id": "W001",
                "query": item["title"],
                "provider": "unit",
                "results": [],
                "error": "",
            }
        ]
        selected_content = detailed_text("人工选中来源披露完整事实。")

        with (
            patch.object(report, "run_web_research", return_value=search_rows),
            patch.object(
                report,
                "_fetch_search_result_content",
                return_value=selected_content,
            ) as fetch_content,
        ):
            researched = report.research_weekly_model_online(
                model,
                progress=lambda _message: None,
            )

        evidence = researched["sections"][0]["items"][0]["webResearch"][
            "lockedSourceEvidence"
        ]
        self.assertEqual(evidence["content"], selected_content)
        self.assertEqual(evidence["verification"], "selected_direct_source")
        self.assertTrue(fetch_content.called)

    def test_final_review_audit_is_renumbered_after_section_reordering(self) -> None:
        items = [make_item(f"W00{index}", index) for index in range(1, 5)]
        for index, item in enumerate(items, start=1):
            item["sourceIds"] = [f"S{index}"]
            item["section"] = "行业资讯"
            item["_reviewAuditId"] = f"W{index:03d}"
            item["reviewDecision"] = "approve"
            item["webResearch"] = {
                "query": f"query-{index}",
                "provider": "unit",
                "results": [{"title": f"result-{index}", "url": f"https://verify.test/{index}"}],
            }
        items[3]["section"] = "政治资讯"
        model = {
            "range": {"start": "2026-07-01", "end": "2026-07-14"},
            "plannedRange": {"start": "2026-07-01", "end": "2026-07-14"},
            "periodStatus": "final",
            "issueDate": "2026-07-14",
            "asOf": "2026-07-14T10:00:00+08:00",
            "sections": [{"name": "行业资讯", "narrative": "", "items": items}],
            "sources": [
                {
                    "sourceId": f"S{index}",
                    "url": f"https://source.test/{index}",
                    "title": f"source-{index}",
                    "publishedAt": "2026-07-12",
                }
                for index in range(1, 5)
            ],
            "toc": [],
        }
        audit = {
            "reviewStatus": "passed",
            "reviewerModel": "unit-reviewer",
            "reviewPromptVersion": "unit",
            "approvedItems": 4,
            "revisedItems": 0,
            "rejectedItems": 0,
            "items": [
                {
                    "id": f"W{index:03d}",
                    "decision": "approve",
                    "title": f"old-audit-{index}",
                    "sourceIds": [f"S{index}"],
                }
                for index in range(1, 5)
            ],
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            usage_path = Path(temp_dir) / "usage.json"
            usage_path.write_text("{}", encoding="utf-8")
            with (
                patch.object(report, "WEEKLY_USAGE_AUDIT", usage_path),
                patch.object(report, "research_weekly_model_online", return_value=model),
                patch.object(
                    report,
                    "write_weekly_items_once",
                    return_value=items,
                ),
                patch.object(
                    report,
                    "edit_weekly_items_once",
                    return_value=(items, audit),
                ),
            ):
                reviewed = report.apply_weekly_ai_review(
                    model,
                    progress=lambda _message: None,
                )

        final_items = [
            item for section in reviewed["sections"] for item in section["items"]
        ]
        self.assertEqual(final_items[0]["title"], items[3]["title"])
        self.assertEqual(reviewed["reviewAudit"]["items"][0]["id"], "W001")
        self.assertEqual(
            reviewed["reviewAudit"]["items"][0]["title"],
            items[3]["title"],
        )
        self.assertEqual(
            reviewed["reviewAudit"]["webSearch"]["queries"][0]["query"],
            "query-4",
        )

    def test_weekly_online_research_fails_closed_when_search_is_unavailable(self) -> None:
        item = make_item("W001", 1)
        model = make_model(item)

        def empty_search(query, _limit):
            return {"query": query, "provider": "", "results": [], "error": "offline"}

        with self.assertRaisesRegex(RuntimeError, "均无可用结果"):
            report.research_weekly_model_online(
                model,
                search_client=empty_search,
                progress=lambda _message: None,
            )

    def test_independent_reviewer_approves_revises_and_rejects_per_item(self) -> None:
        items = [make_item(f"W00{index}", index) for index in range(1, 4)]
        revised_detail = detailed_text("独立审核模型根据原始事实修订了第二条的表述。")
        reviewer_response = {
            "items": [
                {
                    "id": "W001",
                    "decision": "approve",
                    "scores": {"factuality": 5, "detail": 5, "relevance": 5, "language": 5},
                    "reason": "事实、日期与来源一致",
                },
                {
                    "id": "W002",
                    "decision": "revise",
                    "scores": {"factuality": 5, "detail": 5, "relevance": 4, "language": 5},
                    "title": "独立审核后的修订标题",
                    "detail": revised_detail,
                    "reason": "原文需要弱化确定性结论",
                },
                {
                    "id": "W003",
                    "decision": "reject",
                    "scores": {"factuality": 2, "detail": 3, "relevance": 3, "language": 4},
                    "reason": "证据不足以支持正文",
                },
            ]
        }
        repaired_third_detail = detailed_text("强制修复阶段依据原始事实保留并完善了第三条。")
        repaired_third_response = {
            "items": [
                {
                    "id": "W003",
                    "decision": "revise",
                    "scores": {"factuality": 5, "detail": 4, "relevance": 4, "language": 5},
                    "title": "第三条修复后的标题",
                    "detail": repaired_third_detail,
                    "reason": "已按锁定证据修复",
                }
            ]
        }
        progress_events: list[tuple[tuple, dict]] = []

        def progress(*args, **kwargs) -> None:
            progress_events.append((args, kwargs))

        with tempfile.TemporaryDirectory() as temp_dir:
            review_cache = Path(temp_dir) / "review-cache.json"
            with (
                patch.object(report, "WEEKLY_REVIEW_CACHE", review_cache, create=True),
                patch.object(
                    report,
                    "_call_weekly_quality_reviewer_llm",
                    side_effect=[reviewer_response, repaired_third_response],
                ) as reviewer_call,
                patch.object(
                    report,
                    "_call_weekly_writer_llm",
                    side_effect=AssertionError("审核不得复用写作模型调用"),
                ) as writer_call,
            ):
                reviewed, audit = report.review_weekly_items_with_ai(
                    items,
                    progress=progress,
                    bypass_cache=True,
                )

        reviewer_call.assert_called()
        writer_call.assert_not_called()
        self.assertEqual([item["id"] for item in reviewed], ["W001", "W002", "W003"])
        self.assertEqual(reviewed[0]["reviewDecision"], "approve")
        self.assertEqual(reviewed[1]["reviewDecision"], "revise")
        self.assertEqual(reviewed[1]["title"], "独立审核后的修订标题")
        self.assertEqual(reviewed[1]["detail"], revised_detail)
        self.assertEqual(reviewed[1]["sourceIds"], items[1]["sourceIds"])
        self.assertEqual(reviewed[1]["eventAt"], items[1]["eventAt"])
        self.assertEqual(reviewed[2]["reviewDecision"], "revise")
        self.assertEqual(reviewed[2]["title"], "第三条修复后的标题")
        self.assertEqual(audit["approvedItems"], 1)
        self.assertEqual(audit["revisedItems"], 2)
        self.assertEqual(audit["rejectedItems"], 0)
        rejected = [entry for entry in audit["items"] if entry["decision"] == "reject"]
        self.assertEqual(rejected, [])
        self.assertTrue(progress_events)

    def test_reviewer_rejection_triggers_writer_rewrite_and_final_review(self) -> None:
        item = make_item("W001", 1)
        rejected_response = {
            "items": [
                {
                    "id": "W001",
                    "decision": "reject",
                    "scores": {"factuality": 3, "detail": 3, "relevance": 4, "language": 4},
                    "reason": "正文事实组织和详细度不足",
                }
            ]
        }
        rewritten_detail = detailed_text("写作模型依据锁定资料和联网证据重新组织了完整事实。")
        writer_response = {
            "items": [
                {
                    "id": "W001",
                    "status": "ok",
                    "title": "联网证据支持的完整重写标题",
                    "detail": rewritten_detail,
                    "used_fact_ids": ["F001", "F002"],
                }
            ]
        }
        approved_response = {
            "items": [
                {
                    "id": "W001",
                    "decision": "approve",
                    "scores": {"factuality": 5, "detail": 5, "relevance": 4, "language": 5},
                    "reason": "重写后达到发布门槛",
                }
            ]
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            with (
                patch.object(
                    report,
                    "WEEKLY_REVIEW_CACHE",
                    Path(temp_dir) / "review-cache.json",
                    create=True,
                ),
                patch.object(
                    report,
                    "_call_weekly_quality_reviewer_llm",
                    side_effect=[rejected_response, rejected_response, approved_response],
                ) as reviewer_call,
                patch.object(
                    report,
                    "_call_weekly_writer_llm",
                    return_value=writer_response,
                ) as writer_call,
            ):
                reviewed, audit = report.review_weekly_items_with_ai(
                    [item],
                    progress=lambda _message: None,
                    bypass_cache=True,
                )

        self.assertEqual(reviewer_call.call_count, 3)
        writer_call.assert_called_once()
        self.assertEqual(len(reviewed), 1)
        self.assertEqual(reviewed[0]["title"], "联网证据支持的完整重写标题")
        self.assertEqual(reviewed[0]["detail"], rewritten_detail)
        self.assertEqual(reviewed[0]["writerStatus"], "llm_rewritten_after_review")
        self.assertEqual(audit["rejectedItems"], 0)

    def test_single_item_score_only_review_response_is_normalized(self) -> None:
        item = make_item("W001", 1)
        score_only_response = {
            "factuality": 5,
            "detail": 5,
            "relevance": 4,
            "language": 5,
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            with (
                patch.object(
                    report,
                    "WEEKLY_REVIEW_CACHE",
                    Path(temp_dir) / "review-cache.json",
                    create=True,
                ),
                patch.object(
                    report,
                    "_call_weekly_quality_reviewer_llm",
                    return_value=score_only_response,
                ) as reviewer_call,
                patch.object(
                    report,
                    "_call_weekly_writer_llm",
                    side_effect=AssertionError("合格评分不应触发重写"),
                ) as writer_call,
            ):
                reviewed, audit = report.review_weekly_items_with_ai(
                    [item],
                    progress=lambda _message: None,
                    bypass_cache=True,
                )

        reviewer_call.assert_called_once()
        writer_call.assert_not_called()
        self.assertEqual(len(reviewed), 1)
        self.assertEqual(reviewed[0]["reviewDecision"], "approve")
        self.assertEqual(audit["approvedItems"], 1)
        self.assertEqual(audit["rejectedItems"], 0)

    def test_repeated_rejection_uses_verified_replacement_without_dropping_item(self) -> None:
        item = make_item("W001", 1)
        rejected_response = {
            "items": [
                {
                    "id": "W001",
                    "decision": "reject",
                    "scores": {"factuality": 3, "detail": 3, "relevance": 3, "language": 4},
                    "reason": "当前条目仍不适合发布",
                }
            ]
        }
        rewritten_detail = detailed_text("原条目经过重写但仍由独立审核模型判定不适合发布。")
        replacement_detail = detailed_text("备用文章依据锁定资料和联网结果形成了完整可核验正文。")
        writer_responses = [
            {
                "items": [
                    {
                        "id": "W001",
                        "status": "ok",
                        "title": "原条目重写后的标题",
                        "detail": rewritten_detail,
                    }
                ]
            },
            {
                "items": [
                    {
                        "id": "W001",
                        "status": "ok",
                        "title": "备用文章重写后的标题",
                        "detail": replacement_detail,
                    }
                ]
            },
        ]
        approved_response = {
            "items": [
                {
                    "id": "W001",
                    "decision": "approve",
                    "scores": {"factuality": 5, "detail": 5, "relevance": 5, "language": 5},
                    "reason": "备用文章通过全部门禁",
                }
            ]
        }
        replacement = make_item("R001", 9, title="备用文章原始标题")
        replacement["_replacementSource"] = {
            "sourceId": "",
            "title": "备用文章原始标题",
            "url": "https://example.test/replacement",
            "sourceName": "测试来源",
            "publishedAt": "2026-07-12",
        }
        research = {
            "id": "W001",
            "query": "备用文章",
            "provider": "test-search",
            "results": [{"title": "备用文章证据", "url": "https://example.test/evidence"}],
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            with (
                patch.object(
                    report,
                    "WEEKLY_REVIEW_CACHE",
                    Path(temp_dir) / "review-cache.json",
                    create=True,
                ),
                patch.object(
                    report,
                    "_call_weekly_quality_reviewer_llm",
                    side_effect=[
                        rejected_response,
                        rejected_response,
                        rejected_response,
                        approved_response,
                    ],
                ),
                patch.object(
                    report,
                    "_call_weekly_writer_llm",
                    side_effect=writer_responses,
                ),
                patch.object(report, "run_web_research", return_value=[research]),
            ):
                reviewed, audit = report.review_weekly_items_with_ai(
                    [item],
                    progress=lambda _message: None,
                    bypass_cache=True,
                    replacement_candidates=[replacement],
                )

        self.assertEqual(len(reviewed), 1)
        self.assertEqual(reviewed[0]["title"], "备用文章重写后的标题")
        self.assertEqual(reviewed[0]["sourceIds"], item["sourceIds"])
        self.assertEqual(
            reviewed[0]["_replacementSource"]["url"],
            "https://example.test/replacement",
        )
        self.assertEqual(audit["reviewReplacementCount"], 1)
        self.assertEqual(audit["rejectedItems"], 0)

    def test_repeated_rejection_without_replacement_stops_publication(self) -> None:
        item = make_item("W001", 1, title="人工入选新闻原始标题")
        item["webResearch"] = {
            "query": "人工入选新闻原始标题",
            "provider": "unit",
            "results": [
                {
                    "title": "联网核验资料",
                    "url": "https://example.test/evidence",
                    "snippet": "公开资料确认该主体公布业务进展，并说明当前实施范围和后续安排。",
                }
            ],
        }
        rejected_response = {
            "items": [
                {
                    "id": "W001",
                    "decision": "reject",
                    "scores": {"factuality": 3, "detail": 3, "relevance": 3, "language": 4},
                    "issues": ["模型未能形成可通过版本"],
                    "reason": "仍需修复",
                }
            ]
        }
        progress_messages: list[str] = []

        with tempfile.TemporaryDirectory() as temp_dir:
            with (
                patch.object(
                    report,
                    "WEEKLY_REVIEW_CACHE",
                    Path(temp_dir) / "review-cache.json",
                    create=True,
                ),
                patch.object(
                    report,
                    "_call_weekly_quality_reviewer_llm",
                    return_value=rejected_response,
                ) as reviewer_call,
                patch.object(
                    report,
                    "_call_weekly_writer_llm",
                    side_effect=TimeoutError("writer unavailable"),
                ) as writer_call,
            ):
                with self.assertRaisesRegex(RuntimeError, "停止发布并继续修稿"):
                    report.review_weekly_items_with_ai(
                        [item],
                        progress=progress_messages.append,
                        bypass_cache=True,
                    )

        self.assertGreaterEqual(reviewer_call.call_count, 2)
        writer_call.assert_called_once()
        self.assertIn("最终证据约束修复完成", " ".join(progress_messages))

    def test_unreviewed_or_writer_only_item_cannot_pass_quality_gate(self) -> None:
        item = make_item("W001", 1)
        item["writerStatus"] = "llm"
        item["qualityStatus"] = "passed"
        model = make_model(item)

        with self.assertRaisesRegex(ValueError, "审核"):
            report.validate_review_gate(model)
        with self.assertRaisesRegex(ValueError, "审核"):
            report.validate_report_model(model)

        item["reviewDecision"] = "approve"
        report.validate_review_gate(model)
        report.validate_report_model(model)


class QualitySidecarTests(unittest.TestCase):
    def test_every_written_docx_gets_a_bound_quality_sidecar(self) -> None:
        item = make_item("W001", 1)
        item.update(
            {
                "reviewDecision": "approve",
                "reviewReason": "日期、事实与来源一致",
            }
        )
        model = make_model(item)
        model.update(
            {
                "range": {"start": "2026-07-07", "end": "2026-07-15"},
                "plannedRange": {"start": "2026-07-07", "end": "2026-07-17"},
                "periodStatus": "draft",
                "issueDate": "2026-07-17",
                "asOf": "2026-07-15T10:00:00+08:00",
            }
        )
        model["reviewAudit"] = {
            "reviewStatus": "passed",
            "approved": 1,
            "revised": 0,
            "rejected": 0,
            "items": [
                {
                    "id": "W001",
                    "reviewDecision": "approve",
                    "reason": item["reviewReason"],
                }
            ],
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "report.docx"
            write_small_template(template_path)
            with patch.object(report, "SOURCE_WORD_TEMPLATE", template_path):
                report.weekly_to_docx(model, output_path)

            expected_sidecar = Path(str(output_path) + ".quality.json")
            self.assertEqual(report.weekly_quality_sidecar_path(output_path), expected_sidecar)
            self.assertTrue(expected_sidecar.exists())
            payload = json.loads(expected_sidecar.read_text(encoding="utf-8"))

        self.assertEqual(payload["reportFile"], "report.docx")
        self.assertEqual(payload["reviewStatus"], "passed")
        self.assertEqual(payload["window"], {"start": "2026-07-07", "end": "2026-07-15"})
        self.assertEqual(payload["plannedWindow"], {"start": "2026-07-07", "end": "2026-07-17"})
        self.assertEqual(payload["periodStatus"], "draft")
        self.assertEqual(payload["issueDate"], "2026-07-17")
        self.assertEqual(payload["items"][0]["reviewDecision"], "approve")

    def test_limited_report_writes_an_honest_quality_sidecar(self) -> None:
        model = make_model()
        report.record_weekly_limitation(
            model,
            "review",
            "独立审核暂不可用",
            impact="未完成独立审核",
            action="保留锁定来源并以受限模式输出",
            progress=lambda _message: None,
        )
        model = report.finalize_weekly_limited_model(model)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "limited.docx"
            report.weekly_to_emergency_docx(model, output_path, reason="模板不可用")
            sidecar_path = report.weekly_quality_sidecar_path(output_path)
            payload = json.loads(sidecar_path.read_text(encoding="utf-8"))

        self.assertEqual(payload["reviewStatus"], "limited")
        self.assertEqual(payload["generationMode"], "limited")
        self.assertEqual(payload["items"][0]["reviewDecision"], "evidence_repair")
        self.assertEqual(payload["limitations"][0]["stage"], "review")

    def test_limited_mode_repairs_thin_body_from_locked_supplemental_evidence(self) -> None:
        item = make_item("W001", 1, title="测试主体公布网络部署")
        item["detail"] = "测试主体公布网络部署。"
        item["rawDetail"] = item["detail"]
        item["limitedNotice"] = "写作模型暂不可用"
        item["webResearch"] = {
            "supplementalEvidence": {
                "detail": detailed_text("测试主体公布网络部署。"),
            }
        }

        model = report.finalize_weekly_limited_model(make_model(item))

        self.assertTrue(report.summary_has_reference_density(model["sections"][0]["items"][0]["detail"]))
        report.validate_human_template_content(model)

    def test_limited_mode_recovers_previously_validated_writer_cache(self) -> None:
        item = make_item("W001", 1, title="测试主体公布网络部署")
        item["detail"] = item["title"]
        item["rawDetail"] = item["title"]
        item["facts"] = [detailed_text("测试主体公布网络部署。")]
        item["_weeklyWriterCacheKey"] = "validated-draft"
        cached = {
            "validated-draft": {
                "status": "ok",
                "title": item["title"],
                "detail": item["facts"][0],
            }
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            cache_path = Path(temp_dir) / "weekly-cache.json"
            cache_path.write_text(json.dumps(cached, ensure_ascii=False), encoding="utf-8")
            with patch.object(report, "WEEKLY_LLM_CACHE", cache_path):
                model = report.finalize_weekly_limited_model(make_model(item))

        repaired = model["sections"][0]["items"][0]
        self.assertEqual(repaired["writerStatus"], "validated_cache_recovery")
        report.validate_human_template_content(model)

    def test_prepare_human_template_stops_on_irreparable_body(self) -> None:
        item = make_item("W001", 1, title="测试主体公布网络部署")
        item["detail"] = item["title"]
        item["rawDetail"] = item["title"]
        model = make_model(item)

        with self.assertRaisesRegex(ValueError, "已停止发布"):
            report.prepare_human_template_content(model, progress=lambda _message: None)

    def test_build_weekly_model_recovers_from_untouched_selected_copy(self) -> None:
        item = make_item("W001", 1, title="测试主体公布新业务")
        item["detail"] = detailed_text("原始人工入选正文。")
        item["rawDetail"] = item["detail"]
        selected = make_model(item)

        def destructive_review(working, progress=None):
            del progress
            working["sections"][0]["items"][0]["detail"] = ""
            working["sections"][0]["items"][0]["rawDetail"] = ""
            raise RuntimeError("模型重试失败")

        with (
            patch.object(report, "build_review_sheet_weekly_model", return_value=selected),
            patch.object(report, "apply_weekly_ai_review", side_effect=destructive_review),
        ):
            recovered = report.build_weekly_model([])

        recovered_item = recovered["sections"][0]["items"][0]
        self.assertIn("原始人工入选正文", recovered_item["detail"])
        self.assertEqual(len(recovered["sections"][0]["items"]), 1)

    def test_main_uses_emergency_docx_when_standard_template_fails(self) -> None:
        period = report.resolve_weekly_period(
            now=report.parse_report_date("2026-07-15T10:00:00+08:00")
        )
        model = report.build_weekly_limitation_model(
            period,
            stage="selection",
            reason="没有人工入选新闻",
            progress=lambda _message: None,
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            output_path = temp_path / "7月15日周报.docx"
            with (
                patch.object(report, "ROOT", temp_path),
                patch.object(report, "WEEKLY_MD", temp_path / "weekly_report.md"),
                patch.object(report, "WEEKLY_HTML", temp_path / "weekly_report.html"),
                patch.object(report, "TEMPLATE_MD", temp_path / "weekly_report_template.md"),
                patch.object(report, "AGENT_MD_ALIAS", temp_path / "agent_report.md"),
                patch.object(report, "AGENT_HTML_ALIAS", temp_path / "agent_report.html"),
                patch.object(report, "WEEKLY_AI_QUALITY_AUDIT", temp_path / "weekly_quality.json"),
                patch.object(report, "SOURCE_WORD_TEMPLATE", temp_path / "missing-template.docx"),
                patch.object(report, "resolve_weekly_period", return_value=period),
                patch.object(report, "load_results", return_value=[]),
                patch.object(report, "build_weekly_model", return_value=model),
                patch.object(report, "dated_weekly_docx_path", return_value=output_path),
            ):
                report.main()

            self.assertTrue(output_path.exists())
            sidecar_path = report.weekly_quality_sidecar_path(output_path)
            self.assertTrue(sidecar_path.exists())
            payload = json.loads(sidecar_path.read_text(encoding="utf-8"))
            rendered = Document(output_path)
            rendered_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)

        self.assertEqual(payload["reviewStatus"], "limited")
        self.assertIn("战略内参", rendered_text)
        for forbidden in (
            "本期新闻信息局限说明",
            "生成说明",
            "模板渲染",
            "受限模式",
            "没有人工入选新闻",
        ):
            self.assertNotIn(forbidden, rendered_text)
        self.assertTrue(
            any(entry["stage"] == "template_render" for entry in payload["limitations"])
        )


class LlmCacheBypassTests(unittest.TestCase):
    def test_writer_cache_can_be_bypassed_for_a_fresh_llm_review_cycle(self) -> None:
        item = make_item("W001", 1, title="原始标题")
        cached_detail = detailed_text("这是写作缓存中的旧段落。")
        fresh_detail = detailed_text("这是强制绕过缓存后由模型重新生成的段落。")
        cached_result = {"status": "ok", "title": "缓存中的旧标题", "detail": cached_detail}
        fresh_result = {
            "items": [
                {
                    "id": "W001",
                    "status": "ok",
                    "title": "绕过缓存后的新标题",
                    "detail": fresh_detail,
                    "used_fact_ids": ["F001"],
                }
            ]
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            cache_path = Path(temp_dir) / "writer-cache.json"
            cache_path.write_text(json.dumps({"known-key": cached_result}, ensure_ascii=False), encoding="utf-8")
            with (
                patch.object(report, "WEEKLY_LLM_CACHE", cache_path),
                patch.object(report, "_weekly_writer_cache_key", return_value="known-key"),
                patch.object(report, "load_ai_config", return_value={"model": "deepseek-v4"}),
                patch.object(report, "_call_weekly_writer_llm", return_value=fresh_result) as llm_call,
            ):
                cached = report.enrich_weekly_items_with_llm([item], progress=lambda *args, **kwargs: None)
                llm_call.assert_not_called()
                self.assertEqual(cached[0]["writerStatus"], "cache")
                self.assertEqual(cached[0]["title"], "缓存中的旧标题")

                fresh = report.enrich_weekly_items_with_llm(
                    [item],
                    progress=lambda *args, **kwargs: None,
                    bypass_cache=True,
                )

        llm_call.assert_called()
        self.assertEqual(fresh[0]["writerStatus"], "llm")
        self.assertEqual(fresh[0]["title"], "绕过缓存后的新标题")
        self.assertNotEqual(fresh[0]["detail"], cached[0]["detail"])


class HumanReferencePromptTests(unittest.TestCase):
    @staticmethod
    def _sample_payload() -> dict:
        reports = []
        examples = []
        remaining = 161
        for report_index in range(1, 10):
            article_count = 18 if report_index < 9 else remaining
            remaining -= article_count
            source_file = f"人工战略内参{report_index}.pdf"
            reports.append(
                {
                    "source_file": source_file,
                    "report_date": f"2026-0{report_index}-01",
                    "article_count": article_count,
                }
            )
            for article_index in range(1, article_count + 1):
                examples.append(
                    {
                        "source_file": source_file,
                        "report_date": f"2026-0{report_index}-01",
                        "section": "行业资讯",
                        "subject": f"标签{report_index}",
                        "title": f"人工标题{report_index}-{article_index}",
                        "detail": f"人工完整正文{report_index}-{article_index}，包含重点动作、规模和进展。",
                    }
                )
        return {
            "unique_report_count": 9,
            "reports": reports,
            "examples": examples,
        }

    def test_prompt_contains_all_nine_reports_and_all_161_articles(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            sample_path = Path(temp_dir) / "human-examples.json"
            sample_path.write_text(
                json.dumps(self._sample_payload(), ensure_ascii=False),
                encoding="utf-8",
            )
            with patch.object(report, "WEEKLY_HUMAN_EXAMPLES", sample_path):
                prompt = report.weekly_human_examples_prompt()

        self.assertEqual(prompt.count("===== 人工完整周报"), 9)
        self.assertEqual(prompt.count("【第"), 161)
        self.assertIn("人工标题1-1", prompt)
        self.assertIn("人工完整正文9-17", prompt)

    def test_incomplete_reference_payload_fails_closed(self) -> None:
        payload = self._sample_payload()
        payload["examples"].pop()
        with tempfile.TemporaryDirectory() as temp_dir:
            sample_path = Path(temp_dir) / "human-examples.json"
            sample_path.write_text(
                json.dumps(payload, ensure_ascii=False),
                encoding="utf-8",
            )
            with (
                patch.object(report, "WEEKLY_HUMAN_EXAMPLES", sample_path),
                self.assertRaisesRegex(RuntimeError, "完整包含161篇"),
            ):
                report.weekly_human_examples_prompt()

    def test_writer_and_reviewer_requests_both_carry_full_reference_prompt(self) -> None:
        captured_requests = []

        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return False

            def read(self):
                return json.dumps(
                    {"choices": [{"message": {"content": "{\"items\":[]}"}}]}
                ).encode("utf-8")

        captured_timeouts = []

        def fake_urlopen(request, timeout):
            captured_timeouts.append(timeout)
            captured_requests.append(json.loads(request.data.decode("utf-8")))
            return FakeResponse()

        config = {
            "api_key": "test-key",
            "provider": "deepseek",
            "model": "test-model",
            "base_url": "https://example.test/v1",
        }
        marker = "FULL-NINE-REPORTS-161-ARTICLES"
        with (
            patch.object(report, "load_ai_config", return_value=config),
            patch.object(report, "weekly_human_examples_prompt", return_value=marker),
            patch.object(report, "wait_for_internal_ai_slot"),
            patch.object(
                report,
                "urlopen_with_local_proxy_fallback",
                side_effect=fake_urlopen,
            ),
        ):
            report._call_weekly_writer_llm([])
            report._call_weekly_quality_reviewer_llm([])

        self.assertEqual(len(captured_requests), 2)
        self.assertEqual(
            captured_timeouts,
            [report.WEEKLY_WRITER_TIMEOUT_SECONDS, report.WEEKLY_REVIEW_TIMEOUT_SECONDS],
        )
        self.assertGreaterEqual(report.WEEKLY_WRITER_TIMEOUT_SECONDS, 180)
        self.assertGreaterEqual(report.WEEKLY_REVIEW_TIMEOUT_SECONDS, 180)
        for body in captured_requests:
            self.assertIn(marker, body["messages"][0]["content"])
        self.assertIn("真正重要的变化", captured_requests[0]["messages"][0]["content"])
        self.assertIn("像同一位人类编辑", captured_requests[1]["messages"][0]["content"])


if __name__ == "__main__":
    unittest.main()
