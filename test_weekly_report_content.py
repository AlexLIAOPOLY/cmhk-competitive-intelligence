from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from generate_weekly_report import curated_section, localized_weekly_value
from tts_service import build_audio_summary, normalize_for_speech, prepare_tts_text


class WeeklyReportContentTests(unittest.TestCase):
    def test_political_english_fact_is_localized_to_concrete_chinese(self):
        row = {
            "company": "政治新闻",
            "metric": "重大政策/声明",
            "value": "166 foreign-invested enterprises approved for value-added telecom services",
            "detail": "片段中明确提到'166家外资企业获增值电信业务经营试点批复'",
        }
        value = localized_weekly_value(row, limit=120)
        self.assertIn("166家", value)
        self.assertIn("增值电信业务", value)
        self.assertNotIn("相关动态更新", value)

    def test_regulator_fact_is_classified_as_political(self):
        self.assertEqual(
            curated_section(
                {
                    "company": "通信监管机构",
                    "group": "",
                    "metricCategory": "客户经营",
                }
            ),
            "政治资讯",
        )

    def test_offline_audio_summary_covers_multiple_sections(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "report.docx"
            doc = Document()
            doc.add_paragraph("政治资讯")
            doc.add_paragraph("一、政策事件")
            doc.add_paragraph("工信部批准166家外资企业开展增值电信业务经营试点。")
            doc.add_paragraph("行业资讯")
            doc.add_paragraph("二、友商业绩")
            doc.add_paragraph("香港友商公布最新经营收入及网络投资安排。")
            doc.add_paragraph("国际资讯")
            doc.add_paragraph("三、国际运营商动态")
            doc.add_paragraph("国际运营商继续推进人工智能云和网络基础设施投资。")
            doc.save(path)

            with patch("tts_service._generate_audio_summary_with_llm", return_value=None):
                summary = build_audio_summary(path)

        self.assertIn("政治资讯方面", summary)
        self.assertIn("行业资讯方面", summary)
        self.assertIn("国际资讯方面", summary)
        self.assertGreater(len(summary), 120)

    def test_offline_audio_summary_parses_current_unnumbered_body_layout(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "7月31日周报.docx"
            doc = Document()
            for section, tag, title in (
                ("政治资讯", "政策动向", "监管机构公布新政策"),
                ("经济资讯", "宏观经济", "统计机构公布经济数据"),
                ("行业资讯", "行业动态", "运营商推进网络投资"),
                ("本地运营商资讯", "香港友商", "香港友商公布业务进展"),
            ):
                doc.add_paragraph(section)
                doc.add_paragraph(tag)
                doc.add_paragraph(title)
                doc.add_paragraph(
                    "据公开来源于2026年7月24日发布的信息，相关主体公布了本期可核验进展和关键数据。"
                    "正式披露进一步说明业务背景、实施范围和当前安排。"
                    "后续仍需持续跟进执行节奏及市场影响。"
                )
                doc.add_paragraph("发布时间：2026年7月24日　来源：[S1] 测试来源")
            doc.save(path)

            with patch("tts_service._generate_audio_summary_with_llm", return_value=None):
                summary = build_audio_summary(path)

        self.assertIn("政治资讯方面", summary)
        self.assertIn("经济资讯方面", summary)
        self.assertIn("本地运营商资讯方面", summary)
        self.assertGreaterEqual(len(summary), 220)

    def test_carrier_performance_has_a_detailed_offline_audio_fallback(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "7月24日运营商业绩摘要.docx"
            doc = Document()
            for company in ("中国移动", "中国电信", "中国联通", "中国铁塔", "香港电讯"):
                doc.add_paragraph(f"{company}（最新业绩）关键摘要")
                doc.add_paragraph("1. 派息：全年派息保持增长，股东回报政策维持稳定。")
                doc.add_paragraph("2. 资本开支：资本开支继续聚焦网络、算力和人工智能基础设施。")
                doc.add_paragraph(
                    "3. 战略升级：公司围绕通信服务、算力服务和人工智能服务推进协同转型，"
                    "并持续跟进传统业务压力、经营效率和新业务收入兑现情况。"
                )
                doc.add_paragraph("4. 券商观点：机构关注盈利质量与转型进展。")
                doc.add_paragraph("5. 市场反应：市场继续观察收入、利润及投资回报变化。")
            doc.save(path)

            with patch("tts_service._generate_audio_summary_with_llm", return_value=None):
                summary = build_audio_summary(path)

        self.assertIn("本期运营商业绩摘要重点如下", summary)
        self.assertIn("中国移动方面", summary)
        self.assertIn("香港电讯方面", summary)
        self.assertGreaterEqual(len(summary), 220)

    def test_tts_does_not_treat_15gw_as_a_5g_network_term(self):
        normalized = normalize_for_speech("SKT建设15GW数据中心，同时推进5G网络。")
        prepared = prepare_tts_text(normalized)

        self.assertNotIn("1五GW", normalized)
        self.assertIn("15GW", normalized)
        self.assertIn("五G网络", prepared)


if __name__ == "__main__":
    unittest.main()
