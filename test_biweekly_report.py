from __future__ import annotations

import json
import re
import tempfile
import unittest
from datetime import date, datetime, timedelta
from pathlib import Path
from unittest.mock import patch
from zoneinfo import ZoneInfo

from bs4 import BeautifulSoup
from docx import Document

import generate_weekly_report as report


HKT = ZoneInfo("Asia/Hong_Kong")
FIXED_NOW = datetime(2026, 7, 15, 10, 30, tzinfo=HKT)


def make_row(row_id: str, published_at: str | None, **overrides: object) -> dict:
    detail = (
        "该主体公布最新经营进展，披露核心业务保持增长，并继续投入网络、人工智能及企业服务。"
        "公开资料同时说明客户结构和产品组合正在调整，相关变化将影响后续市场竞争、资源配置与服务策略。"
    )
    row = {
        "id": row_id,
        "rowRef": row_id,
        "company": f"测试主体{row_id}",
        "group": "hong-kong",
        "metricCategory": "客户经营",
        "metric": "5G用户数",
        "value": detail,
        "detail": detail,
        "disclosureDate": published_at or "",
        "generatedAt": "",
        "sourceType": "public-crawl",
        "sources": [
            {
                "label": "测试来源",
                "url": f"https://example.test/{row_id}",
                "publishedAt": published_at or "",
            }
        ],
    }
    row.update(overrides)
    return row


class FrozenDateTime(datetime):
    @classmethod
    def now(cls, tz=None):
        if tz is None:
            return FIXED_NOW.replace(tzinfo=None)
        return FIXED_NOW.astimezone(tz)


class WeeklyIssuePeriodTests(unittest.TestCase):
    @staticmethod
    def write_period_config(path: Path) -> None:
        path.write_text(
            json.dumps(
                {
                    "periodStart": "2026-07-07",
                    "periodEnd": "2026-07-17",
                    "issueDate": "2026-07-17",
                    "cadenceDays": 14,
                }
            ),
            encoding="utf-8",
        )

    def test_current_issue_keeps_planned_range_but_marks_early_run_as_draft(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "period.json"
            self.write_period_config(config_path)
            draft = report.resolve_weekly_period(now=FIXED_NOW, config_path=config_path, environ={})
            final = report.resolve_weekly_period(
                now=datetime(2026, 7, 17, 9, 0, tzinfo=HKT),
                config_path=config_path,
                environ={},
            )

        self.assertEqual(draft.planned_range, {"start": "2026-07-07", "end": "2026-07-17"})
        self.assertEqual(draft.effective_range, {"start": "2026-07-07", "end": "2026-07-15"})
        self.assertEqual(draft.status, "draft")
        self.assertEqual(final.effective_range, final.planned_range)
        self.assertEqual(final.status, "final")

    def test_explicit_issue_boundaries_filter_both_curated_and_future_rows(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "period.json"
            self.write_period_config(config_path)
            period = report.resolve_weekly_period(
                now=datetime(2026, 7, 17, 12, 0, tzinfo=HKT),
                config_path=config_path,
                environ={},
            )
        rows = [
            make_row("before", "2026-07-06"),
            make_row("start", "2026-07-07"),
            make_row("end", "2026-07-17"),
            make_row("after", "2026-07-18"),
        ]

        included, audit = report.filter_biweekly_rows(rows, period=period)

        self.assertEqual([row["id"] for row in included], ["start", "end"])
        self.assertEqual(audit["windowStart"], "2026-07-07")
        self.assertEqual(audit["windowEnd"], "2026-07-17")

    def test_early_draft_rejects_dates_after_its_actual_as_of_day(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "period.json"
            self.write_period_config(config_path)
            period = report.resolve_weekly_period(now=FIXED_NOW, config_path=config_path, environ={})
        rows = [make_row("today", "2026-07-15"), make_row("not_yet", "2026-07-16")]

        included, audit = report.filter_biweekly_rows(rows, period=period)

        self.assertEqual([row["id"] for row in included], ["today"])
        self.assertEqual(audit["windowEnd"], "2026-07-15")
        self.assertEqual(audit["plannedWindowEnd"], "2026-07-17")
        self.assertEqual(audit["periodStatus"], "draft")

    def test_schedule_advances_to_the_next_fourteen_day_issue(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "period.json"
            self.write_period_config(config_path)
            next_period = report.resolve_weekly_period(
                now=datetime(2026, 7, 18, 8, 0, tzinfo=HKT),
                config_path=config_path,
                environ={},
            )

        self.assertEqual(next_period.planned_range, {"start": "2026-07-18", "end": "2026-07-31"})
        self.assertEqual(next_period.effective_range, {"start": "2026-07-18", "end": "2026-07-18"})
        self.assertEqual(next_period.issue_date.date(), date(2026, 7, 31))
        self.assertEqual(next_period.status, "draft")

    def test_early_issue_filename_keeps_its_as_of_date_without_draft_label(self) -> None:
        path = report.dated_weekly_docx_path(
            date(2026, 7, 17),
            draft_as_of=date(2026, 7, 15),
        )
        self.assertTrue(path.name.startswith("7月17日周报（截至7月15日）"))
        self.assertNotIn("草稿", path.name)

    def test_invalid_issue_period_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config_path = Path(temp_dir) / "period.json"
            config_path.write_text(
                json.dumps(
                    {
                        "periodStart": "2026-07-17",
                        "periodEnd": "2026-07-07",
                        "issueDate": "2026-07-17",
                    }
                ),
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "结束日"):
                report.resolve_weekly_period(now=FIXED_NOW, config_path=config_path, environ={})


class BiweeklyDateFilteringTests(unittest.TestCase):
    def test_fourteen_calendar_day_boundary_is_inclusive(self) -> None:
        boundary = (FIXED_NOW - timedelta(days=13)).replace(hour=0, minute=0, second=0, microsecond=0)
        rows = [
            make_row("at_boundary", boundary.isoformat()),
            make_row("one_second_too_old", (boundary - timedelta(seconds=1)).isoformat()),
            make_row("recent", (FIXED_NOW - timedelta(days=2)).isoformat()),
            make_row("future", (FIXED_NOW + timedelta(days=1)).isoformat()),
        ]

        included, audit = report.filter_biweekly_rows(rows, now=FIXED_NOW, window_days=14)

        self.assertEqual([row["id"] for row in included], ["at_boundary", "recent"])
        self.assertIsInstance(audit, dict)

    def test_unknown_publication_date_is_excluded(self) -> None:
        recent_non_publication_timestamp = (FIXED_NOW - timedelta(hours=1)).isoformat()
        unknown = make_row(
            "unknown",
            None,
            disclosureDate="",
            generatedAt=recent_non_publication_timestamp,
            fetched_at_hkt=recent_non_publication_timestamp,
            sources=[{"label": "测试来源", "url": "https://example.test/unknown"}],
        )

        self.assertIsNone(report.resolve_row_publication_date(unknown))
        included, audit = report.filter_biweekly_rows([unknown], now=FIXED_NOW, window_days=14)

        self.assertEqual(included, [])
        self.assertIsInstance(audit, dict)

    def test_report_date_parser_accepts_exact_dates_and_rejects_vague_dates(self) -> None:
        exact_values = [
            "2026-07-02T00:00:00+08:00",
            "2026/07/02",
            "2026年7月2日",
            "July 2, 2026",
            "2 July 2026",
        ]
        for value in exact_values:
            with self.subTest(value=value):
                parsed = report.parse_report_date(value)
                self.assertIsNotNone(parsed)
                parsed_date = parsed.date() if isinstance(parsed, datetime) else parsed
                self.assertEqual(parsed_date, date(2026, 7, 2))

        for value in ("2026年4月后", "持续更新", "2026-07", "2026"):
            with self.subTest(value=value):
                self.assertIsNone(report.parse_report_date(value))

    def test_source_gap_rows_are_excluded_even_with_a_recent_date(self) -> None:
        recent = (FIXED_NOW - timedelta(days=1)).isoformat()
        gap_rows = [
            make_row("gap_status", recent, status="source_gap"),
            make_row("gap_confirmed", recent, qualityStatus="source_gap_confirmed"),
            make_row(
                "gap_text",
                recent,
                metric="公开披露缺口",
                value="未发现可核验公开披露，记录为 source-gap 待后续跟进。",
                detail="未发现可核验公开披露，记录为 source-gap 待后续跟进。",
            ),
        ]
        self.assertTrue(all(report.is_source_gap_row(row) for row in gap_rows))

        included, _ = report.filter_biweekly_rows(
            [make_row("normal", recent), *gap_rows],
            now=FIXED_NOW,
            window_days=14,
        )

        self.assertEqual([row["id"] for row in included], ["normal"])

    def test_curated_model_uses_the_same_fourteen_day_filter(self) -> None:
        boundary = (FIXED_NOW - timedelta(days=13)).replace(hour=0, minute=0, second=0, microsecond=0)
        rows = [
            make_row("included", boundary.isoformat()),
            make_row("stale", (boundary - timedelta(seconds=1)).isoformat()),
            make_row("unknown", None, disclosureDate="", generatedAt=""),
        ]
        with tempfile.TemporaryDirectory() as temp_dir:
            audit_path = Path(temp_dir) / "usage.json"
            with (
                patch.object(report, "datetime", FrozenDateTime),
                patch.object(report, "load_curated_rows", return_value=rows),
                patch.object(report, "WEEKLY_USAGE_AUDIT", audit_path),
                patch.object(
                    report,
                    "enrich_weekly_items_with_llm",
                    side_effect=lambda items, *args, **kwargs: items,
                    create=True,
                ),
            ):
                model = report.build_curated_weekly_model()

        self.assertIsNotNone(model)
        assert model is not None
        self.assertEqual(model["range"], {"start": "2026-07-02", "end": "2026-07-15"})
        self.assertEqual([source["row"] for source in model["sources"]], ["included"])


class BiweeklyContentQualityTests(unittest.TestCase):
    def test_strategic_reference_uses_the_six_required_sections_in_order(self) -> None:
        self.assertEqual(
            report.SECTION_ORDER,
            [
                "政治资讯",
                "经济资讯",
                "行业资讯",
                "本地运营商资讯",
                "社会资讯",
                "国际资讯",
            ],
        )

    def test_short_title_like_text_is_expanded_to_a_real_paragraph(self) -> None:
        short_text = "运营商发布新5G套餐。"

        first = report.ensure_detailed_paragraph(short_text, min_chars=120, max_chars=300)
        second = report.ensure_detailed_paragraph(short_text, min_chars=120, max_chars=300)
        meaningful_length = len(re.sub(r"\s+", "", first))

        self.assertEqual(first, second, "LLM失败回退必须是确定性的")
        self.assertIn(short_text, first)
        self.assertGreaterEqual(meaningful_length, 120)
        self.assertLessEqual(len(first), 300)
        self.assertNotEqual(first.strip(), short_text)

    def test_llm_timeout_repairs_online_then_fails_closed_instead_of_dropping_item(self) -> None:
        original = {
            "id": "item-1",
            "row": "row-1",
            "index": 1,
            "localIndex": 1,
            "tag": "5G套餐",
            "title": "运营商发布新5G套餐",
            "detail": "运营商发布新5G套餐。",
            "eventAt": "2026-07-14T09:00:00+08:00",
            "sourceIds": ["S1", "S2"],
        }
        progress_events: list[tuple[tuple, dict]] = []

        def capture_progress(*args, **kwargs) -> None:
            progress_events.append((args, kwargs))

        repair_search = [
            {
                "id": "W001",
                "query": "repair",
                "provider": "unit",
                "results": [
                    {
                        "title": "官方详情",
                        "url": "https://example.test/detail",
                        "snippet": "运营商发布新5G套餐并披露网络投资安排。",
                    }
                ],
                "error": "",
            }
        ]
        with (
            patch.object(
                report,
                "_call_weekly_writer_llm",
                side_effect=TimeoutError("simulated timeout"),
                create=True,
            ) as llm_call,
            patch.object(report, "run_web_research", return_value=repair_search) as web_search,
        ):
            with self.assertRaisesRegex(RuntimeError, "不能静默删减"):
                report.enrich_weekly_items_with_llm([original], progress=capture_progress)

        self.assertTrue(llm_call.called, "此用例必须真正走到LLM超时回退分支")
        web_search.assert_called_once()
        self.assertTrue(progress_events, "LLM写作、联网修复和失败关闭必须向页面输出可见进度")
        progress_text = " ".join(str(args[0]) for args, _ in progress_events if args)
        self.assertIn("并行修复", progress_text)
        self.assertGreaterEqual(report.WEEKLY_WRITER_RETRY_WORKERS, 4)
        self.assertLessEqual(report.WEEKLY_WRITER_TIMEOUT_SECONDS, 60)

    def test_social_livelihood_rows_are_classified_as_social_news(self) -> None:
        cases = [
            {
                "company": "社会资讯",
                "group": "hong-kong",
                "metricCategory": "社会民生",
                "metric": "失业率",
            },
            {
                "company": "社会资讯",
                "group": "hong-kong",
                "metricCategory": "综合信息",
                "metric": "本地生活咨询",
            },
        ]

        for row in cases:
            with self.subTest(row=row):
                self.assertEqual(report.curated_section(row), "社会资讯")

    def test_macroeconomic_rows_are_classified_as_economic_news(self) -> None:
        cases = [
            {
                "company": "宏观指标",
                "group": "hong-kong",
                "metricCategory": "宏观经济",
                "metric": "GDP",
            },
            {
                "company": "经济资讯",
                "group": "mainland",
                "metricCategory": "经济数据",
                "metric": "CPI",
            },
        ]

        for row in cases:
            with self.subTest(row=row):
                self.assertEqual(report.curated_section(row), "经济资讯")

    def test_hong_kong_operator_rows_are_classified_as_local_operator_news(self) -> None:
        cases = [
            {
                "company": "HKT",
                "group": "hong-kong",
                "metricCategory": "客户经营",
                "metric": "5G用户数",
            },
            {
                "company": "HKBN",
                "group": "hong-kong",
                "metricCategory": "业务动态",
                "metric": "企业业务",
            },
        ]

        for row in cases:
            with self.subTest(row=row):
                self.assertEqual(report.curated_section(row), "本地运营商资讯")

        self.assertEqual(
            report.strategic_section_for_content(
                "行业资讯",
                title="香港宽频公布新一轮企业服务部署",
            ),
            "本地运营商资讯",
        )


class RecentArticleQualityTests(unittest.TestCase):
    def test_explicit_date_parser_does_not_truncate_days_after_twenty(self) -> None:
        cases = {
            "2026-07-23T05:52:23+00:00": "2026-07-23",
            "2026/07/24 09:30": "2026-07-24",
            "Jul 22,2026": "2026-07-22",
            "22 July, 2026": "2026-07-22",
        }

        for value, expected in cases.items():
            with self.subTest(value=value):
                parsed = report._explicit_dates_in_text(value)
                self.assertEqual([item.date().isoformat() for item in parsed], [expected])

    def test_navigation_link_is_not_treated_as_a_news_article(self) -> None:
        soup = BeautifulSoup(
            '<article><time>July 24, 2026</time><h3>National economy update</h3>'
            '<a href="/terms">Terms of Service</a></article>',
            "html.parser",
        )

        self.assertEqual(report._article_title_from_card(soup.a, soup.article), "")

    def test_latest_releases_navigation_title_is_rejected(self) -> None:
        title = "National Bureau of Statistics of China >> Latest Releases"
        self.assertTrue(report._is_navigation_article_title(title))
        self.assertTrue(report._is_navigation_article_title("Latest Release More"))
        self.assertFalse(
            report._cached_recent_article_is_usable(
                {
                    "title": title,
                    "url": "https://www.stats.gov.cn/english/PressRelease/",
                    "publishedAt": "2026-07-24",
                    "rawDetail": "有效正文" * 100,
                }
            )
        )

    def test_partner_article_is_not_admitted_as_independent_news(self) -> None:
        self.assertTrue(
            report._is_promotional_article_evidence(
                "Partner Article This promotional feature discusses a vendor solution."
            )
        )
        self.assertFalse(
            report._cached_recent_article_is_usable(
                {
                    "title": "5G-A: A mobile foundation for embodied AI",
                    "url": "https://totaltele.com/example",
                    "publishedAt": "2026-07-22",
                    "rawDetail": "Partner Article " + "推广正文" * 100,
                }
            )
        )

    def test_generic_contribution_label_uses_the_real_card_heading(self) -> None:
        expected = "The second fiber migration: Why Germany’s FTTH pioneers are moving to XGS-PON"
        soup = BeautifulSoup(
            f'<article><time>July 21, 2026</time><h3>{expected}</h3>'
            '<a href="/article">Contributed Article</a></article>',
            "html.parser",
        )

        self.assertEqual(report._article_title_from_card(soup.a, soup.article), expected)

    def test_article_page_recovers_real_title_from_metadata(self) -> None:
        expected = "The second fiber migration: Why Germany’s FTTH pioneers are moving to XGS-PON"
        html = (
            f'<html><head><meta property="og:title" content="{expected} | Total Telecom"></head>'
            '<body><main><time datetime="2026-07-21">21 July 2026</time>'
            '<h2>Contributed Article</h2><p>'
            + ("Verified network migration evidence. " * 20)
            + "</p></main></body></html>"
        )
        source = {
            "title": "Contributed Article",
            "url": "https://totaltele.com/article",
            "publishedAt": "2026-07-21",
            "section": "行业资讯",
            "row": 9,
        }

        with patch.object(report, "_fetch_public_html", return_value=html):
            item = report._fetch_article_evidence(source)

        self.assertIsNotNone(item)
        self.assertEqual(item["title"], expected)
        self.assertEqual(item["pagePublishedAt"], "2026-07-21")

    def test_article_page_with_mismatched_publication_date_is_rejected(self) -> None:
        html = (
            '<html><head><meta property="article:published_time" content="2007-01-04"></head>'
            '<body><main><h1>National economy update</h1><p>'
            + ("Historical statistical information. " * 20)
            + "</p></main></body></html>"
        )
        source = {
            "title": "National economy update",
            "url": "https://www.stats.gov.cn/english/nbs/200701/example.html",
            "publishedAt": "2026-07-24",
            "section": "经济资讯",
            "row": 7,
        }

        with patch.object(report, "_fetch_public_html", return_value=html):
            self.assertIsNone(report._fetch_article_evidence(source))

    def test_recent_list_page_cannot_be_admitted_as_article(self) -> None:
        source = {
            "title": "Market Prices of Important Means of Production",
            "url": "https://www.stats.gov.cn/english/PressRelease/",
            "publishedAt": "2026-07-24",
        }

        with patch.object(report, "_fetch_public_html") as fetch:
            self.assertIsNone(report._fetch_article_evidence(source))
        fetch.assert_not_called()
        self.assertTrue(report._is_recent_list_url(source["url"]))

    def test_cached_article_requires_matching_page_publication_date(self) -> None:
        article = {
            "title": "Verified telecom infrastructure investment update",
            "url": "https://totaltele.com/verified-update/",
            "publishedAt": "2026-07-23",
            "pagePublishedAt": "2026-07-22",
            "rawDetail": "Verified evidence. " * 30,
        }

        self.assertFalse(report._cached_recent_article_is_usable(article))
        article["pagePublishedAt"] = "2026-07-23"
        self.assertTrue(report._cached_recent_article_is_usable(article))

    def test_dirty_fresh_cache_triggers_live_replacement_discovery(self) -> None:
        now = datetime(2026, 7, 24, 10, 30, tzinfo=HKT)
        replacement = {
            "title": "Verified telecom infrastructure investment update",
            "url": "https://totaltele.com/verified-update/",
            "publishedAt": "2026-07-23",
            "pagePublishedAt": "2026-07-23",
            "section": "行业资讯",
            "tag": "行业动态",
            "row": 9,
            "rawDetail": "Verified evidence. " * 30,
            "detail": "Verified evidence. " * 30,
        }
        cached = {
            "version": report.RECENT_ARTICLE_CACHE_VERSION,
            "fetchedAt": now.isoformat(),
            "windowStart": "2026-07-11",
            "windowEnd": "2026-07-24",
            "articles": [
                {
                    "title": "Terms of Service",
                    "url": "https://example.test/terms",
                    "publishedAt": "2026-07-24",
                    "rawDetail": "navigation " * 30,
                }
            ],
        }
        results = [{"row": 9, "raw_records": [{"url": "https://totaltele.com/"}]}]
        progress: list[str] = []
        with tempfile.TemporaryDirectory() as temp_dir:
            cache_path = Path(temp_dir) / "recent.json"
            cache_path.write_text(json.dumps(cached), encoding="utf-8")
            with (
                patch.object(report, "WEEKLY_EVENT_CACHE", cache_path),
                patch.object(report, "_discover_articles_from_list", return_value=[replacement]) as discover,
                patch.object(report, "_fetch_article_evidence", return_value=replacement),
            ):
                articles, audit = report.discover_recent_articles(
                    results,
                    now=now,
                    progress=progress.append,
                )

        discover.assert_called_once()
        self.assertEqual([item["title"] for item in articles], [replacement["title"]])
        self.assertEqual(audit["cacheRejectedArticles"], 1)
        self.assertTrue(any("重新联网发现正确文章" in message for message in progress))

    def test_unwritable_items_are_replaced_without_reducing_report_count(self) -> None:
        articles = [
            {
                "title": f"Telecom verified article {index:02d}",
                "url": f"https://example.test/article-{index}",
                "publishedAt": "2026-07-23",
                "section": "行业资讯",
                "tag": "行业动态",
                "row": 9,
                "sourceName": "测试来源",
                "rawDetail": f"Verified evidence for article {index}. " * 20,
                "detail": f"Verified evidence for article {index}. " * 20,
            }
            for index in range(1, 8)
        ]
        calls = 0

        def fake_enrich(items, *args, **kwargs):
            nonlocal calls
            calls += 1
            output = [dict(item) for item in items]
            for index, item in enumerate(output):
                item["writerStatus"] = "fallback" if calls == 1 and index < 2 else "llm"
                if item["writerStatus"] == "llm":
                    item["title"] = "已完成写作：" + item["title"]
            return output

        with tempfile.TemporaryDirectory() as temp_dir:
            usage_path = Path(temp_dir) / "usage.json"
            with (
                patch.object(
                    report,
                    "discover_recent_articles",
                    return_value=(articles, {"verifiedArticles": len(articles)}),
                ),
                patch.object(report, "enrich_weekly_items_with_llm", side_effect=fake_enrich),
                patch.object(report, "WEEKLY_USAGE_AUDIT", usage_path),
            ):
                model = report.build_recent_evidence_weekly_model([])
            usage = json.loads(usage_path.read_text(encoding="utf-8"))

        output_items = [item for section in model["sections"] for item in section["items"]]
        self.assertEqual(len(output_items), report.WEEKLY_SECTION_LIMITS["行业资讯"])
        self.assertEqual(usage["dateAudit"]["writerReplacementCount"], 2)
        self.assertEqual(len(model["sources"]), len(output_items))
        self.assertTrue(all(item["writerStatus"] == "llm" for item in output_items))


class BiweeklyDocxDirectoryTests(unittest.TestCase):
    @staticmethod
    def _write_small_template(path: Path) -> None:
        doc = Document()
        doc.add_paragraph("中国移动香港公司")
        doc.add_paragraph("中国移动香港公司战略部    YYYY年M月D日")
        doc.add_paragraph("目 录")
        doc.add_paragraph("")
        for section_name in report.SECTION_ORDER:
            doc.add_paragraph(section_name)
            doc.add_paragraph("1.【标签】一句话事件标题")
        doc.add_paragraph("")
        doc.add_paragraph("政治资讯")
        doc.add_paragraph("标签")
        doc.add_paragraph("一、一句话事件标题")
        doc.add_paragraph("事件事实正文。")
        doc.add_paragraph("")
        doc.save(path)

    @staticmethod
    def _large_model() -> tuple[dict, list[str]]:
        sections = []
        toc = []
        expected_entries = []
        index = 1
        for section_name in report.SECTION_ORDER:
            items = []
            for local_index in range(1, 8):
                title = f"{section_name}测试事件{local_index}"
                item = {
                    "index": index,
                    "localIndex": local_index,
                    "tag": "测试标签",
                    "title": title,
                    "detail": (
                        f"据测试来源于2026年7月15日发布的信息，{title}已通过公开资料确认。"
                        + "该信息反映相关进展。" * 14
                    ),
                    "eventAt": FIXED_NOW.isoformat(),
                    "sourceIds": [f"S{index}"],
                }
                items.append(item)
                toc.append({"index": index, "section": section_name, "tag": item["tag"], "title": title})
                expected_entries.append(f"{report.chinese_order(index)}、{title}")
                index += 1
            sections.append({"name": section_name, "narrative": "", "items": items})
        return (
            {
                "company": "中国移动香港公司",
                "department": "中国移动香港公司战略部",
                "generatedDate": "2026年7月15日",
                "title": "战略内参",
                "range": {"start": "2026-07-02", "end": "2026-07-15"},
                "toc": toc,
                "sections": sections,
                "sources": [],
            },
            expected_entries,
        )

    def test_docx_toc_contains_every_item_even_when_template_slots_are_exhausted(self) -> None:
        model, expected_entries = self._large_model()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "small-template.docx"
            output_path = temp_path / "rendered.docx"
            self._write_small_template(template_path)

            with patch.object(report, "SOURCE_WORD_TEMPLATE", template_path):
                rendered = report.render_into_source_template(model)
            rendered.save(output_path)
            reopened = Document(output_path)

        texts = [paragraph.text.strip() for paragraph in reopened.paragraphs]
        toc_title_index = texts.index("目 录")
        body_start = texts.index("政治资讯", toc_title_index + 1)
        toc_slice = texts[toc_title_index + 1 : body_start]
        toc_sections = [text for text in toc_slice if re.fullmatch(r"【[^\n】]+】", text)]
        toc_entries = [
            text
            for text in toc_slice
            if re.match(r"^[一二三四五六七八九十百零]+、", text)
        ]

        self.assertEqual(toc_sections, [f"【{name}】" for name in report.SECTION_ORDER])
        self.assertEqual(toc_entries, expected_entries)
        self.assertTrue(all("【测试标签】" not in entry for entry in toc_entries))
        self.assertTrue(all(not re.match(r"^\d+[.、]", entry) for entry in toc_entries))


if __name__ == "__main__":
    unittest.main()
