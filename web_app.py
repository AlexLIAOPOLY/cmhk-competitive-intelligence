from __future__ import annotations

import base64
import csv
import ipaddress
import json
import mimetypes
import os
import queue
import re
import subprocess
import sys
import threading
import time
import urllib.error
import urllib.request
import uuid
from datetime import datetime
from html import escape
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse

import crawl
from crawl_run_registry import (
    heartbeat_crawl_run,
    latest_crawl_run_summary,
    load_crawl_run_log,
    load_index as load_crawl_run_index,
    load_run_history as load_crawl_run_history,
    mark_crawl_run_interrupted,
    reconcile_interrupted_crawl_runs,
    register_crawl_run,
    start_crawl_run,
)
from ai_config import INTERNAL_AI_BASE_URL, is_internal_ai_base_url, load_ai_config, save_ai_config
from ai_rate_limit import reset_internal_ai_priority, set_internal_ai_priority, wait_for_internal_ai_slot
from company_metrics import build_company_metrics_payload
from executive_company_benchmarks import build_company_benchmarks
from extractors import row_fields
from rag_llm import ask_llm_with_rag, estimate_tokens, list_knowledge_datasets, stream_llm_with_rag
from agent import available_agent_skills, stream_agent
from agent_memory import delete_memory, load_memories
from agent_production import dataset_lineage, list_agent_runs
from chart_renderer import generated_chart_path
from tts_service import (
    AUDIO_DIR,
    audio_info_for_report,
    audio_paths_for_report,
    delete_audio_for_report,
    rename_audio_for_report,
    synthesize_report_audio,
)
from subscription_service import (
    FREQUENCY_LABELS,
    NEWS_CATEGORY_LABELS,
    REPORT_CADENCE_LABEL,
    REPORT_MODE_LABELS,
    SubscriptionService,
    encode_strategic_news_digest,
    filter_news_by_categories,
    news_category_summary,
)
from cmhk_auth import AuthService
from project_monitor_card_actions import CardActionHandler


ROOT = Path(__file__).resolve().parent
AUTH = AuthService(ROOT)
CRAWL_PIPELINE_LOCK = threading.Lock()
CRAWL_PIPELINE_STATE: dict[str, object] = {}
INTELLIGENCE_INSIGHT_REFRESH_LOCK = threading.Lock()
SCHEDULER_OVERVIEW_LOCK = threading.Lock()
SCHEDULER_OVERVIEW_CACHE: dict[str, object] = {}
SCHEDULER_OVERVIEW_CACHE_SECONDS = 90
TASK_HEARTBEAT_INTERVAL_SECONDS = 10
STATIC_DIR = ROOT / "web" / "static"
COMPETITOR_WORKBENCH_DATA_PATH = STATIC_DIR / "competitor-workbench-data.json"
RESULTS_DIR = ROOT / "results"
CURATION_LATEST_PATH = ROOT / "curation_data" / "latest.json"
CURATION_CANDIDATE_FACTS_PATH = ROOT / "curation_data" / "candidate_facts.jsonl"
CURATION_AGENT_TRACE_PATH = ROOT / "curation_data" / "agent_trace.jsonl"
STRATEGIC_BRIEFING_DIR = ROOT / "strategy_briefing"
STRATEGIC_BRIEFING_RUNS_DIR = STRATEGIC_BRIEFING_DIR / "runs"
LOCAL_TEMPLATE_PATH = Path("/Users/liaowang/Downloads/模板.docx")
REPO_TEMPLATE_PATH = ROOT / "weekly_report_template.docx"
TEMPLATE_PATH = LOCAL_TEMPLATE_PATH if LOCAL_TEMPLATE_PATH.exists() else REPO_TEMPLATE_PATH
REPORT_FILE_RE = re.compile(
    r"^\d{1,2}月\d{1,2}日周报(?:（(?:草稿，)?截至\d{1,2}月\d{1,2}日）)?(?: \(\d+\))?\.docx$"
)
REPORT_METADATA_PATH = ROOT / "report_file_metadata.json"
EXCLUDED_REPORT_NAMES = {
    "test_out.docx",
    "weekly_report.docx",
    "weekly_report_from_word_template.docx",
    "weekly_report_template.docx",
    "carrier_performance_template.docx",
    "模板.docx",
}
REFERENCE_FILES = {"weekly_report.md", "weekly_report.html", "final_audit.md", "coverage_report.tsv", "run_log.tsv"}
UPLOAD_DATASET_PREFIX = "user-upload"
UPLOAD_ALLOWED_SUFFIXES = {".txt", ".md", ".csv", ".tsv", ".json", ".docx", ".pdf"}
UPLOAD_MAX_BYTES = 8 * 1024 * 1024
CHAT_IMAGE_MAX_BYTES = 8 * 1024 * 1024
CHAT_AUDIO_MAX_BYTES = 20 * 1024 * 1024
CHAT_STT_MODEL = (os.environ.get("CMHK_STT_MODEL") or "Qwen3ASR").strip()
CHAT_AUDIO_MIME_EXTENSIONS = {
    "audio/webm": "webm",
    "audio/mp4": "m4a",
    "audio/mpeg": "mp3",
    "audio/mpga": "mpga",
    "audio/m4a": "m4a",
    "audio/x-m4a": "m4a",
    "audio/wav": "wav",
    "audio/x-wav": "wav",
}
CHAT_THREADS_DIR = ROOT / "agent_chat_threads"
CHAT_THREADS_PATH = CHAT_THREADS_DIR / "threads.json"
CHAT_THREADS_LOCK = threading.Lock()
CHAT_TITLE_TASK_LOCK = threading.Lock()
CHAT_TITLE_PENDING: dict[str, str] = {}
CHAT_TITLE_ACTIVE: set[str] = set()
CHAT_APPROVAL_LOCK = threading.Lock()
CHAT_APPROVAL_WAITERS: dict[tuple[str, str], dict[str, object]] = {}
CHAT_STARTER_POOL = (
    {
        "icon": "performance",
        "tone": "blue",
        "title": "查询香港资费数据",
        "detail": "例：csl、3HK、SmarTone 5G 月费",
        "prompt": "请查询并对比香港 csl、1O1O、3HK 和 SmarTone 当前 5G 套餐的月费、本地数据量、合约期和促销优惠，用表格展示并标明来源。",
    },
    {
        "icon": "performance",
        "tone": "violet",
        "title": "分析香港电信趋势",
        "detail": "例：移动用户、数据用量与宽频接入",
        "prompt": "请基于香港官方与已核验数据，分析近十年移动用户数、移动数据用量、宽频接入线和 5G 发展趋势，说明增速、拐点及其对香港电信市场的影响，并用图表展示。",
    },
    {
        "icon": "cloud",
        "tone": "teal",
        "title": "预测香港市场走势",
        "detail": "例：移动数据用量、5G 与宽频趋势",
        "prompt": "请基于已核验的香港历史数据，预测未来四个季度的移动数据用量、移动用户数、5G 发展和宽频接入趋势，给出基准、乐观和谨慎情景，并说明预测依据、不确定性和风险。",
    },
    {
        "icon": "policy",
        "tone": "orange",
        "title": "解读香港政策影响",
        "detail": "例：频谱分配、SIM 实名制与消费者保护",
        "prompt": "请梳理香港近期频谱分配、5G、SIM 实名制、电信监管和消费者保护政策的变化，分析对 CMHK 产品、网络、营销和运营的影响，并提出应对建议。",
    },
)


def subscription_service() -> SubscriptionService:
    return SubscriptionService(runtime_root=ROOT)


def is_loopback_client(address: str) -> bool:
    try:
        parsed = ipaddress.ip_address(address)
    except ValueError:
        return False
    if isinstance(parsed, ipaddress.IPv6Address) and parsed.ipv4_mapped:
        parsed = parsed.ipv4_mapped
    return parsed.is_loopback


def sample_chat_starters(limit: int = 4) -> list[dict]:
    count = max(1, min(int(limit), len(CHAT_STARTER_POOL)))
    return [dict(item) for item in CHAT_STARTER_POOL[:count]]


def request_runtime_context(handler: BaseHTTPRequestHandler) -> dict:
    now = datetime.now().astimezone()
    client_ip = ""
    try:
        client_ip = str(handler.client_address[0] or "")
    except Exception:
        client_ip = ""
    forwarded_for = str(handler.headers.get("X-Forwarded-For") or "").split(",")[0].strip()
    real_ip = str(handler.headers.get("X-Real-IP") or "").strip()
    visible_ip = forwarded_for or real_ip or client_ip or "unknown"
    if visible_ip in {"127.0.0.1", "::1", "localhost"} or visible_ip.startswith(("10.", "192.168.", "172.16.", "172.17.", "172.18.", "172.19.", "172.20.", "172.21.", "172.22.", "172.23.", "172.24.", "172.25.", "172.26.", "172.27.", "172.28.", "172.29.", "172.30.", "172.31.")):
        location_hint = "本机或内网访问；按服务端时区和用户工作环境推断为 Hong Kong / Asia_Hong_Kong"
    else:
        location_hint = "公网 IP；未接入第三方 GeoIP，不能精确到城市"
    return {
        "current_time": now.isoformat(timespec="seconds"),
        "timezone": now.tzname() or "local",
        "utc_offset": now.strftime("%z"),
        "client_ip": client_ip,
        "forwarded_for": forwarded_for,
        "real_ip": real_ip,
        "visible_ip": visible_ip,
        "location_hint": location_hint,
    }


def analyze_chat_image(payload: dict) -> dict:
    config = load_ai_config(include_key=True)
    requested_model = str(payload.get("model") or config.get("model") or "").strip()
    vision_pattern = r"(?:vision|multimodal|omni|(?:^|[-_.])vl(?:[-_.]|$)|qwen[^/]*vl|internvl|llava|gpt-4o|gpt-4\.1|gemini|claude-3|kimi[-_.]?k2\.5)"
    model = requested_model
    if not re.search(vision_pattern, model, flags=re.I):
        model = str(os.environ.get("CMHK_CHAT_IMAGE_MODEL") or "Kimi-K2.5").strip()
    if not re.search(vision_pattern, model, flags=re.I):
        raise ValueError("未配置可用的图片识别模型")
    data_url = str(payload.get("image") or "").strip()
    match = re.fullmatch(r"data:(image/(?:png|jpeg|webp|gif));base64,([A-Za-z0-9+/=\s]+)", data_url, flags=re.I)
    if not match:
        raise ValueError("只支持 PNG、JPG、WebP 或 GIF 图片")
    raw = base64.b64decode(re.sub(r"\s+", "", match.group(2)), validate=True)
    if not raw or len(raw) > CHAT_IMAGE_MAX_BYTES:
        raise ValueError("图片不能为空且不能超过 8 MB")
    base_url = str(config.get("base_url") or "").strip().rstrip("/")
    api_key = str(config.get("api_key") or "").strip()
    if not is_internal_ai_base_url(base_url) or not api_key:
        raise ValueError("公司内网模型配置不完整")
    question = str(payload.get("question") or "请分析这张图片").strip()[:1200]
    body = {
        "model": model,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": f"请准确识别图片中与下列问题相关的内容，输出可供后续分析的中文事实描述，不要猜测。问题：{question}"},
                {"type": "image_url", "image_url": {"url": data_url}},
            ],
        }],
        "max_tokens": 900,
        "stream": False,
    }
    request = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    wait_for_internal_ai_slot("chat-image-analyze")
    with urllib.request.urlopen(request, timeout=90) as response:
        result = json.loads(response.read().decode("utf-8"))
    message = ((result.get("choices") or [{}])[0].get("message") or {})
    content = message.get("content") or message.get("reasoning_content") or ""
    if isinstance(content, list):
        content = "\n".join(str(item.get("text") or "") for item in content if isinstance(item, dict))
    elif isinstance(content, dict):
        content = content.get("text") or content.get("content") or ""
    description = str(content).strip()
    if not description:
        raise ValueError("视觉模型没有返回可用的图片描述")
    return {"description": description, "model": model}


def _competitor_insight_content(content: object) -> str:
    """Normalize the model response container without applying semantic gates."""
    if isinstance(content, list):
        text = "".join(str(item.get("text") or "") for item in content if isinstance(item, dict))
    elif isinstance(content, dict):
        text = str(content.get("text") or content.get("content") or "")
    else:
        text = str(content or "")
    if not text.strip():
        raise RuntimeError("AI 未返回可用洞察")
    return text


def _parse_competitor_insight_items(content: object) -> list[str]:
    """Best-effort three-row display parsing; format drift must not reject a usable answer."""
    text = _competitor_insight_content(content).strip()
    text = re.sub(r"^```(?:json|text|markdown)?\s*|\s*```$", "", text, flags=re.I)
    labels = ("竞争格局", "公司定位", "业务含义")
    aliases = {"竞争格局": 0, "公司分化": 1, "公司定位": 1, "业务含义": 2}
    labelled: dict[int, str] = {}
    candidates: list[str] = []
    for raw_line in text.splitlines():
        line = re.sub(r"^\s*(?:#{1,6}\s*|[-*•]\s+|\d+[.)、]\s*)", "", raw_line).strip()
        line = re.sub(r"^\*\*(.*?)\**$", r"\1", line).strip()
        if not line or re.fullmatch(r"\|?\s*:?-{2,}[-| :]*", line):
            continue
        match = re.match(r"^(?:一|二|三)?[、.\s]*(竞争格局|公司分化|公司定位|业务含义)[：|｜]\s*(.+)$", line)
        if match:
            labelled.setdefault(aliases[match.group(1)], match.group(2).strip())
            continue
        if len(line) < 12 and not re.search(r"[\d。！？!?；;，,]", line):
            continue
        if not (line.startswith("|") and line.endswith("|")):
            candidates.append(line)
    if len(candidates) < 3:
        sentences = [part.strip() for part in re.split(r"(?<=[。！？!?])\s*", " ".join(candidates) or text) if part.strip()]
        if len(sentences) > len(candidates):
            candidates = sentences
    result: list[str] = []
    candidate_index = 0
    for index, label in enumerate(labels):
        value = labelled.get(index, "")
        while not value and candidate_index < len(candidates):
            candidate = candidates[candidate_index]
            candidate_index += 1
            value = candidate
        if not value:
            continue
        value = re.sub(r"^(?:竞争格局|公司分化|公司定位|业务含义)[：|｜]\s*", "", value).strip()
        if len(value) > 180:
            value = value[:179].rstrip("，,；; ") + "…"
        result.append(f"{label}｜{value}")
    return result


def generate_competitor_insight(payload: dict, stream_callback=None) -> dict:
    request_id = str(payload.get("requestId") or "")[:80]
    companies = [str(value)[:80] for value in (payload.get("companies") or []) if str(value).strip()]
    metric = payload.get("metric") if isinstance(payload.get("metric"), dict) else {}
    years = [int(value) for value in (payload.get("years") or [])]
    if len(set(companies)) != len(companies) or len(set(years)) != len(years):
        raise ValueError("竞对或年份包含重复选择")
    if not (2 <= len(companies) <= 6) or not (2 <= len(years) <= 10):
        raise ValueError("竞对、年份或表格数据不完整")
    metric_key = str(metric.get("key") or "")[:120]
    if not metric_key or not COMPETITOR_WORKBENCH_DATA_PATH.exists():
        raise ValueError("竞对指标或权威数据集不可用")
    canonical = json.loads(COMPETITOR_WORKBENCH_DATA_PATH.read_text(encoding="utf-8"))
    evidence_version = str(payload.get("evidenceVersion") or "")
    if evidence_version and evidence_version != str(canonical.get("evidenceVersion") or ""):
        raise ValueError("竞对数据版本已更新，请刷新后重试")
    allowed = set(companies)
    normalized = [
        {
            "company": str(row.get("company") or "")[:80],
            "year": int(row.get("year") or 0),
            "value": float(row.get("value")),
            "unit": str(row.get("unit") or "")[:80],
            "comparator": str(row.get("comparator") or "=")[:8],
            "period": str(row.get("period") or "")[:40],
            "period_end": str(row.get("periodEnd") or "")[:20],
            "scope": str(row.get("scope") or "")[:300],
            "basis": str(row.get("basis") or "")[:120],
            "status": str(row.get("status") or "")[:80],
            "source": str(row.get("source") or "")[:500],
            "note": str(row.get("note") or "")[:500],
        }
        for row in (canonical.get("cells") or [])
        if isinstance(row, dict)
        and str(row.get("company") or "") in allowed
        and str(row.get("metric") or "") == metric_key
        and int(row.get("year") or 0) in years
    ]
    normalized.sort(key=lambda row: (row["year"], row["company"]))
    if not normalized or len(normalized) > 80:
        raise ValueError("当前比较范围没有可用的权威数据")
    if len({row["unit"] for row in normalized}) != 1:
        raise ValueError("所选数据单位不一致，不能直接比较")
    per_company_years = {company: {row["year"] for row in normalized if row["company"] == company} for company in companies}
    if any(len(company_years) < 2 for company_years in per_company_years.values()):
        raise ValueError("每家竞对至少需要两个有效年度")
    common_years = sorted(set.intersection(*(set(value) for value in per_company_years.values())))
    if len(common_years) < 2:
        raise ValueError("共同可比年度不足两个，暂不生成 AI 解析")
    comparison_rows = [row for row in normalized if row["year"] in common_years]
    table = "\n".join(
        "\t".join(str(row[key]) for key in ("company", "year", "comparator", "value", "unit", "period", "period_end", "scope", "basis", "status", "source", "note"))
        for row in comparison_rows
    )
    config = load_ai_config(include_key=True)
    base_url = str(config.get("base_url") or INTERNAL_AI_BASE_URL).strip().rstrip("/")
    api_key = str(config.get("api_key") or "").strip()
    model = str(config.get("model") or "").strip()
    if not base_url or not api_key or not model or not is_internal_ai_base_url(base_url):
        raise RuntimeError("公司内网 AI 配置不完整")
    metric_label = str(metric.get("label") or metric_key)[:120]
    body = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是电信行业竞争策略分析师。只输出三行精炼的简体中文纯文本，每行约45—90字，禁止前言、编号、Markdown标题、表格和结语。三行依次以‘竞争格局｜’‘公司定位｜’‘业务含义｜’开头，基于表格数据判断趋势、公司间位置与业务意义。保留大于、至少、约等原始比较符，共建共享数值不得相加，不得补数或引用外部知识。"},
            {"role": "user", "content": f"指标：{metric_label}\n证据版本：{str(canonical.get('evidenceVersion') or '')}\n所选公司：{'、'.join(companies)}\n共同数据年度：{','.join(str(year) for year in common_years)}\n列：公司、年度、比较符、数值、单位、披露期、期末日、范围、口径、核验状态、官方来源、备注\n{table}"},
        ],
        "temperature": 0.1,
        "max_tokens": 500,
        "chat_template_kwargs": {"enable_thinking": False},
        "stream": stream_callback is not None,
    }
    request = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    priority_token = set_internal_ai_priority("interactive")
    try:
        if stream_callback:
            stream_callback({"type": "status", "stage": "queue", "message": "请求已进入内网 AI 前台队列"})
        wait_for_internal_ai_slot(
            "competitor-insight",
            wait_callback=(
                lambda remaining: stream_callback(
                    {
                        "type": "status",
                        "stage": "queue",
                        "message": f"内网 AI 繁忙，已为本次洞察保留队列（约 {max(1, round(remaining))} 秒）",
                    }
                )
                if stream_callback
                else None
            ),
        )
        if stream_callback:
            stream_callback({"type": "status", "stage": "generating", "message": "内网 AI 已连接，正在生成真实结果"})
        with urllib.request.urlopen(request, timeout=90 if stream_callback else 60) as response:
            if stream_callback:
                content_parts = []
                reasoning_started = False
                for raw_line in response:
                    line = raw_line.decode("utf-8", errors="replace").strip()
                    if not line.startswith("data:"):
                        continue
                    payload_text = line.removeprefix("data:").strip()
                    if not payload_text or payload_text == "[DONE]":
                        continue
                    try:
                        event = json.loads(payload_text)
                    except json.JSONDecodeError:
                        continue
                    delta = ((event.get("choices") or [{}])[0].get("delta") or {})
                    reasoning_delta = delta.get("reasoning_content")
                    if reasoning_delta and not reasoning_started:
                        reasoning_started = True
                        stream_callback({"type": "status", "stage": "reasoning", "message": "内网 AI 正在分析所选数据"})
                    content_delta = delta.get("content")
                    if isinstance(content_delta, str) and content_delta:
                        content_parts.append(content_delta)
                        stream_callback({"type": "delta", "text": content_delta})
                raw_content = "".join(content_parts)
            else:
                result = json.loads(response.read().decode("utf-8"))
                message = ((result.get("choices") or [{}])[0].get("message") or {})
                raw_content = message.get("content") or message.get("reasoning_content") or ""
    finally:
        reset_internal_ai_priority(priority_token)
    insight = _competitor_insight_content(raw_content)
    insights = _parse_competitor_insight_items(insight)
    return {"requestId": request_id, "insight": insight, "insights": insights, "model": model}


def transcribe_chat_audio(payload: dict) -> dict:
    data_url = str(payload.get("audio") or "").strip()
    match = re.fullmatch(
        r"data:(audio/(?:webm|mp4|mpeg|mpga|m4a|x-m4a|wav|x-wav))(?:;codecs=[^;,]+)?;base64,([A-Za-z0-9+/=\s]+)",
        data_url,
        flags=re.I,
    )
    if not match:
        raise ValueError("只支持 WebM、M4A、MP3 或 WAV 语音")
    try:
        raw = base64.b64decode(re.sub(r"\s+", "", match.group(2)), validate=True)
    except Exception as exc:
        raise ValueError("语音数据格式无效") from exc
    if not raw:
        raise ValueError("没有录到可识别的语音")
    if len(raw) > CHAT_AUDIO_MAX_BYTES:
        raise ValueError("单次语音不能超过 20 MB")

    mime_type = match.group(1).lower()
    extension = CHAT_AUDIO_MIME_EXTENSIONS.get(mime_type)
    if not extension:
        raise ValueError("不支持当前录音格式")
    try:
        converted = subprocess.run(
            [
                os.environ.get("CMHK_FFMPEG_BIN") or "ffmpeg",
                "-hide_banner",
                "-loglevel",
                "error",
                "-i",
                "pipe:0",
                "-vn",
                "-ac",
                "1",
                "-ar",
                "16000",
                "-c:a",
                "pcm_s16le",
                "-f",
                "wav",
                "pipe:1",
            ],
            input=raw,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=20,
            check=False,
        )
    except FileNotFoundError as exc:
        raise ValueError("服务器音频解码组件未就绪，请联系管理员安装 FFmpeg") from exc
    except subprocess.TimeoutExpired as exc:
        raise ValueError("录音解码超时，请缩短录音后重试") from exc
    if converted.returncode != 0 or len(converted.stdout) <= 44:
        raise ValueError("录音格式无法解码，请重新录音")
    raw = converted.stdout
    mime_type = "audio/wav"
    extension = "wav"

    config = load_ai_config(include_key=True)
    base_url = str(config.get("base_url") or "").strip().rstrip("/")
    api_key = str(config.get("api_key") or "").strip()
    if not is_internal_ai_base_url(base_url) or not api_key:
        raise ValueError("公司内网模型配置不完整")

    boundary = f"----CMHKVoice{uuid.uuid4().hex}"
    line_break = b"\r\n"
    body = bytearray()
    for name, value in (("model", CHAT_STT_MODEL), ("language", "zh")):
        body.extend(f"--{boundary}".encode("ascii") + line_break)
        body.extend(f'Content-Disposition: form-data; name="{name}"'.encode("ascii") + line_break + line_break)
        body.extend(str(value).encode("utf-8") + line_break)
    body.extend(f"--{boundary}".encode("ascii") + line_break)
    body.extend(
        f'Content-Disposition: form-data; name="file"; filename="voice.{extension}"'.encode("ascii")
        + line_break
    )
    body.extend(f"Content-Type: {mime_type}".encode("ascii") + line_break + line_break)
    body.extend(raw + line_break)
    body.extend(f"--{boundary}--".encode("ascii") + line_break)

    request = urllib.request.Request(
        f"{base_url}/audio/transcriptions",
        data=bytes(body),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
            "Accept": "application/json",
        },
        method="POST",
    )
    wait_for_internal_ai_slot("chat-audio-transcription")
    with urllib.request.urlopen(request, timeout=90) as response:
        result = json.loads(response.read().decode("utf-8"))
    transcript = str(result.get("text") or result.get("transcript") or "").strip()
    if not transcript:
        raise ValueError("语音模型没有识别出文字，请靠近麦克风后重试")
    if transcript.rstrip("。！？!?，, ").strip() in {"嗯", "呃", "啊", "唔"}:
        raise ValueError("只识别到很短的语气词，请确认麦克风输入后靠近说话并重试")
    return {"text": transcript, "model": CHAT_STT_MODEL}


def _now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _clean_chat_message(item: dict) -> dict | None:
    if not isinstance(item, dict):
        return None
    role = "assistant" if item.get("role") == "assistant" else "user"
    content = str(item.get("content") or "").strip()
    if not content:
        return None
    clean = {"role": role, "content": content[:20000]}
    for timestamp_key in ("createdAt", "completedAt"):
        timestamp = str(item.get(timestamp_key) or "").strip()
        if timestamp:
            try:
                normalized = timestamp[:-1] + "+00:00" if timestamp.endswith(("Z", "z")) else timestamp
                datetime.fromisoformat(normalized)
            except ValueError:
                continue
            clean[timestamp_key] = timestamp[:40]
    if role == "user":
        display_content = str(item.get("displayContent") or "").strip()
        if display_content:
            clean["displayContent"] = display_content[:4000]
        image_preview = item.get("imagePreview")
        if isinstance(image_preview, dict):
            data_url = str(image_preview.get("dataUrl") or "")
            if len(data_url) <= 1_500_000 and re.match(r"^data:image/(?:png|jpeg|webp|gif);base64,", data_url, flags=re.I):
                clean["imagePreview"] = {
                    "name": str(image_preview.get("name") or "已发送图片")[:200],
                    "dataUrl": data_url,
                }
    else:
        references = item.get("references")
        links = item.get("links")
        suggestions = item.get("suggestions")
        timeline = item.get("timeline")
        metrics = item.get("metrics")
        if isinstance(references, list):
            clean["references"] = references[:30]
        if isinstance(links, list):
            clean["links"] = links[:30]
        if isinstance(suggestions, list):
            clean["suggestions"] = [str(s).strip()[:160] for s in suggestions if str(s).strip()][:3]
        if isinstance(timeline, list):
            clean_timeline = []
            timeline_size = 0
            for raw_event in timeline[:240]:
                if not isinstance(raw_event, dict):
                    continue
                event_type = str(raw_event.get("type") or "")
                if event_type == "text":
                    event = {"type": "text", "text": str(raw_event.get("text") or "")[:20000]}
                    if not event["text"]:
                        continue
                elif event_type in {"tool_call_start", "tool_call_result"}:
                    event = {
                        "type": event_type,
                        "id": str(raw_event.get("id") or "")[:160],
                        "name": str(raw_event.get("name") or "")[:120],
                        "processText": str(raw_event.get("processText") or "")[:500],
                        "args": str(raw_event.get("args") or "")[:6000],
                        "content": str(raw_event.get("content") or "")[:16000],
                    }
                else:
                    continue
                event_size = len(json.dumps(event, ensure_ascii=False))
                if timeline_size + event_size > 120000:
                    break
                timeline_size += event_size
                clean_timeline.append(event)
            if clean_timeline:
                clean["timeline"] = clean_timeline
        if isinstance(metrics, dict):
            clean["metrics"] = {
                "inputTokens": max(0, int(metrics.get("inputTokens") or 0)),
                "outputTokens": max(0, int(metrics.get("outputTokens") or 0)),
                "totalTokens": max(0, int(metrics.get("totalTokens") or 0)),
                "durationMs": max(0, int(metrics.get("durationMs") or 0)),
                "estimated": bool(metrics.get("estimated")),
            }
    return clean


def load_chat_threads() -> list[dict]:
    if not CHAT_THREADS_PATH.exists():
        return []
    try:
        data = json.loads(CHAT_THREADS_PATH.read_text(encoding="utf-8"))
    except Exception:
        return []
    threads = data.get("threads") if isinstance(data, dict) else data
    if not isinstance(threads, list):
        return []
    return [item for item in threads if isinstance(item, dict) and item.get("id")]


def save_chat_threads(threads: list[dict]) -> None:
    CHAT_THREADS_DIR.mkdir(parents=True, exist_ok=True)
    CHAT_THREADS_PATH.write_text(
        json.dumps({"threads": threads[:200]}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def chat_thread_summaries() -> list[dict]:
    threads = sorted(load_chat_threads(), key=lambda item: str(item.get("updatedAt") or ""), reverse=True)
    threads = sorted(threads, key=lambda item: 0 if item.get("pinned") else 1)
    summaries = []
    for thread in threads:
        messages = thread.get("messages") if isinstance(thread.get("messages"), list) else []
        last = next((m for m in reversed(messages) if isinstance(m, dict) and m.get("content")), {})
        preview = str(last.get("content") or "")
        preview = re.sub(r"[*_`#]+", "", preview)
        preview = re.sub(r"!?\[([^\]]*)\]\([^)]*\)", r"\1", preview)
        preview = re.sub(r"\s+", " ", preview).strip()[:120]
        summaries.append(
            {
                "id": thread.get("id"),
                "title": thread.get("title") or "未命名对话",
                "createdAt": thread.get("createdAt"),
                "updatedAt": thread.get("updatedAt"),
                "messageCount": len(messages),
                "preview": preview,
                "pinned": bool(thread.get("pinned")),
            }
        )
    return summaries


def _sanitize_thread_title(raw: str) -> str:
    title = re.sub(r"^[\"'“”‘’\s]+|[\"'“”‘’\s]+$", "", str(raw or ""))
    title = re.sub(r"^(标题|对话标题|主题)[:：]\s*", "", title)
    title = re.sub(r"[\r\n\t]+", " ", title)
    title = re.sub(r"\s+", " ", title).strip(" -_，。,.")
    return title[:24] or "新对话"


def _fallback_thread_title(first_user: str) -> str:
    text = re.sub(r"\s+", " ", str(first_user or "")).strip()
    if text in {"你好", "您好", "hi", "hello", "看看", "测试"}:
        return "初次咨询"
    text = re.sub(r"^(请|帮我|麻烦|能不能|可以|给我)", "", text).strip()
    return _sanitize_thread_title(text[:18] or "新对话")


def _thread_title_source(messages: list[dict]) -> str:
    generic = {"你好", "您好", "hi", "hello", "看看", "测试"}
    users = [str(item.get("content") or "").strip() for item in messages if item.get("role") == "user"]
    for text in users:
        normalized = re.sub(r"\s+", " ", text).strip()
        if len(normalized) >= 6 and normalized.lower() not in generic:
            return normalized
    return users[0] if users else ""


def generate_chat_thread_title(first_user: str) -> str:
    first_user = str(first_user or "").strip()
    if not first_user:
        return "新对话"
    config = load_ai_config(include_key=True)
    api_key = str(config.get("api_key") or "").strip()
    if not api_key:
        return _fallback_thread_title(first_user)
    provider = str(config.get("provider") or "deepseek").lower()
    model = (
        os.environ.get("CMHK_CHAT_TITLE_MODEL", "").strip()
        or "Qwen3-30B-A3B-Instruct-2507"
    )
    base_url = str(config.get("base_url") or INTERNAL_AI_BASE_URL).rstrip("/")
    prompt = (
        "请根据用户第一条问题生成一个中文对话标题。"
        "要求：6到12个汉字或短词；不要照抄原句；不要加引号、标点、解释或前缀。\n\n"
        f"用户第一问：{first_user[:500]}"
    )
    body = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你只输出简洁中文标题。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.1,
        "max_tokens": 48,
    }
    req = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        wait_for_internal_ai_slot("chat-thread-title")
        with urllib.request.urlopen(req, timeout=12) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        content = str(payload.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        if not content:
            return _fallback_thread_title(first_user)
        title = _sanitize_thread_title(content)
        if title in {"新对话", "未命名对话"}:
            return _fallback_thread_title(first_user)
        return title
    except Exception as exc:
        print(f"chat thread title generation failed: {exc}", flush=True)
        return _fallback_thread_title(first_user)


def _schedule_chat_thread_title(thread_id: str, title_source: str) -> None:
    """Refresh an AI chat title in the background without blocking message saves."""
    source = str(title_source or "").strip()
    if not thread_id or not source:
        return
    with CHAT_TITLE_TASK_LOCK:
        CHAT_TITLE_PENDING[thread_id] = source
        if thread_id in CHAT_TITLE_ACTIVE:
            return
        CHAT_TITLE_ACTIVE.add(thread_id)

    def worker() -> None:
        try:
            while True:
                with CHAT_TITLE_TASK_LOCK:
                    current_source = CHAT_TITLE_PENDING.pop(thread_id, "")
                if not current_source:
                    return
                generated_title = generate_chat_thread_title(current_source)
                with CHAT_TITLE_TASK_LOCK:
                    has_newer_source = bool(CHAT_TITLE_PENDING.get(thread_id))
                with CHAT_THREADS_LOCK:
                    threads = load_chat_threads()
                    target = next((item for item in threads if str(item.get("id")) == thread_id), None)
                    if target and target.get("titlePending"):
                        target["title"] = generated_title
                        target["titlePending"] = has_newer_source
                        save_chat_threads(threads)
                if not has_newer_source:
                    return
        finally:
            with CHAT_TITLE_TASK_LOCK:
                CHAT_TITLE_ACTIVE.discard(thread_id)
                should_restart = bool(CHAT_TITLE_PENDING.get(thread_id))
            if should_restart:
                _schedule_chat_thread_title(thread_id, CHAT_TITLE_PENDING.get(thread_id, ""))

    threading.Thread(target=worker, name=f"chat-title-{thread_id[:8]}", daemon=True).start()


def get_chat_thread(thread_id: str) -> dict | None:
    for thread in load_chat_threads():
        if str(thread.get("id")) == thread_id:
            return thread
    return None


def upsert_chat_thread(payload: dict) -> dict:
    messages = [_clean_chat_message(item) for item in payload.get("messages", []) if isinstance(item, dict)]
    messages = [item for item in messages if item]
    title = str(payload.get("title") or "").strip()
    thread_id = str(payload.get("id") or "").strip() or uuid.uuid4().hex[:12]
    now = _now_iso()
    existing_title = ""
    existing_title_pending = False
    for thread in load_chat_threads():
        if str(thread.get("id")) == thread_id and thread.get("title"):
            existing_title = str(thread.get("title"))
            existing_title_pending = bool(thread.get("titlePending"))
            break
    if title:
        title = _sanitize_thread_title(title)
    title_source = _thread_title_source(messages)
    first_user = next((item["content"] for item in messages if item["role"] == "user"), "")
    placeholder_titles = {
        "新对话",
        "未命名对话",
        "你好",
        "看看",
        first_user[:24],
        _fallback_thread_title(first_user),
    }
    title_pending = False
    if title:
        title_pending = False
    elif existing_title and not existing_title_pending and existing_title not in placeholder_titles:
        title = existing_title
    else:
        title = _fallback_thread_title(title_source or first_user)
        title_pending = True
    with CHAT_THREADS_LOCK:
        threads = load_chat_threads()
        existing = next((item for item in threads if str(item.get("id")) == thread_id), None)
        record = {
            "id": thread_id,
            "title": title[:80],
            "titlePending": title_pending,
            "createdAt": (existing or {}).get("createdAt") or now,
            "updatedAt": now,
            "messages": messages[-80:],
            "agentContextKey": str(payload.get("agentContextKey") or ""),
            "loadedSkillIds": [str(item) for item in payload.get("loadedSkillIds", []) if str(item)],
            "pinned": bool((existing or {}).get("pinned")),
        }
        threads = [item for item in threads if str(item.get("id")) != thread_id]
        threads.insert(0, record)
        save_chat_threads(threads)
    if title_pending:
        _schedule_chat_thread_title(thread_id, title_source or first_user)
    return record


def delete_chat_thread(thread_id: str) -> bool:
    with CHAT_THREADS_LOCK:
        threads = load_chat_threads()
        next_threads = [item for item in threads if str(item.get("id")) != thread_id]
        save_chat_threads(next_threads)
    return len(next_threads) != len(threads)


def set_chat_thread_pinned(thread_id: str, pinned: bool) -> dict | None:
    with CHAT_THREADS_LOCK:
        threads = load_chat_threads()
        updated = None
        for thread in threads:
            if str(thread.get("id")) == thread_id:
                thread["pinned"] = bool(pinned)
                thread["updatedAt"] = _now_iso()
                updated = thread
                break
        save_chat_threads(threads)
    return updated


def reference_path(name: str) -> Path | None:
    raw = str(name or "").strip().lstrip("/")
    clean = Path(raw).name
    if clean in REFERENCE_FILES:
        return ROOT / clean
    if re.fullmatch(r"row_\d+\.json", clean):
        return RESULTS_DIR / clean
    if raw.startswith("agent_knowledge/"):
        target = (ROOT / raw).resolve()
        knowledge_root = (ROOT / "agent_knowledge").resolve()
        if knowledge_root in target.parents and target.exists() and target.is_file():
            return target
    return None


def decode_text_bytes(body: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5", "cp950"):
        try:
            return body.decode(encoding)
        except UnicodeDecodeError:
            continue
    return body.decode("utf-8", errors="replace")


def read_display_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            from docx import Document

            doc = Document(str(path))
            parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as exc:
            return f"Word 文档预览失败：{exc}"
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
        except Exception as exc:
            return f"PDF 预览失败：{exc}"
    raw = decode_text_bytes(path.read_bytes())
    if suffix == ".json":
        try:
            raw = json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
        except Exception:
            pass
    return raw


def settings_rows() -> list[dict]:
    rows = []
    for source_row in crawl.parse_latest_sheet():
        row_no = str(source_row["row"])
        available_entities = list(dict.fromkeys(source_row.get("entities") or []))
        available_fields = list(dict.fromkeys(row_fields(int(row_no))))
        rows.append(
            {
                "row": row_no,
                "block": source_row.get("block", ""),
                "object": source_row.get("object", ""),
                "package": source_row.get("package", ""),
                "need": source_row.get("need", ""),
                "sources": source_row.get("sources", ""),
                "entities": available_entities,
                "fields": available_fields,
                "enabled": True,
                "selectedEntities": available_entities,
                "selectedFields": available_fields,
            }
        )
    return rows


def build_settings_payload() -> dict:
    rows = settings_rows()
    enabled = [row for row in rows if row["enabled"]]
    return {
        "source": "飞书主表",
        "rows": rows,
        "summary": {
            "totalRows": len(rows),
            "enabledRows": len(enabled),
            "selectedEntities": sum(len(row["selectedEntities"]) for row in enabled),
            "selectedFields": sum(len(row["selectedFields"]) for row in enabled),
        },
    }


def json_response(handler: BaseHTTPRequestHandler, payload: dict, status: int = 200) -> None:
    body = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
    try:
        handler.send_response(status)
        handler.send_header("Content-Type", "application/json; charset=utf-8")
        handler.send_header("Content-Length", str(len(body)))
        handler.send_header("Cache-Control", "no-store")
        handler.end_headers()
        handler.wfile.write(body)
    except (BrokenPipeError, ConnectionResetError, OSError):
        # Status polling and stopped browser requests can disconnect while a
        # response is being written.  The request is already over, so avoid
        # turning a harmless client disconnect into a server traceback.
        return


def start_ndjson_response(handler: BaseHTTPRequestHandler) -> None:
    handler.send_response(200)
    handler.send_header("Content-Type", "application/x-ndjson; charset=utf-8")
    handler.send_header("Cache-Control", "no-store, no-transform")
    handler.send_header("X-Accel-Buffering", "no")
    handler.send_header("Connection", "close")
    handler.end_headers()


def write_ndjson_event(handler: BaseHTTPRequestHandler, payload: dict) -> None:
    body = (json.dumps(payload, ensure_ascii=False, separators=(",", ":")) + "\n").encode("utf-8")
    handler.wfile.write(body)
    handler.wfile.flush()


def public_intelligence_error_message(exc: Exception) -> str:
    """Return a stable user-facing message without leaking model output or gates."""
    raw = str(exc or "").strip()
    internal_markers = (
        "AI分析", "AI跨库发现", "新洞察", "模型未返回", "内网模型",
        "Expecting value", "JSON", "Traceback", "SyntaxError", "内容：",
        "输入之外的数字", "必须精炼", "门禁",
    )
    if not raw or any(marker in raw for marker in internal_markers):
        return "本次AI结果未通过数据校验，已保留当前版本，请点击重试。"
    return raw[:120]


def load_curation_status() -> dict:
    if not CURATION_LATEST_PATH.exists():
        return {}
    try:
        payload = json.loads(CURATION_LATEST_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    return payload if isinstance(payload, dict) else {}


def build_curation_rejection_visuals() -> dict:
    status = load_curation_status()
    accepted = int(status.get("accepted") or 0)
    reported_rejected = int(status.get("rejected") or 0)
    quality_rejected = 0
    evidence_gaps = 0
    review = int(status.get("review") or 0)
    reasons: dict[str, int] = {}
    if CURATION_CANDIDATE_FACTS_PATH.exists():
        try:
            for line in CURATION_CANDIDATE_FACTS_PATH.read_text(encoding="utf-8").splitlines():
                if not line.strip():
                    continue
                item = json.loads(line)
                if item.get("decision") != "rejected":
                    continue
                if item.get("status") != "ok":
                    evidence_gaps += 1
                    continue
                quality_rejected += 1
                for reason in item.get("reasons") or []:
                    reason_text = str(reason or "").strip()
                    if reason_text:
                        reasons[reason_text] = reasons.get(reason_text, 0) + 1
        except Exception:
            reasons = {}
    if quality_rejected + evidence_gaps == 0 and reported_rejected:
        quality_rejected = reported_rejected
    quality_total = accepted + quality_rejected + review
    total = accepted + reported_rejected + review
    top_reasons = sorted(
        [{"label": key, "value": value} for key, value in reasons.items()],
        key=lambda item: item["value"],
        reverse=True,
    )[:6]
    return {
        "accepted": accepted,
        "rejected": quality_rejected,
        "qualityRejected": quality_rejected,
        "evidenceGaps": evidence_gaps,
        "reportedRejected": reported_rejected,
        "review": review,
        "total": total,
        "qualityTotal": quality_total,
        "rejectRate": round((quality_rejected / quality_total) * 100) if quality_total else 0,
        "passRate": round((accepted / quality_total) * 100) if quality_total else 0,
        "reasons": top_reasons,
        "runId": status.get("run_id", ""),
        "completedAt": status.get("completed_at", ""),
    }


def build_crawl_result_visuals() -> dict:
    run_log_path = ROOT / "run_log.json"
    if not run_log_path.exists():
        return {
            "success": 0,
            "failed": 0,
            "fallback": 0,
            "total": 0,
            "successRate": 0,
            "completedAt": "",
        }
    try:
        rows = json.loads(run_log_path.read_text(encoding="utf-8"))
    except Exception:
        rows = []
    if not isinstance(rows, list):
        rows = []

    success = 0
    failed = 0
    fallback = 0
    for item in rows:
        if not isinstance(item, dict):
            continue
        status = int(item.get("http_status") or 0)
        used_fallback = str(item.get("evidence_fallback_used") or "").lower() in {
            "1",
            "true",
            "yes",
        }
        if used_fallback:
            fallback += 1
            failed += 1
        elif 200 <= status < 400:
            success += 1
        else:
            failed += 1
    total = success + failed
    return {
        "success": success,
        "failed": failed,
        "fallback": fallback,
        "total": total,
        "successRate": round((success / total) * 100) if total else 0,
        "completedAt": time.strftime(
            "%Y-%m-%d %H:%M:%S",
            time.localtime(run_log_path.stat().st_mtime),
        ),
    }


def load_agent_trace(limit: int = 300) -> list[dict]:
    if not CURATION_AGENT_TRACE_PATH.exists():
        return []
    rows: list[dict] = []
    try:
        lines = CURATION_AGENT_TRACE_PATH.read_text(encoding="utf-8").splitlines()
    except Exception:
        return []
    for line in lines[-limit:]:
        if not line.strip():
            continue
        try:
            item = json.loads(line)
        except Exception:
            continue
        if isinstance(item, dict):
            rows.append(item)
    return rows


def read_request_json(handler: BaseHTTPRequestHandler) -> dict:
    length = int(handler.headers.get("Content-Length") or 0)
    if not length:
        return {}
    raw = handler.rfile.read(length)
    return json.loads(raw.decode("utf-8") or "{}")


def _chat_approval_key(request_id: str, action_id: str) -> tuple[str, str]:
    return (str(request_id or "")[:160], str(action_id or "")[:240])


def register_chat_approval(request_id: str, action_id: str) -> None:
    key = _chat_approval_key(request_id, action_id)
    with CHAT_APPROVAL_LOCK:
        CHAT_APPROVAL_WAITERS[key] = {"event": threading.Event(), "decision": ""}


def resolve_chat_approval(request_id: str, action_id: str, decision: str) -> bool:
    key = _chat_approval_key(request_id, action_id)
    normalized = "allow" if decision == "allow" else "deny"
    with CHAT_APPROVAL_LOCK:
        waiter = CHAT_APPROVAL_WAITERS.get(key)
        if not waiter:
            return False
        waiter["decision"] = normalized
        signal = waiter.get("event")
    if isinstance(signal, threading.Event):
        signal.set()
    return True


def wait_for_chat_approval(request_id: str, action_id: str) -> str:
    key = _chat_approval_key(request_id, action_id)
    with CHAT_APPROVAL_LOCK:
        waiter = CHAT_APPROVAL_WAITERS.get(key)
    if not waiter:
        return "deny"
    signal = waiter.get("event")
    if isinstance(signal, threading.Event):
        signal.wait()
    with CHAT_APPROVAL_LOCK:
        resolved = CHAT_APPROVAL_WAITERS.pop(key, waiter)
    return "allow" if resolved.get("decision") == "allow" else "deny"


def stream_agent_with_approvals(
    message: str,
    *,
    request_id: str,
    approved_action_ids: list[str] | None = None,
    decision_waiter=wait_for_chat_approval,
    agent_factory=stream_agent,
    **agent_kwargs,
):
    """Pause one SSE turn for approval, then resume it with the user's decision."""
    approved = {str(item) for item in (approved_action_ids or []) if str(item).strip()}
    while True:
        events = agent_factory(message, approved_action_ids=sorted(approved), **agent_kwargs)
        restart = False
        try:
            for raw_event in events:
                event = dict(raw_event or {})
                if event.get("type") != "action_confirmation":
                    yield event
                    continue
                action_id = str(event.get("actionId") or "")
                if not action_id:
                    yield event
                    continue
                register_chat_approval(request_id, action_id)
                event["requestId"] = request_id
                yield event
                decision = decision_waiter(request_id, action_id)
                with CHAT_APPROVAL_LOCK:
                    CHAT_APPROVAL_WAITERS.pop(_chat_approval_key(request_id, action_id), None)
                yield {
                    "type": "approval_result",
                    "requestId": request_id,
                    "actionId": action_id,
                    "decision": decision,
                    "label": str(event.get("label") or "执行操作"),
                }
                if decision == "allow":
                    approved.add(action_id)
                    restart = True
                else:
                    yield {"type": "delta", "text": f"已取消执行：{event.get('label') or '该操作'}。"}
                    yield {"type": "done"}
                break
        finally:
            close = getattr(events, "close", None)
            if callable(close):
                close()
        if restart:
            continue
        return


def safe_dataset_slug(value: str) -> str:
    stem = Path(value or "upload").stem or "upload"
    slug = re.sub(r"[^A-Za-z0-9_.\-\u4e00-\u9fff]+", "-", stem).strip("-._")
    return slug[:48] or "upload"


def clean_upload_text(value: object, limit: int = 500) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()[:limit]


def split_upload_tags(value: object, limit: int = 12) -> list[str]:
    tags: list[str] = []
    for part in re.split(r"[,，;；\n]+", str(value or "")):
        tag = clean_upload_text(part, 32)
        if tag and tag not in tags:
            tags.append(tag)
        if len(tags) >= limit:
            break
    return tags


def write_uploaded_knowledge_dataset(payload: dict) -> dict:
    filename = str(payload.get("filename") or "").strip()
    encoded = str(payload.get("contentBase64") or "").strip()
    title = clean_upload_text(payload.get("title"), 80)
    summary = clean_upload_text(payload.get("summary"), 600)
    scope = clean_upload_text(payload.get("scope"), 600)
    source_type = clean_upload_text(payload.get("sourceType") or payload.get("source_type"), 40) or "user_uploaded_file"
    quality_note = clean_upload_text(payload.get("quality"), 600)
    user_tags = split_upload_tags(payload.get("tags"))
    if not filename:
        raise ValueError("缺少文件名")
    if not title:
        raise ValueError("请填写知识库名称")
    if not summary:
        raise ValueError("请填写知识库说明")
    suffix = Path(filename).suffix.lower()
    if suffix not in UPLOAD_ALLOWED_SUFFIXES:
        raise ValueError("暂不支持该文件类型；请上传 txt、md、csv、tsv、json、docx 或 pdf。")
    if not encoded:
        raise ValueError("上传文件内容为空")
    try:
        raw = base64.b64decode(encoded, validate=True)
    except Exception as exc:
        raise ValueError(f"文件内容解码失败：{exc}") from exc
    if not raw:
        raise ValueError("上传文件内容为空")
    if len(raw) > UPLOAD_MAX_BYTES:
        raise ValueError("文件过大，当前单文件上限为 8MB。")

    now = datetime.now().astimezone()
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    slug = safe_dataset_slug(title or filename)
    dataset_id = f"{UPLOAD_DATASET_PREFIX}-{timestamp}-{slug}"
    folder = ROOT / "agent_knowledge" / dataset_id
    folder.mkdir(parents=True, exist_ok=False)

    original_name = f"original{suffix}"
    original_path = folder / original_name
    original_path.write_bytes(raw)

    extracted_text = read_display_text(original_path).strip()
    if not extracted_text:
        extracted_text = decode_text_bytes(raw).strip()
    if not extracted_text:
        raise ValueError("文件已保存但未能提取可检索文本，请换用文本、CSV、JSON、Word 或可复制文字的 PDF。")

    knowledge_path = folder / "uploaded_knowledge.md"
    knowledge_path.write_text(
        "\n".join(
            [
                f"# {title}",
                "",
                f"- 上传时间：{now.isoformat(timespec='seconds')}",
                f"- 原始文件：{original_name}",
                f"- 文件大小：{len(raw)} bytes",
                f"- 知识库说明：{summary}",
                f"- 范围/口径：{scope or '用户未填写'}",
                f"- 来源类型：{source_type}",
                f"- 标签：{', '.join(user_tags) if user_tags else '用户未填写'}",
                f"- 质量备注：{quality_note or '用户未填写'}",
                "",
                "## 可检索正文",
                "",
                extracted_text[:300000],
            ]
        ),
        encoding="utf-8",
    )
    readme_path = folder / "README.md"
    readme_path.write_text(
        "\n".join(
            [
                f"# {title}",
                "",
                "该数据集由前端上传文件生成。只有用户在数据库按钮中选中本数据集时，后端才会把它发送给小竞AI检索。",
                "",
                f"- 数据集 id：`{dataset_id}`",
                f"- 说明：{summary}",
                f"- 范围/口径：{scope or '用户未填写'}",
                f"- 来源类型：{source_type}",
                f"- 标签：{', '.join(user_tags) if user_tags else '用户未填写'}",
                f"- 质量备注：{quality_note or '用户未填写'}",
                f"- 原始文件：`{original_name}`",
                "- 检索入口：`uploaded_knowledge.md`",
            ]
        ),
        encoding="utf-8",
    )
    manifest = {
        "id": dataset_id,
        "title": title,
        "summary": summary,
        "source_type": source_type,
        "scope": scope or "用户手动上传给小竞AI的知识库文件",
        "tags": ["user-upload", "knowledge-base", *user_tags],
        "keywords": [title, Path(filename).stem, filename, summary, *user_tags, "用户上传", "知识库"],
        "entrypoints": ["README.md", "uploaded_knowledge.md"],
        "updated_at": now.isoformat(timespec="seconds"),
        "quality": quality_note or "user_uploaded_unverified; visible to AI only when selected in the database picker",
        "original_file": original_name,
    }
    (folder / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    dataset = next((item for item in list_knowledge_datasets() if item.get("id") == dataset_id), manifest)
    return {"dataset": dataset, "folder": folder.relative_to(ROOT).as_posix()}


def load_report_metadata() -> dict:
    if not REPORT_METADATA_PATH.exists():
        return {}
    try:
        data = json.loads(REPORT_METADATA_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    return data if isinstance(data, dict) else {}


def save_report_metadata(data: dict) -> None:
    REPORT_METADATA_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def is_report_path(path: Path) -> bool:
    if not path.exists() or not path.is_file() or path.suffix.lower() != ".docx":
        return False
    if path.name.startswith("~$") or path.name in EXCLUDED_REPORT_NAMES:
        return False
    try:
        path.relative_to(ROOT)
    except ValueError:
        return False
    return path.parent == ROOT or ROOT / "archives" in path.parents


def quality_sidecar_for_report(path: Path) -> Path:
    return Path(str(path) + ".quality.json")


def report_audio_metadata(report_path: Path) -> dict:
    audio_path = next(
        (path for path in audio_paths_for_report(report_path) if path.exists()),
        None,
    )
    if not audio_path:
        return {"exists": False}
    audio_stat = audio_path.stat()
    return {
        "exists": True,
        "url": f"/audio/{quote(audio_path.name)}?v={audio_stat.st_mtime_ns}",
    }


def file_info(path: Path, url: str = None) -> dict:
    stat = path.stat()
    rel_path = str(path.relative_to(ROOT))
    metadata = load_report_metadata().get(rel_path, {})
    compact_audio = report_audio_metadata(path)
    return {
        "name": path.name,
        "size": stat.st_size,
        "mtime": stat.st_mtime,
        "mtimeText": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_mtime)),
        "url": url or f"/outputs/{quote(path.name)}",
        "path_str": rel_path,
        "note": metadata.get("note", "") if isinstance(metadata, dict) else "",
        "reportType": "carrier-performance" if "业绩摘要" in path.name else "weekly",
        # Full subtitle cues and spoken text are loaded only when the user
        # plays one report.  Embedding every historical transcript made the
        # ten-second status poll grow to megabytes.
        "audio": compact_audio,
    }


def is_report_file_name(name: str) -> bool:
    return name.endswith(".docx") and "/" not in name and "\\" not in name and name not in EXCLUDED_REPORT_NAMES


def current_report_files() -> list[Path]:
    files = [path for path in ROOT.glob("*.docx") if is_report_path(path)]
    return sorted(files, key=lambda p: p.stat().st_mtime, reverse=True)


def report_target_from_rel(path_str: str) -> Path | None:
    if not path_str or path_str.startswith("/") or ".." in Path(path_str).parts:
        return None
    target = ROOT / path_str
    try:
        target.relative_to(ROOT)
    except ValueError:
        return None
    return target if is_report_path(target) else None


def update_report_file(payload: dict) -> dict:
    target = report_target_from_rel(str(payload.get("path") or ""))
    if not target:
        raise ValueError("文件不存在或不允许修改")
    new_name = Path(str(payload.get("name") or "").strip()).name
    if not new_name:
        raise ValueError("文件名不能为空")
    if not new_name.endswith(".docx"):
        new_name += ".docx"
    if not is_report_file_name(new_name):
        raise ValueError("文件名只能是 Word 文档，不能包含路径字符")
    new_note = re.sub(r"\s+", " ", str(payload.get("note") or "")).strip()[:500]
    new_target = target.with_name(new_name)
    if new_target != target and new_target.exists():
        raise ValueError("同名文件已存在")
    old_quality = quality_sidecar_for_report(target)
    new_quality = quality_sidecar_for_report(new_target)
    if new_target != target and new_quality.exists():
        raise ValueError("同名质量审计文件已存在")

    metadata = load_report_metadata()
    old_rel = str(target.relative_to(ROOT))
    if new_target != target:
        target.rename(new_target)
        if old_quality.exists():
            old_quality.rename(new_quality)
            try:
                quality_payload = json.loads(new_quality.read_text(encoding="utf-8"))
                if isinstance(quality_payload, dict):
                    quality_payload["reportFile"] = new_target.name
                    new_quality.write_text(
                        json.dumps(quality_payload, ensure_ascii=False, indent=2),
                        encoding="utf-8",
                    )
            except Exception:
                pass
        rename_audio_for_report(target, new_target)
        existing = metadata.pop(old_rel, {})
    else:
        existing = metadata.get(old_rel, {})
    new_rel = str(new_target.relative_to(ROOT))
    if not isinstance(existing, dict):
        existing = {}
    existing["note"] = new_note
    existing["updatedAt"] = time.strftime("%Y-%m-%d %H:%M:%S")
    metadata[new_rel] = existing
    save_report_metadata(metadata)
    return build_status()


def delete_report_files(paths: list[str]) -> dict:
    metadata = load_report_metadata()
    deleted = 0
    for path_str in paths:
        target = report_target_from_rel(str(path_str))
        if not target:
            continue
        rel_path = str(target.relative_to(ROOT))
        target.unlink()
        quality_sidecar = quality_sidecar_for_report(target)
        if quality_sidecar.exists():
            quality_sidecar.unlink()
        delete_audio_for_report(target)
        metadata.pop(rel_path, None)
        deleted += 1
    save_report_metadata(metadata)
    return {"deleted": deleted, "status": build_status()}


def current_crawl_result_files() -> list[Path]:
    """Return only result files listed by the latest crawl coverage report."""
    coverage_path = ROOT / "coverage_report.tsv"
    if coverage_path.exists():
        try:
            with coverage_path.open(encoding="utf-8", newline="") as fh:
                rows = list(csv.DictReader(fh, delimiter="\t"))
            current: list[Path] = []
            results_root = RESULTS_DIR.resolve()
            for row in rows:
                raw_path = str(row.get("result_file") or "").strip()
                if not raw_path:
                    continue
                candidate = (ROOT / raw_path).resolve()
                if candidate.parent == results_root and candidate.name.startswith("row_") and candidate.exists():
                    current.append(candidate)
            if current:
                return current
        except (OSError, csv.Error):
            pass
    return sorted(RESULTS_DIR.glob("row_*.json"), key=lambda p: int(p.stem.split("_")[1]))


def build_status() -> dict:
    result_files = current_crawl_result_files()
    running_tasks = [
        task
        for task in load_unified_task_index(limit=1000)
        if str(task.get("run_status") or "") == "running"
    ]
    
    outputs = [file_info(path) for path in current_report_files()]
    
    ok_count = 0
    partial_count = 0
    failed_count = 0
    block_counts: dict[str, int] = {}
    source_type_counts: dict[str, int] = {}
    jurisdiction_counts: dict[str, int] = {}
    method_counts: dict[str, int] = {}
    entity_counts: dict[str, int] = {}
    field_total = 0
    missing_total = 0
    raw_total = 0
    
    valid_results_count = 0
    for path in result_files:
        valid_results_count += 1
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        if data.get("status") == "ok":
            ok_count += 1
        elif data.get("status") == "partial":
            partial_count += 1
        else:
            failed_count += 1
        block = str(data.get("need") or data.get("block") or "未分类")
        if "香港" in block:
            block = "香港本地"
        elif any(token in block for token in ["国际", "全球", "欧盟", "国家"]):
            block = "国际监管"
        elif any(token in block for token in ["收入", "ARPU", "EBITDA", "利润", "客户"]):
            block = "经营指标"
        elif any(token in block for token in ["套餐", "资费", "产品", "服务"]):
            block = "产品资费"
        else:
            block = "运营动态"
        block_counts[block] = block_counts.get(block, 0) + 1
        selected_fields = data.get("selected_fields") or []
        missing_fields = data.get("missing_fields") or []
        if isinstance(selected_fields, list):
            field_total += len(selected_fields)
        if isinstance(missing_fields, list):
            missing_total += len(missing_fields)
        for entity in data.get("entities") or []:
            name = str(entity).strip()
            if name:
                entity_counts[name] = entity_counts.get(name, 0) + 1
        for record in data.get("raw_records") or []:
            if not isinstance(record, dict):
                continue
            raw_total += 1
            source_type = str(record.get("source_type") or "unknown")
            source_type_counts[source_type] = source_type_counts.get(source_type, 0) + 1
            jurisdiction = str(record.get("jurisdiction") or "unknown")
            jurisdiction_counts[jurisdiction] = jurisdiction_counts.get(jurisdiction, 0) + 1
            method = str(record.get("method") or "unknown")
            method_counts[method] = method_counts.get(method, 0) + 1
            
    # Calculate latest timestamp from crawler output rather than HTML reports
    latest_crawl_time = max((path.stat().st_mtime for path in result_files if path.exists()), default=None)
    settings = build_settings_payload()
    latest_news_funnel = build_latest_news_funnel()
    today_news_rounds = build_today_news_rounds()
    
    # Sort outputs by mtime descending
    outputs.sort(key=lambda x: x["mtime"], reverse=True)
        
    return {
        "template": {
            "path": str(TEMPLATE_PATH),
            "exists": TEMPLATE_PATH.exists(),
            "mtimeText": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(TEMPLATE_PATH.stat().st_mtime))
            if TEMPLATE_PATH.exists()
            else "",
        },
        "results": {
            "count": valid_results_count,
            "ok": ok_count,
            "partial": partial_count,
        },
        "visuals": {
            "crawl": build_crawl_result_visuals(),
            "quality": {
                "ok": ok_count,
                "partial": partial_count,
                "failed": failed_count,
                "fieldTotal": field_total,
                "missingFields": missing_total,
                "rawSources": raw_total,
            },
            "blocks": sorted(
                [{"label": key, "value": value} for key, value in block_counts.items()],
                key=lambda item: item["value"],
                reverse=True,
            ),
            "sourceTypes": sorted(
                [{"label": key, "value": value} for key, value in source_type_counts.items()],
                key=lambda item: item["value"],
                reverse=True,
            )[:6],
            "jurisdictions": sorted(
                [{"label": key, "value": value} for key, value in jurisdiction_counts.items()],
                key=lambda item: item["value"],
                reverse=True,
            )[:6],
            "methods": sorted(
                [{"label": key, "value": value} for key, value in method_counts.items()],
                key=lambda item: item["value"],
                reverse=True,
            )[:6],
            "rejection": build_curation_rejection_visuals(),
            "newsFunnel": latest_news_funnel,
            "todayNewsRounds": today_news_rounds,
            "entities": sorted(
                [{"label": key, "value": value} for key, value in entity_counts.items()],
                key=lambda item: item["value"],
                reverse=True,
            )[:8],
            "outputs": [
                {
                    "name": item["name"],
                    "mtime": item["mtime"],
                    "mtimeText": item["mtimeText"],
                    "audio": bool(item.get("audio", {}).get("exists")),
                }
                for item in outputs[:8]
            ],
        },
        "outputs": outputs,
        "settings": settings["summary"],
        "ai": load_ai_config(include_key=False),
        "tasks": {
            "runningCount": len(running_tasks),
            "hasRunning": bool(running_tasks),
        },
        "latestOutputText": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(latest_crawl_time)) if latest_crawl_time else "未生成",
    }


def build_scheduler_overview(*, force: bool = False) -> dict[str, object]:
    """Return a read-only, cached view of every effective crawl schedule and its downstream jobs."""
    now_monotonic = time.monotonic()
    with SCHEDULER_OVERVIEW_LOCK:
        cached_at = float(SCHEDULER_OVERVIEW_CACHE.get("cached_at_monotonic") or 0)
        cached_payload = SCHEDULER_OVERVIEW_CACHE.get("payload")
        if not force and isinstance(cached_payload, dict) and now_monotonic - cached_at < SCHEDULER_OVERVIEW_CACHE_SECONDS:
            return dict(cached_payload)

        import scheduler

        now = datetime.now(scheduler.HKT)
        state = scheduler.load_state()
        _due, rows = scheduler.due_rows(now, state)
        active_rows = [item for item in rows if item.get("status") != "disabled"]
        frequency_counts = {"daily": 0, "weekly": 0, "monthly": 0, "other": 0}
        for item in active_rows:
            frequency = str(item.get("frequency") or "")
            key = "daily" if frequency.startswith("每天") else "weekly" if frequency.startswith("每周") else "monthly" if frequency.startswith("每月") else "other"
            frequency_counts[key] += 1

        row_numbers = {int(item.get("row") or 0) for item in active_rows}
        source_groups = [
            {"id": "local", "label": "香港本地竞对", "count": len(row_numbers.intersection(range(2, 19)))},
            {"id": "benchmark", "label": "全球标杆运营商", "count": len(row_numbers.intersection(range(19, 22)))},
            {"id": "hong-kong-news", "label": "香港重点资讯", "count": len(row_numbers.intersection(range(22, 26)))},
            {"id": "international", "label": "国际政策与行业", "count": len(row_numbers.intersection(range(26, 35)))},
        ]

        run_history = load_crawl_run_history(task_kind="")
        latest_main = next((item for item in run_history if str(item.get("trigger") or "") == "定时爬虫"), {})
        latest_news = next((item for item in run_history if str(item.get("task_kind") or "") == "strategic-news"), {})
        latest_intelligence = next((item for item in run_history if str(item.get("task_kind") or "") == "executive-intelligence-refresh"), {})
        next_runs = sorted(
            {str(item.get("next_run_hkt") or "") for item in active_rows if item.get("next_run_hkt")}
        )
        payload: dict[str, object] = {
            "ok": True,
            "checked_at_hkt": now.isoformat(timespec="seconds"),
            "timezone": "Asia/Hong_Kong",
            "configured_rows": len(active_rows),
            "frequency_counts": frequency_counts,
            "source_groups": source_groups,
            "next_runs": next_runs,
            "latest": {
                "main_crawl": latest_main,
                "strategic_news": latest_news,
                "four_database_refresh": latest_intelligence,
            },
            "pipeline": {
                "main_crawl": ["页面抓取", "Agent证据审核", "飞书归档", "页面变化线索"],
                "four_databases": ["local", "international", "cloud", "macro"],
                "four_database_stages": ["数据库刷新", "质量门禁", "17项AI洞察", "主页与公开页发布"],
                "strategic_news": ["线索补缺", "确定性门禁", "AI语义审核", "历史语义去重", "写入与推送"],
            },
        }
        SCHEDULER_OVERVIEW_CACHE.clear()
        SCHEDULER_OVERVIEW_CACHE.update({"cached_at_monotonic": now_monotonic, "payload": payload})
        return dict(payload)


def load_strategic_news_run(slot: str) -> dict:
    normalized_slot = str(slot or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}@\d{2}:\d{2}(?:[-A-Za-z0-9_.]+)?", normalized_slot):
        return {}
    path = STRATEGIC_BRIEFING_RUNS_DIR / f"{normalized_slot.replace(':', '-')}.json"
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return payload if isinstance(payload, dict) else {}


def strategic_news_items_for_crawl_run(run: object) -> list[dict]:
    """Return the published news rows that belong to one archived strategic run."""
    if not isinstance(run, dict) or str(run.get("task_kind") or "") != "strategic-news":
        return []
    summary = run.get("operational_summary")
    slot = str(summary.get("slot") or "") if isinstance(summary, dict) else ""
    payload = load_strategic_news_run(slot)
    review_sheet = payload.get("review_sheet")
    raw_items = review_sheet.get("new_items") if isinstance(review_sheet, dict) else []
    if not isinstance(raw_items, list):
        return []
    items: list[dict] = []
    for raw in raw_items[:100]:
        if not isinstance(raw, dict):
            continue
        raw_url = str(raw.get("url") or "").strip()
        parsed_url = urlparse(raw_url)
        items.append(
            {
                "newsId": str(raw.get("news_id") or ""),
                "title": str(raw.get("title") or ""),
                "summary": str(raw.get("summary") or ""),
                "category": str(raw.get("category") or ""),
                "region": str(raw.get("region") or ""),
                "source": str(raw.get("source") or ""),
                "publishedAt": str(raw.get("published_at") or ""),
                "url": raw_url if parsed_url.scheme in {"http", "https"} else "",
                "inclusionReason": str(raw.get("inclusion_reason") or ""),
                "businessImpact": str(raw.get("business_impact") or ""),
            }
        )
    return items


def _strategic_process_item(raw: object) -> dict:
    if not isinstance(raw, dict):
        return {}
    raw_url = str(raw.get("url") or raw.get("source_url") or "").strip()
    parsed_url = urlparse(raw_url)
    return {
        "newsId": str(raw.get("news_id") or ""),
        "sourceTitle": str(raw.get("source_title") or raw.get("title") or ""),
        "sourceSummary": str(
            raw.get("source_summary")
            or raw.get("snippet")
            or raw.get("summary")
            or ""
        ),
        "source": str(raw.get("source") or raw.get("source_domain") or ""),
        "url": raw_url if parsed_url.scheme in {"http", "https"} else "",
        "publishedAt": str(raw.get("published_at") or raw.get("source_date") or ""),
        "module": str(raw.get("module") or raw.get("category") or ""),
        "matchedKeywords": str(raw.get("matched_keywords") or raw.get("keywords") or ""),
        "status": str(raw.get("status") or ""),
        "shouldInclude": raw.get("should_include") if isinstance(raw.get("should_include"), bool) else None,
        "aiTitle": str(raw.get("ai_title") or ""),
        "aiSummary": str(raw.get("ai_summary") or ""),
        "category": str(raw.get("category") or ""),
        "region": str(raw.get("region") or ""),
        "decisionPath": str(raw.get("decision_path") or ""),
        "signalType": str(raw.get("signal_type") or ""),
        "businessImpact": str(raw.get("business_impact") or ""),
        "exclusionCode": str(raw.get("exclusion_code") or ""),
        "reason": str(raw.get("reason") or raw.get("inclusion_reason") or ""),
        "duplicateOf": str(raw.get("duplicate_of") or ""),
        "errors": [str(item) for item in raw.get("errors") or []],
        "query": str(raw.get("query") or ""),
        "searchOrigin": str(raw.get("search_origin") or ""),
    }


def strategic_news_process_items_for_crawl_run(run: object) -> dict:
    """Expose per-object records for each strategic-news node, not only totals."""
    empty = {"discoveryItems": [], "aiReviewItems": [], "dedupeItems": []}
    if not isinstance(run, dict) or str(run.get("task_kind") or "") != "strategic-news":
        return empty
    summary = run.get("operational_summary")
    slot = str(summary.get("slot") or "") if isinstance(summary, dict) else ""
    payload = load_strategic_news_run(slot)
    discovery = payload.get("news_discovery") if isinstance(payload.get("news_discovery"), dict) else {}
    raw_discovery = discovery.get("items") if isinstance(discovery.get("items"), list) else []
    if not raw_discovery:
        latest_path = STRATEGIC_BRIEFING_DIR / "news_discovery_latest.json"
        try:
            latest = json.loads(latest_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            latest = {}
        generated_at = str(latest.get("generated_at") or "") if isinstance(latest, dict) else ""
        if slot and generated_at[:16] == slot.replace("@", "T")[:16]:
            raw_discovery = latest.get("items") if isinstance(latest.get("items"), list) else []
    review_sheet = payload.get("review_sheet") if isinstance(payload.get("review_sheet"), dict) else {}
    raw_ai = review_sheet.get("ai_review_items") if isinstance(review_sheet.get("ai_review_items"), list) else []
    if not raw_ai and raw_discovery:
        try:
            from strategic_briefing import reconstruct_ai_review_items

            raw_ai = reconstruct_ai_review_items(raw_discovery)
        except Exception:
            raw_ai = []
    raw_dedupe = review_sheet.get("semantic_review_items") if isinstance(review_sheet.get("semantic_review_items"), list) else []
    return {
        "discoveryItems": [item for raw in raw_discovery[:300] if (item := _strategic_process_item(raw))],
        "aiReviewItems": [item for raw in raw_ai[:300] if (item := _strategic_process_item(raw))],
        "dedupeItems": [item for raw in raw_dedupe[:300] if (item := _strategic_process_item(raw))],
    }


def build_today_news_rounds(today_key: str = "") -> list[dict]:
    day = str(today_key or "").strip() or datetime.now().astimezone().strftime("%Y-%m-%d")
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", day):
        return []
    rounds: list[dict] = []
    for hour, label in (("09", "上午"), ("15", "下午")):
        paths = sorted(
            STRATEGIC_BRIEFING_RUNS_DIR.glob(f"{day}@{hour}-*.json"),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        payload: dict = {}
        for path in paths:
            try:
                candidate = json.loads(path.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                continue
            if isinstance(candidate, dict) and isinstance(candidate.get("review_sheet"), dict):
                payload = candidate
                break
        if not payload:
            continue
        review_sheet = payload.get("review_sheet") or {}
        dashboard_summary = payload.get("dashboard_summary") or {}
        news_discovery = payload.get("news_discovery") or {}
        discovered = max(
            0,
            int(
                dashboard_summary.get("discovered")
                or news_discovery.get("result_count")
                or review_sheet.get("input_count")
                or 0
            ),
        )
        confirmed = max(
            0,
            int(
                dashboard_summary.get("confirmed")
                or review_sheet.get("batch_count")
                or max(0, discovered - int(review_sheet.get("filtered_count") or 0))
            ),
        )
        new_count = max(
            0,
            int(dashboard_summary.get("new_count") or review_sheet.get("new_count") or 0),
        )
        history_duplicates = max(
            0,
            min(
                confirmed,
                int(
                    dashboard_summary.get("history_duplicates")
                    or review_sheet.get("semantic_duplicate_count")
                    or 0
                ),
            ),
        )
        deduplicated = max(0, confirmed - history_duplicates)
        status = str(dashboard_summary.get("status") or "").strip()
        if not status:
            status = "已完成" if payload.get("status") == "completed" else "已归档"
        stages = [
            {
                "key": "discovered",
                "label": "检索发现",
                "value": discovered,
                "detail": "固定监控与 Agentic Search 汇总的候选新闻。",
            },
            {
                "key": "confirmed",
                "label": "AI确认",
                "value": confirmed,
                "detail": "AI 结合竞对、政策及战略相关性完成审核。",
            },
            {
                "key": "deduplicated",
                "label": "历史去重",
                "value": deduplicated,
                "detail": f"与历史记录比对，排除 {history_duplicates} 条重复事件。",
            },
            {
                "key": "new",
                "label": "新增入库",
                "value": new_count,
                "detail": "已写入飞书、纳入今日信息资产的新增记录。",
            },
        ]
        rounds.append(
            {
                "key": f"{day}-{hour}",
                "label": label,
                "time": f"{hour}:00",
                "status": status,
                "discovered": discovered,
                "confirmed": confirmed,
                "historyDuplicates": history_duplicates,
                "newCount": new_count,
                "note": str(dashboard_summary.get("note") or "").strip(),
                "categories": _group_latest_news_categories(
                    review_sheet.get("new_category_counts")
                ),
                "impacts": _group_latest_news_impacts(review_sheet.get("new_items")),
                "stages": stages,
            }
        )
    return rounds


def _group_latest_news_categories(raw_counts: object) -> list[dict]:
    if not isinstance(raw_counts, dict):
        return []
    label_map = {
        "基础设施/网络/技术类": "网络与技术",
        "宏观经济&国际形势&地缘政治&其他国际性质关注词汇": "宏观与国际",
        "市场/产品类": "市场与产品",
        "竞争对手": "竞对动态",
    }
    grouped: dict[str, int] = {}
    for raw_label, raw_value in raw_counts.items():
        label = label_map.get(str(raw_label), str(raw_label).strip() or "其他")
        try:
            value = max(0, int(raw_value or 0))
        except (TypeError, ValueError):
            value = 0
        if value:
            grouped[label] = grouped.get(label, 0) + value
    label_priority = {
        "竞对动态": 0,
        "政策监管": 1,
        "网络与技术": 2,
        "市场与产品": 3,
        "宏观与国际": 4,
    }
    items = sorted(
        [{"label": label, "value": value} for label, value in grouped.items()],
        key=lambda item: (-item["value"], label_priority.get(item["label"], 99), item["label"]),
    )
    if len(items) <= 4:
        return items
    return items[:3] + [{"label": "其他", "value": sum(item["value"] for item in items[3:])}]


def _group_latest_news_impacts(new_items: object) -> list[dict]:
    if not isinstance(new_items, list):
        return []
    counts: dict[str, int] = {}
    for item in new_items:
        if not isinstance(item, dict):
            continue
        label = str(item.get("business_impact") or "").strip()
        if label:
            counts[label] = counts.get(label, 0) + 1
    return sorted(
        [{"label": label, "value": value} for label, value in counts.items()],
        key=lambda item: (-item["value"], item["label"]),
    )


def build_latest_news_funnel() -> dict:
    for run in load_crawl_run_index():
        if str(run.get("task_kind") or "") != "strategic-news":
            continue
        if str(run.get("run_status") or "") != "completed":
            continue
        summary = run.get("operational_summary")
        if not isinstance(summary, dict):
            continue
        discovered = max(0, int(summary.get("discovered") or 0))
        ai_confirmed = max(0, int(summary.get("ai_retained") or 0))
        history_duplicates = max(
            0,
            min(ai_confirmed, int(summary.get("history_duplicates") or 0)),
        )
        new_count = max(0, int(summary.get("new_count") or 0))
        deduplicated = max(0, ai_confirmed - history_duplicates)
        slot = str(summary.get("slot") or "")
        run_payload = load_strategic_news_run(slot)
        review_sheet = run_payload.get("review_sheet")
        if not isinstance(review_sheet, dict):
            review_sheet = {}
        categories = _group_latest_news_categories(review_sheet.get("new_category_counts"))
        impacts = _group_latest_news_impacts(review_sheet.get("new_items"))
        source_count = max(0, int(review_sheet.get("new_source_count") or 0))
        slot_match = re.match(
            r"^(\d{4})-(\d{2})-(\d{2})@(\d{2}:\d{2})",
            slot,
        )
        slot_label = (
            f"{int(slot_match.group(2))}月{int(slot_match.group(3))}日 {slot_match.group(4)}"
            if slot_match
            else ""
        )
        return {
            "scope": str(run.get("scope") or ""),
            "label": slot_label,
            "completedAt": str(run.get("completed_at_hkt") or ""),
            "historyDuplicates": history_duplicates,
            "summary": {
                "discovered": discovered,
                "confirmed": ai_confirmed,
                "newCount": new_count,
                "sourceCount": source_count,
            },
            "categories": categories,
            "impacts": impacts,
            "stages": [
                {
                    "key": "discovered",
                    "label": "检索发现",
                    "value": discovered,
                    "removed": 0,
                    "rate": 100,
                    "note": "候选池",
                    "detail": (
                        "固定页面、正式关键词、定时页面线索与 Agentic 补缺搜索合并后，"
                        f"在本轮时间窗内共得到 {discovered} 条候选。"
                    ),
                },
                {
                    "key": "confirmed",
                    "label": "AI确认",
                    "value": ai_confirmed,
                    "removed": max(0, discovered - ai_confirmed),
                    "rate": round(ai_confirmed / discovered * 100) if discovered else 0,
                    "note": (
                        f"保留 {round(ai_confirmed / discovered * 100)}%"
                        if discovered
                        else "等待审核"
                    ),
                    "detail": (
                        "AI逐条判断竞对或战略相关性、具体事件、发布时间及来源证据；"
                        f"确认 {ai_confirmed} 条，未确认 {max(0, discovered - ai_confirmed)} 条。"
                    ),
                },
                {
                    "key": "deduplicated",
                    "label": "历史去重",
                    "value": deduplicated,
                    "removed": history_duplicates,
                    "rate": round(deduplicated / ai_confirmed * 100) if ai_confirmed else 0,
                    "note": f"排除 {history_duplicates} 条",
                    "detail": (
                        "对全部飞书历史记录执行事件级语义去重；"
                        f"识别并排除 {history_duplicates} 条重复事件，剩余 {deduplicated} 条。"
                    ),
                },
                {
                    "key": "new",
                    "label": "本轮新增",
                    "value": new_count,
                    "removed": max(0, deduplicated - new_count),
                    "rate": round(new_count / deduplicated * 100) if deduplicated else 0,
                    "note": "已写入飞书",
                    "detail": (
                        f"最终 {new_count} 条完成飞书写入和逐格回读，"
                        "全部归档成功后才发送群通知。"
                    ),
                },
            ],
        }
    return {
        "scope": "",
        "label": "",
        "completedAt": "",
        "historyDuplicates": 0,
        "summary": {},
        "categories": [],
        "impacts": [],
        "stages": [],
    }


def start_scheduler_with_backend() -> None:
    """Run the Feishu scheduler and strategic briefing monitor with the APP backend."""
    import threading as scheduler_threading
    import traceback as scheduler_traceback

    disable_all = os.environ.get(
        "CMHK_DISABLE_EMBEDDED_SCHEDULER", ""
    ).strip().lower() in {"1", "true", "yes"}
    disable_frequency = os.environ.get(
        "CMHK_DISABLE_FREQUENCY_SCHEDULER", ""
    ).strip().lower() in {"1", "true", "yes"}

    def scheduler_worker() -> None:
        import scheduler

        while True:
            try:
                scheduler.main()
            except Exception:
                scheduler_traceback.print_exc()
            time.sleep(10)

    if not disable_all and not disable_frequency:
        thread = scheduler_threading.Thread(target=scheduler_worker, name="feishu-frequency-scheduler", daemon=True)
        thread.start()
        print("Feishu frequency scheduler started with APP backend", flush=True)

    if (
        not disable_all
        and os.environ.get(
            "CMHK_DISABLE_STRATEGIC_BRIEFING_MONITOR", ""
        ).strip().lower()
        not in {"1", "true", "yes"}
    ):
        def strategic_briefing_worker() -> None:
            import strategic_briefing

            while True:
                try:
                    strategic_briefing.main()
                except Exception:
                    scheduler_traceback.print_exc()
                time.sleep(10)

        briefing_thread = scheduler_threading.Thread(
            target=strategic_briefing_worker,
            name="strategic-briefing-monitor",
            daemon=True,
        )
        briefing_thread.start()

        # News discovery runs inside strategic_briefing._run_scan so discovery,
        # review-sheet synchronization and group reporting form one ordered task.
        # A second 06:00/13:30 worker would race the Feishu write and report stale counts.
        print("Strategic briefing monitor started with APP backend", flush=True)


def run_crawl() -> dict:
    started = time.time()
    crawl_env = os.environ.copy()
    crawl_env.pop("CMHK_ROWS", None)
    crawl_env["CMHK_CRAWL_TRIGGER"] = "手动全量"
    crawl_env["CMHK_CRAWL_SCOPE"] = "全量（第2-34行）"
    proc = subprocess.run(
        [sys.executable, str(ROOT / "crawl.py")],
        cwd=str(ROOT),
        env=crawl_env,
        text=True,
        capture_output=True,
        timeout=1200,
    )
    main_sync = None
    performance_sync = None
    metrics_refresh = None
    agent_trace_sync = None
    if proc.returncode == 0 and (ROOT / "write_payload.json").exists():
        main_sync = subprocess.run(
            [sys.executable, str(ROOT / "daily_crawl_and_write.py"), "--sync-only"],
            cwd=str(ROOT),
            env=crawl_env,
            text=True,
            capture_output=True,
            timeout=600,
        )
        subprocess.run(
            [sys.executable, str(ROOT / "update_sources_from_crawl.py")],
            cwd=str(ROOT),
            text=True,
            capture_output=True,
            timeout=60,
        )
        performance_sync = run_carrier_performance_sync()
        metrics_refresh = run_company_metrics_refresh()
        if main_sync.returncode == 0 and metrics_refresh["ok"]:
            sync_result = json_object_from_output(main_sync.stdout)
            log_sheet_id = str(sync_result.get("log_sheet_id") or "")
            agent_run_id = str(load_curation_status().get("run_id") or "")
            if log_sheet_id and agent_run_id:
                agent_trace_sync = append_agent_trace_to_feishu_log(log_sheet_id, agent_run_id)
    result = {
        "ok": proc.returncode == 0
        and (main_sync is None or main_sync.returncode == 0)
        and (performance_sync is None or performance_sync["ok"])
        and (metrics_refresh is None or metrics_refresh["ok"])
        and (agent_trace_sync is None or agent_trace_sync["ok"]),
        "returnCode": proc.returncode,
        "durationMs": round((time.time() - started) * 1000),
        "stdout": proc.stdout.strip(),
        "stderr": proc.stderr.strip(),
        "mainFeishuSync": None
        if main_sync is None
        else {
            "ok": main_sync.returncode == 0,
            "stdout": main_sync.stdout.strip(),
            "stderr": main_sync.stderr.strip(),
        },
        "carrierPerformanceSync": performance_sync,
        "companyMetricsRefresh": metrics_refresh,
        "agentTraceFeishuSync": agent_trace_sync,
        "status": build_status(),
    }
    result["crawlRunRegistry"] = register_crawl_run(
        crawl_return_code=proc.returncode,
        duration_ms=result["durationMs"],
        sync_result=json_object_from_output(main_sync.stdout) if main_sync and main_sync.returncode == 0 else {},
        metrics_refresh=metrics_refresh or {},
        trace_sync=agent_trace_sync or {},
        trigger="api-crawl",
    )
    return result


def run_company_metrics_refresh() -> dict:
    started = time.time()
    # A full web crawl must act on high-priority evidence gaps, not merely record
    # them. Keep the retry bounded to one round and six rows.
    command = [
        sys.executable,
        str(ROOT / "run_data_curation.py"),
        "--recrawl-gaps",
        "--max-recrawl-rows",
        "6",
        "--max-recrawl-rounds",
        "1",
        "--ai-workers",
        os.environ.get("CMHK_AI_WORKERS", "3"),
        "--search-verify-workers",
        os.environ.get("CMHK_SEARCH_VERIFY_WORKERS", "4"),
    ]
    search_verify_online = os.environ.get("CMHK_SEARCH_VERIFY_ONLINE", "1").lower() not in {
        "0",
        "false",
        "no",
        "off",
    }
    if search_verify_online:
        command.extend(
            [
                "--search-verify-online",
                "--search-verify-online-limit",
                os.environ.get("CMHK_SEARCH_VERIFY_ONLINE_LIMIT", "0"),
            ]
        )
    proc = subprocess.run(
        command,
        cwd=str(ROOT),
        text=True,
        capture_output=True,
        timeout=2400,
    )
    payload = build_company_metrics_payload()
    return {
        "ok": proc.returncode == 0,
        "returnCode": proc.returncode,
        "durationMs": round((time.time() - started) * 1000),
        "stdout": proc.stdout.strip(),
        "stderr": proc.stderr.strip(),
        "summary": payload.get("summary", {}),
    }


def stream_company_metrics_refresh(
    handler: BaseHTTPRequestHandler,
    extra_args: list[str] | None = None,
) -> dict:
    started = time.time()
    command = [
        sys.executable,
        "-u",
        str(ROOT / "run_data_curation.py"),
        "--recrawl-gaps",
        "--max-recrawl-rows",
        "6",
        "--max-recrawl-rounds",
        "1",
        "--ai-workers",
        os.environ.get("CMHK_AI_WORKERS", "3"),
        "--search-verify-workers",
        os.environ.get("CMHK_SEARCH_VERIFY_WORKERS", "4"),
        *(extra_args or []),
    ]
    search_verify_online = os.environ.get("CMHK_SEARCH_VERIFY_ONLINE", "1").lower() not in {
        "0",
        "false",
        "no",
        "off",
    }
    if search_verify_online:
        command.extend(
            [
                "--search-verify-online",
                "--search-verify-online-limit",
                os.environ.get("CMHK_SEARCH_VERIFY_ONLINE_LIMIT", "0"),
            ]
        )
    env = os.environ.copy()
    env["PYTHONUNBUFFERED"] = "1"
    write_sse(
        handler,
        {
            "type": "agent_trace",
            "trace": {
                "ts": datetime.now().astimezone().isoformat(timespec="seconds"),
                "node": "多 Agent 编排器",
                "phase": "tool_call",
                "event_type": "tool_call",
                "message": "启动 LangGraph 多 Agent 数据整理进程。",
                "tool": "run_data_curation.py",
                "input": {
                    "command": command,
                    "workflow": [
                        "证据接收",
                        "来源分类",
                        "事实抽取",
                        "主体校验",
                        "质量审计",
                        "冲突仲裁",
                        "搜索验证",
                        "缺口规划",
                        "Supervisor 工具决策",
                        "定向补爬（最多 6 行、1 轮）",
                        "发布",
                    ],
                },
            },
        },
    )
    proc = subprocess.Popen(
        command,
        cwd=str(ROOT),
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        bufsize=1,
    )
    line_queue: queue.Queue[str | None] = queue.Queue()
    output_lines: list[str] = []

    def read_output() -> None:
        if proc.stdout:
            for raw_line in proc.stdout:
                line_queue.put(raw_line.rstrip("\n"))
        line_queue.put(None)

    threading.Thread(target=read_output, daemon=True).start()
    finished_reading = False
    while not finished_reading:
        try:
            line = line_queue.get(timeout=10)
        except queue.Empty:
            elapsed = round(time.time() - started)
            write_sse(
                handler,
                {
                    "type": "agent_trace",
                    "trace": {
                        "ts": datetime.now().astimezone().isoformat(timespec="seconds"),
                        "node": "多 Agent 编排器",
                        "phase": "observe",
                        "event_type": "agent",
                        "message": f"Agent 仍在处理，已运行 {elapsed} 秒；正在等待当前工具或模型返回。",
                        "output": {"elapsedSeconds": elapsed, "processId": proc.pid},
                    },
                },
            )
            continue
        if line is None:
            finished_reading = True
            continue
        if not line:
            continue
        output_lines.append(line)
        write_sse(handler, sse_payload_from_process_line(line))

    proc.wait()
    payload = build_company_metrics_payload()
    duration_ms = round((time.time() - started) * 1000)
    write_sse(
        handler,
        {
            "type": "agent_trace",
            "trace": {
                "ts": datetime.now().astimezone().isoformat(timespec="seconds"),
                "node": "多 Agent 编排器",
                "phase": "tool_result",
                "event_type": "tool_result",
                "message": "LangGraph 多 Agent 数据整理进程已结束。",
                "tool": "run_data_curation.py",
                "result": {
                    "returnCode": proc.returncode,
                    "durationMs": duration_ms,
                    "summary": payload.get("summary", {}),
                },
            },
        },
    )
    return {
        "ok": proc.returncode == 0,
        "returnCode": proc.returncode,
        "durationMs": duration_ms,
        "stdout": "\n".join(output_lines),
        "stderr": "",
        "summary": payload.get("summary", {}),
    }


def run_carrier_performance_sync() -> dict:
    env = os.environ.copy()
    proc = subprocess.run(
        [sys.executable, str(ROOT / "sync_carrier_performance_feishu.py")],
        cwd=str(ROOT),
        env=env,
        text=True,
        capture_output=True,
        timeout=180,
    )
    return {
        "ok": proc.returncode == 0,
        "returnCode": proc.returncode,
        "stdout": proc.stdout.strip(),
        "stderr": proc.stderr.strip(),
    }


def build_weekly_report_generation_preview(now: datetime | None = None) -> dict[str, object]:
    """Describe the exact input window and approved rows a new weekly run would use."""
    from generate_weekly_report import resolve_weekly_period
    from news_review_sheet import load_weekly_report_candidates

    period = resolve_weekly_period(now)
    effective_range = period.effective_range
    rows, selection_audit = load_weekly_report_candidates(
        effective_range["start"],
        effective_range["end"],
    )
    return {
        "windowStart": effective_range["start"],
        "windowEnd": effective_range["end"],
        "newsCount": len(rows),
        "acceptedRows": int(selection_audit.get("acceptedRows") or 0),
        "excludedRows": int(selection_audit.get("excludedRows") or 0),
        "refreshedAt": period.as_of.isoformat(timespec="seconds"),
        "selectionSource": str(selection_audit.get("selectionSource") or ""),
    }


def run_report_generation() -> dict:
    started = time.time()
    proc = subprocess.run(
        [sys.executable, str(ROOT / "generate_weekly_report.py")],
        cwd=str(ROOT),
        text=True,
        capture_output=True,
        timeout=900,
    )
    status = build_status()
    audio_result = None
    if proc.returncode == 0 and status.get("outputs"):
        try:
            latest_path = latest_output_path(status, "weekly")
            audio_result = synthesize_report_audio(latest_path, force=True)
            status = build_status()
        except Exception as exc:
            audio_result = {"ok": False, "error": str(exc)}
    return {
        "ok": proc.returncode == 0,
        "returnCode": proc.returncode,
        "durationMs": round((time.time() - started) * 1000),
        "stdout": proc.stdout.strip(),
        "stderr": proc.stderr.strip(),
        "audio": audio_result,
        "status": status,
    }


def run_carrier_performance_generation() -> dict:
    started = time.time()
    env = os.environ.copy()
    proc = subprocess.run(
        [sys.executable, str(ROOT / "generate_carrier_performance_report.py")],
        cwd=str(ROOT),
        text=True,
        capture_output=True,
        timeout=900,
    )
    status = build_status()
    audio_result = None
    if proc.returncode == 0 and status.get("outputs"):
        try:
            latest_path = latest_output_path(status, "carrier-performance")
            audio_result = synthesize_report_audio(latest_path, force=True)
            status = build_status()
        except Exception as exc:
            audio_result = {"ok": False, "error": str(exc)}
    return {
        "ok": proc.returncode == 0,
        "returnCode": proc.returncode,
        "durationMs": round((time.time() - started) * 1000),
        "stdout": proc.stdout.strip(),
        "stderr": proc.stderr.strip(),
        "audio": audio_result,
        "status": status,
    }


def latest_output_path(status: dict, report_type: str) -> Path:
    output = next((item for item in status.get("outputs", []) if item.get("reportType") == report_type), None)
    if not output:
        raise FileNotFoundError(f"未找到最新输出：{report_type}")
    return ROOT / output["path_str"]


def push_latest_subscription_content(
    service: SubscriptionService,
    *,
    target_open_id: str = "",
    confirm_bulk: bool = False,
) -> dict:
    """Send each active subscription's latest formal content without a second form."""
    summary = service.list_summary()
    active = [item for item in summary.get("subscribers", []) if item.get("status") == "active"]
    if target_open_id:
        active = [item for item in active if item.get("open_id") == target_open_id]
        if not active:
            raise ValueError("该订阅者未启用，无法人工推送")
    elif not confirm_bulk:
        raise ValueError("一键推送必须在后台完成二次确认")
    if not active:
        raise ValueError("当前没有有效订阅者")

    selected_services = {
        item
        for subscriber in active
        for item in (subscriber.get("services") or [])
        if item in {"weekly", "performance", "news"}
    }
    if not selected_services:
        raise ValueError("接收范围内没有已启用的订阅内容")

    status = build_status()
    content: dict[str, dict[str, str]] = {}
    for service_key, report_type in (("weekly", "weekly"), ("performance", "carrier-performance")):
        if service_key not in selected_services:
            continue
        try:
            content[service_key] = {
                "mode": "pdf_audio",
                "path": str(latest_output_path(status, report_type).relative_to(ROOT)),
            }
        except (FileNotFoundError, ValueError):
            continue
    latest_news: list[dict] = []
    if "news" in selected_services:
        from strategic_briefing import latest_reviewed_news

        # Category filtering must happen before the per-recipient limit, so use
        # the complete verified pool rather than truncating it globally first.
        latest_news = latest_reviewed_news()
    if not content and not latest_news:
        raise ValueError("当前没有可供人工推送的最新正式内容")

    results = []
    for service_key in ("weekly", "performance"):
        if service_key not in content:
            continue
        item = content[service_key]
        results.append(service.push(
            service=service_key,
            mode=item["mode"],
            path=item.get("path", ""),
            title=item.get("title", ""),
            body=item.get("body", ""),
            target_open_id=target_open_id,
            confirm_bulk=confirm_bulk,
        ))
    if latest_news:
        for subscriber in active:
            if "news" not in (subscriber.get("services") or []):
                continue
            item_limit = int(subscriber.get("news_item_limit") or 10)
            news_categories = subscriber.get("news_categories")
            news_items = filter_news_by_categories(
                latest_news,
                news_categories,
                limit=item_limit,
            )
            results.append(service.push(
                service="news",
                mode="text",
                title=f"CMHK战略新闻｜最新{len(news_items)}条｜{news_category_summary(news_categories)}",
                body=encode_strategic_news_digest(news_items),
                target_open_id=str(subscriber.get("open_id") or ""),
            ))
    return {
        "batch_id": f"manual-latest-{uuid.uuid4().hex[:12]}",
        "target_open_id": target_open_id,
        "service_count": len(results),
        "recipient_count": sum(int(item.get("recipient_count") or 0) for item in results),
        "verified_count": sum(int(item.get("verified_count") or 0) for item in results),
        "failed_count": sum(int(item.get("failed_count") or 0) for item in results),
        "results": results,
    }


def write_sse(handler: BaseHTTPRequestHandler, payload: dict) -> bool:
    body = json.dumps(payload, ensure_ascii=False)
    log_paths = [
        getattr(handler, "_crawl_stream_log_path", None),
        getattr(handler, "_crawl_stream_mirror_path", None),
    ]
    for log_path in dict.fromkeys(path for path in log_paths if path):
        try:
            with Path(log_path).open("a", encoding="utf-8") as fh:
                fh.write(body + "\n")
        except OSError:
            pass
    try:
        handler.wfile.write(f"data: {body}\n\n".encode("utf-8"))
        handler.wfile.flush()
        return True
    except (BrokenPipeError, ConnectionResetError, OSError):
        # The crawl and its post-processing must outlive the browser's SSE
        # connection. A refresh, navigation or laptop sleep must not skip
        # Feishu sync, curation or run registration.
        return False


def json_object_from_output(output: str) -> dict:
    match = re.search(r"\{.*\}\s*$", output, re.S)
    if not match:
        return {}
    try:
        value = json.loads(match.group(0))
    except json.JSONDecodeError:
        return {}
    return value if isinstance(value, dict) else {}


def append_agent_trace_to_feishu_log(sheet_id: str, run_id: str) -> dict:
    env = os.environ.copy()
    proc = subprocess.run(
        [
            sys.executable,
            str(ROOT / "daily_crawl_and_write.py"),
            "--append-agent-trace",
            sheet_id,
            run_id,
        ],
        cwd=str(ROOT),
        env=env,
        text=True,
        capture_output=True,
        timeout=300,
    )
    return {
        "ok": proc.returncode == 0,
        "returnCode": proc.returncode,
        "stdout": proc.stdout.strip(),
        "stderr": proc.stderr.strip(),
        "result": json_object_from_output(proc.stdout),
    }


def sse_payload_from_process_line(text: str) -> dict:
    if text.startswith("AGENT_TRACE="):
        try:
            return {"type": "agent_trace", "trace": json.loads(text.split("=", 1)[1])}
        except Exception:
            return {"type": "log", "text": text}
    return {"type": "log", "text": text}


def stream_report_generation(
    handler: BaseHTTPRequestHandler,
    script_name: str,
    report_type: str,
    script_args: list[str] | None = None,
) -> None:
    handler.send_response(200)
    handler.send_header("Content-Type", "text/event-stream; charset=utf-8")
    handler.send_header("Cache-Control", "no-cache")
    handler.end_headers()

    started = time.time()
    proc = subprocess.Popen(
        [sys.executable, "-u", str(ROOT / script_name), *(script_args or [])],
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        bufsize=1,
    )
    handler._task_worker_pid = proc.pid
    handler._task_monitor_phase = "报告生成"
    handler._task_monitor_detail = "报告生成进程正在执行。"
    created_path = None
    if proc.stdout:
        for line in proc.stdout:
            text = line.strip()
            write_sse(handler, sse_payload_from_process_line(text))
            if text.startswith("->"):
                candidate = Path(text[2:].strip())
                if candidate.exists() and candidate.name.endswith(".docx") and "template" not in candidate.name:
                    created_path = candidate
    proc.wait()
    handler._task_worker_pid = 0

    status = build_status()
    audio_result = None
    if proc.returncode == 0 and status.get("outputs"):
        try:
            latest_path = created_path if created_path and created_path.exists() else latest_output_path(status, report_type)
            write_sse(handler, {"type": "log", "text": "报告生成完成。开始生成语音摘要..."})
            code = "import sys, json\nfrom pathlib import Path\nfrom tts_service import synthesize_report_audio\ntry:\n    res = synthesize_report_audio(Path(sys.argv[1]), force=sys.argv[2] == 'True')\n    print(json.dumps({'ok': True, 'result': res}))\nexcept Exception as e:\n    print(json.dumps({'ok': False, 'error': str(e)}))"
            handler._task_monitor_phase = "生成语音摘要"
            handler._task_monitor_detail = "报告文档已完成，正在调用公司内部语音模型。"
            proc_audio = subprocess.Popen(
                [sys.executable, "-c", code, str(latest_path), "True"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            handler._task_worker_pid = proc_audio.pid
            audio_stdout, audio_stderr = proc_audio.communicate()
            handler._task_worker_pid = 0
            try:
                out = json.loads(audio_stdout)
                if not out.get("ok"):
                    raise Exception(out.get("error"))
                audio_result = out.get("result")
                if not audio_result.get("ok"):
                    raise Exception(audio_result.get("error"))
            except Exception as e:
                raise Exception(f"Audio generation failed: {audio_stderr} | {e}")
            status = build_status()
            write_sse(handler, {"type": "log", "text": "✅ 语音摘要生成完成。"})
        except Exception as exc:
            audio_result = {"ok": False, "error": str(exc)}
            write_sse(handler, {"type": "log", "text": f"❌ 语音摘要生成失败: {exc}"})

    audio_failed = isinstance(audio_result, dict) and audio_result.get("ok") is False
    task_ok = proc.returncode == 0
    warning_detail = str((audio_result or {}).get("error") or "") if audio_failed else ""
    report_label = "周报" if report_type == "weekly" else "业绩摘要"
    if task_ok and audio_failed:
        message = f"{report_label}已生成；语音摘要未完成：{warning_detail}"
    elif task_ok:
        message = "报告及语音摘要均已生成。"
    else:
        message = f"{report_label}生成进程未成功完成。"
    write_sse(
        handler,
        {
            "type": "done",
            "ok": task_ok,
            "completedWithWarnings": task_ok and audio_failed,
            "reportGenerated": task_ok,
            "error": "" if task_ok else message,
            "warning": warning_detail,
            "message": message,
            "durationMs": round((time.time() - started) * 1000),
            "audio": audio_result,
            "status": status,
        },
    )


def report_overview() -> str:
    md_path = ROOT / "weekly_report.md"
    if not md_path.exists():
        return "当前还没有生成周报。你可以先点击“生成周报”，系统会按 Word 模板输出正式 Word 周报。"
    text = md_path.read_text(encoding="utf-8", errors="ignore")
    lines = [line.strip("#- 　\t ") for line in text.splitlines() if line.strip()]
    useful = [line for line in lines if line and not line.startswith("来源")][:8]
    status = build_status()
    intro = (
        "这里的周报是“战略内参周报”：把公开信息监测数据按模板整理成正式汇报文件，"
        "主要用于快速查看政策、行业、社会和国际资讯中的重点变化。"
    )
    if not useful:
        return f"{intro} 当前已有输出文件，最近生成时间是 {status['latestOutputText']}。"
    return f"{intro} 当前最近生成时间是 {status['latestOutputText']}。报告开头内容包括：" + "；".join(useful[:5]) + "。"


def output_overview() -> str:
    status = build_status()
    outputs = status.get("outputs", [])
    if not outputs:
        return "当前还没有输出文件。点击“生成周报”后会生成正式 Word 周报。"
    names = "、".join(item["name"] for item in outputs)
    return f"当前可用输出文件有：{names}。这里仅展示正式 Word 周报，用于下载和提交。"


def check_local_action(message: str) -> dict | None:
    return None

    status = build_status()
    status_intent = any(
        key in text
        for key in [
            "系统状态",
            "运行状态",
            "当前状态",
            "检查系统",
            "检查后端",
            "结果文件状态",
            "输出文件状态",
            "现在有多少文件",
            "现在有哪些文件",
        ]
    ) or text in {"状态", "检查", "现在", "文件"}
    if status_intent:
        return {
            "content": (
                f"当前已有 {status['results']['count']} 个结果文件，"
                f"ok {status['results']['ok']} 个，partial {status['results']['partial']} 个。"
                f"模板文件{'存在' if status['template']['exists'] else '不存在'}，"
                f"最近输出时间是 {status['latestOutputText']}。"
            ),
        }

    if "模板" in text or "格式" in text:
        return {
            "content": (
                "当前生成流程会优先读取本地上传的模板，若无则使用库里的默认模板 weekly_report_template.docx，"
                "保留封面、目录位置、页眉页脚和图片资源，只替换目录与正文段落文字。"
            ),
        }

    if "openai" in lowered or "api" in lowered or "ai" in lowered:
        return {
            "content": (
                "这个助手已接入 OpenAI Responses API 的调用代码，并会先对本地周报、爬取结果和审计文件做 RAG 检索。"
                "当前运行环境需要设置 OPENAI_API_KEY 后才能真正调用模型。"
            ),
        }

    return None



def _curation_quality_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def load_curation_quality_records(run_id: str) -> dict:
    run_id = str(run_id or "").strip()
    if not 8 <= len(run_id) <= 80 or any(
        not (character.isalnum() or character in "_-") for character in run_id
    ):
        return {"ok": False, "error": "无效的审核运行编号。", "runId": run_id}

    snapshot_path = ROOT / "curation_data" / "runs" / f"{run_id}_candidate_facts.jsonl"
    source_path = snapshot_path
    source_kind = "run-snapshot"

    if not source_path.exists():
        latest_run_id = ""
        try:
            latest_payload = json.loads(CURATION_LATEST_PATH.read_text(encoding="utf-8"))
            latest_run_id = str(latest_payload.get("run_id") or "")
        except (OSError, ValueError, TypeError):
            latest_run_id = ""
        if latest_run_id == run_id and CURATION_CANDIDATE_FACTS_PATH.exists():
            source_path = CURATION_CANDIDATE_FACTS_PATH
            source_kind = "current-run"
        else:
            return {
                "ok": False,
                "pending": True,
                "error": "本轮逐条质量明细尚未生成；质量审计完成后会自动出现。",
                "runId": run_id,
                "records": [],
            }

    records: list[dict] = []
    try:
        with source_path.open("r", encoding="utf-8") as handle:
            for raw_line in handle:
                raw_line = raw_line.strip()
                if not raw_line:
                    continue
                item = json.loads(raw_line)
                if not isinstance(item, dict):
                    continue

                raw_sources = item.get("sources")
                sources: list[dict] = []
                if isinstance(raw_sources, list):
                    for source in raw_sources:
                        if isinstance(source, dict):
                            url = _curation_quality_text(
                                source.get("url")
                                or source.get("source_url")
                                or source.get("href")
                                or source.get("link")
                            )
                            title = _curation_quality_text(
                                source.get("title")
                                or source.get("name")
                                or source.get("source")
                            )
                            source_type = _curation_quality_text(
                                source.get("type")
                                or source.get("source_type")
                                or source.get("tier")
                            )
                            sources.append({"url": url, "title": title, "type": source_type})
                        else:
                            sources.append({"url": _curation_quality_text(source), "title": "", "type": ""})

                verification = item.get("search_verification")
                if not isinstance(verification, dict):
                    verification = {}
                online_search = verification.get("online_search")
                if not isinstance(online_search, dict):
                    online_search = {}

                raw_votes = verification.get("votes")
                raw_conflicts = verification.get("conflicts")
                verification_summary = {
                    "status": _curation_quality_text(verification.get("status")),
                    "decision": _curation_quality_text(verification.get("decision")),
                    "vote_count": int(verification.get("vote_count") or (len(raw_votes) if isinstance(raw_votes, list) else 0)),
                    "majority_count": int(verification.get("majority_count") or 0),
                    "majority_source_types": verification.get("majority_source_types") or [],
                    "conflict_count": int(verification.get("conflict_count") or (len(raw_conflicts) if isinstance(raw_conflicts, list) else 0)),
                    "online_search": {
                        "enabled": bool(online_search.get("enabled")),
                        "provider": _curation_quality_text(online_search.get("provider")),
                        "result_count": int(online_search.get("result_count") or 0),
                        "duration_ms": int(online_search.get("duration_ms") or 0),
                        "query": _curation_quality_text(online_search.get("query")),
                    },
                }

                reasons = item.get("reasons")
                if not isinstance(reasons, list):
                    reasons = [reasons] if reasons else []

                records.append(
                    {
                        "id": _curation_quality_text(item.get("id")),
                        "company": _curation_quality_text(item.get("company")),
                        "metric": _curation_quality_text(item.get("metric")),
                        "value": _curation_quality_text(item.get("value")),
                        "basis": _curation_quality_text(item.get("basis")),
                        "note": _curation_quality_text(item.get("note")),
                        "status": _curation_quality_text(item.get("status")),
                        "entity_supported": bool(item.get("entity_supported")),
                        "metric_supported": bool(item.get("metric_supported")),
                        "value_supported": bool(item.get("value_supported")),
                        "confidence": item.get("confidence"),
                        "source_score": item.get("source_score"),
                        "source_tier": _curation_quality_text(item.get("source_tier")),
                        "row_ref": item.get("row_ref"),
                        "sources": sources,
                        "quality_score": item.get("quality_score"),
                        "decision": _curation_quality_text(item.get("decision")),
                        "reasons": [_curation_quality_text(reason) for reason in reasons if reason is not None],
                        "search_verification": verification_summary,
                    }
                )
    except (OSError, ValueError, TypeError) as exc:
        return {"ok": False, "error": f"逐条质量明细读取失败：{exc}", "runId": run_id}

    decisions: dict[str, int] = {}
    source_tiers: dict[str, int] = {}
    quality_values: list[float] = []
    for record in records:
        decision = record.get("decision") or "unknown"
        decisions[decision] = decisions.get(decision, 0) + 1
        source_tier = record.get("source_tier") or "unknown"
        source_tiers[source_tier] = source_tiers.get(source_tier, 0) + 1
        try:
            quality_values.append(float(record.get("quality_score")))
        except (TypeError, ValueError):
            pass

    return {
        "ok": True,
        "runId": run_id,
        "source": source_kind,
        "records": records,
        "summary": {
            "total": len(records),
            "decisions": decisions,
            "sourceTiers": source_tiers,
            "averageQuality": round(sum(quality_values) / len(quality_values), 4) if quality_values else None,
        },
    }



# Unified task archive v144
TASK_RUNS_DIR = ROOT / "task_runs"
TASK_RUNS_LOG_DIR = TASK_RUNS_DIR / "logs"
TASK_RUNS_INDEX_PATH = TASK_RUNS_DIR / "index.json"
PROJECT_MONITOR_STATE_PATH = ROOT / "var" / "project_monitor" / "state.json"
PROJECT_MONITOR_ACTIONS_PATH = ROOT / "var" / "project_monitor" / "card_actions.json"
PROJECT_MONITOR_WEB_ACTIONS_PATH = ROOT / "var" / "project_monitor" / "web_actions.jsonl"
TASK_RUNS_LOCK = threading.Lock()
GENERAL_TASK_MAX_AUTO_RETRIES = max(1, int(os.environ.get("CMHK_TASK_AUTO_RETRY_MAX", "3")))
GENERAL_TASK_RETRY_DELAY_SECONDS = max(
    1.0, float(os.environ.get("CMHK_TASK_AUTO_RETRY_DELAY_SECONDS", "5"))
)


def _task_atomic_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(path.name + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    os.replace(temporary, path)


def _task_read_local_index() -> list[dict]:
    if not TASK_RUNS_INDEX_PATH.exists():
        return []
    try:
        payload = json.loads(TASK_RUNS_INDEX_PATH.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        return []
    if isinstance(payload, dict):
        payload = payload.get("tasks")
    return [item for item in payload if isinstance(item, dict)] if isinstance(payload, list) else []


def _task_log_stats(record: dict) -> tuple[int, int]:
    relative_path = str(record.get("log_path") or "")
    path = ROOT / relative_path if relative_path else None
    if not path or not path.exists():
        return 0, 0
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        return len(content.splitlines()), path.stat().st_size
    except OSError:
        return 0, 0


def _task_public_record(record: dict) -> dict:
    public = dict(record)
    lines, size = _task_log_stats(record)
    public["lines"] = lines
    public["bytes"] = size
    public["run_status"] = str(public.get("run_status") or "completed")
    if public["run_status"] == "running" and int(public.get("backend_pid") or 0) != os.getpid():
        public["run_status"] = "failed"
        public["status_detail"] = "后端重启后任务已中断"
    return public


def start_general_task_run(
    kind: str,
    title: str,
    scope: str,
    script_name: str = "",
    *,
    recovery_of: str = "",
    retry_count: int = 0,
    target_path: str = "",
) -> dict:
    now = datetime.now().astimezone()
    raw_id = now.strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:8]
    task_id = "task:" + raw_id
    log_path = TASK_RUNS_LOG_DIR / (raw_id + ".log")
    kind_labels = {
        "weekly-report": "周报生成",
        "carrier-performance": "业绩摘要",
        "audio-generation": "音频生成",
    }
    record = {
        "task_id": task_id,
        "task_run_id": raw_id,
        "kind": kind,
        "kind_label": kind_labels.get(kind, "后台任务"),
        "title": title,
        "scope": scope,
        "script": script_name,
        "run_status": "running",
        "started_at_hkt": now.isoformat(timespec="seconds"),
        "completed_at_hkt": "",
        "duration_ms": 0,
        "backend_pid": os.getpid(),
        "worker_pid": 0,
        "phase": "任务启动",
        "progress_detail": "后台已接收任务，正在准备执行。",
        "heartbeat_at_hkt": now.isoformat(timespec="seconds"),
        "log_path": str(log_path.relative_to(ROOT)),
        "recovery_of": str(recovery_of or ""),
        "retry_count": max(0, int(retry_count or 0)),
        "auto_recovered": bool(recovery_of),
        "target_path": str(target_path or ""),
    }
    with TASK_RUNS_LOCK:
        TASK_RUNS_DIR.mkdir(parents=True, exist_ok=True)
        TASK_RUNS_LOG_DIR.mkdir(parents=True, exist_ok=True)
        tasks = _task_read_local_index()
        tasks.insert(0, record)
        _task_atomic_json(TASK_RUNS_INDEX_PATH, {"tasks": tasks[:500]})
        _task_atomic_json(TASK_RUNS_DIR / (raw_id + ".json"), record)
        log_path.write_text(
            "[" + now.isoformat(timespec="seconds") + "] 任务启动：" + title + "\n",
            encoding="utf-8",
        )
    return record


def append_general_task_log(task_id: str, text: object) -> None:
    raw_id = str(task_id or "").removeprefix("task:")
    if not raw_id or any(not (char.isalnum() or char in "_-") for char in raw_id):
        return
    log_path = TASK_RUNS_LOG_DIR / (raw_id + ".log")
    value = str(text or "")
    if not value:
        return
    with TASK_RUNS_LOCK:
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with log_path.open("a", encoding="utf-8") as handle:
            handle.write(value)
            if not value.endswith("\n"):
                handle.write("\n")


def finish_general_task_run(task_id: str, ok: bool, detail: str = "") -> dict | None:
    now = datetime.now().astimezone()
    raw_id = str(task_id or "").removeprefix("task:")
    updated = None
    with TASK_RUNS_LOCK:
        tasks = _task_read_local_index()
        for task in tasks:
            if str(task.get("task_id") or "") != task_id:
                continue
            started_text = str(task.get("started_at_hkt") or "")
            try:
                started = datetime.fromisoformat(started_text)
                duration_ms = max(0, int((now - started).total_seconds() * 1000))
            except (TypeError, ValueError):
                duration_ms = 0
            task["run_status"] = "completed" if ok else "failed"
            task["completed_at_hkt"] = now.isoformat(timespec="seconds")
            task["duration_ms"] = duration_ms
            task["status_detail"] = detail
            task["worker_pid"] = 0
            task["phase"] = "已完成" if ok else "失败"
            task["progress_detail"] = detail or ("任务全部步骤已完成。" if ok else "任务执行失败，请查看日志。")
            task["heartbeat_at_hkt"] = now.isoformat(timespec="seconds")
            updated = dict(task)
            break
        if updated:
            _task_atomic_json(TASK_RUNS_INDEX_PATH, {"tasks": tasks[:500]})
            _task_atomic_json(TASK_RUNS_DIR / (raw_id + ".json"), updated)
    return updated


def heartbeat_general_task_run(
    task_id: str,
    phase: str,
    detail: str,
    *,
    worker_pid: int = 0,
    append_log: bool = True,
) -> dict | None:
    """Persist report-task liveness even when the browser connection disappears."""
    now = datetime.now().astimezone()
    raw_id = str(task_id or "").removeprefix("task:")
    updated = None
    with TASK_RUNS_LOCK:
        tasks = _task_read_local_index()
        for task in tasks:
            if str(task.get("task_id") or "") != task_id or task.get("run_status") != "running":
                continue
            task.update(
                {
                    "backend_pid": os.getpid(),
                    "worker_pid": int(worker_pid or 0),
                    "phase": str(phase or task.get("phase") or "执行中"),
                    "progress_detail": str(detail or task.get("progress_detail") or "任务仍在执行。"),
                    "heartbeat_at_hkt": now.isoformat(timespec="seconds"),
                }
            )
            updated = dict(task)
            break
        if updated:
            _task_atomic_json(TASK_RUNS_INDEX_PATH, {"tasks": tasks[:500]})
            _task_atomic_json(TASK_RUNS_DIR / (raw_id + ".json"), updated)
            if append_log:
                log_path = TASK_RUNS_LOG_DIR / (raw_id + ".log")
                with log_path.open("a", encoding="utf-8") as handle:
                    handle.write(
                        f"[监控心跳 {now.strftime('%H:%M:%S')}] 阶段：{updated['phase']}；"
                        f"状态：{updated['progress_detail']}\n"
                    )
    return updated


def reconcile_interrupted_general_tasks() -> list[dict]:
    """Persistently close report tasks left running by an earlier backend process."""
    now = datetime.now().astimezone()
    reconciled: list[dict] = []
    with TASK_RUNS_LOCK:
        tasks = _task_read_local_index()
        for task in tasks:
            if task.get("run_status") != "running":
                continue
            raw_id = str(task.get("task_run_id") or str(task.get("task_id") or "").removeprefix("task:"))
            log_path = ROOT / str(task.get("log_path") or "")
            completed = now
            if log_path.exists():
                try:
                    completed = datetime.fromtimestamp(log_path.stat().st_mtime, now.tzinfo)
                except OSError:
                    completed = now
            try:
                started = datetime.fromisoformat(str(task.get("started_at_hkt") or ""))
                duration_ms = max(0, int((completed - started).total_seconds() * 1000))
            except (TypeError, ValueError):
                duration_ms = 0
            detail = "后台服务已重新启动，原任务执行进程已不存在；已按最后心跳明确收尾。"
            task.update(
                {
                    "run_status": "failed",
                    "interrupted": True,
                    "status_detail": detail,
                    "completed_at_hkt": completed.isoformat(timespec="seconds"),
                    "duration_ms": duration_ms,
                    "worker_pid": 0,
                    "phase": "已中断",
                    "progress_detail": detail,
                    "heartbeat_at_hkt": completed.isoformat(timespec="seconds"),
                }
            )
            if log_path:
                log_path.parent.mkdir(parents=True, exist_ok=True)
                with log_path.open("a", encoding="utf-8") as handle:
                    handle.write("[任务中断] " + detail + "\n")
            _task_atomic_json(TASK_RUNS_DIR / (raw_id + ".json"), task)
            reconciled.append(dict(task))
        if reconciled:
            _task_atomic_json(TASK_RUNS_INDEX_PATH, {"tasks": tasks[:500]})
    return reconciled


def pending_interrupted_general_task_retries() -> list[dict]:
    """Include interruptions recorded by an older service before auto-retry existed."""
    return [
        dict(task)
        for task in _task_read_local_index()
        if task.get("run_status") == "failed"
        and bool(task.get("interrupted"))
        and not task.get("recovery_disposition")
    ]


def _mark_general_task_recovery_disposition(task_id: str, disposition: str) -> None:
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    raw_id = str(task_id or "").removeprefix("task:")
    with TASK_RUNS_LOCK:
        tasks = _task_read_local_index()
        updated = None
        for task in tasks:
            if str(task.get("task_id") or "") != task_id:
                continue
            task["recovery_disposition"] = str(disposition or "")
            task["recovery_scheduled_at_hkt"] = now
            updated = dict(task)
            break
        if updated:
            _task_atomic_json(TASK_RUNS_INDEX_PATH, {"tasks": tasks[:500]})
            _task_atomic_json(TASK_RUNS_DIR / (raw_id + ".json"), updated)


def _latest_report_for_recovered_task(kind: str) -> Path:
    report_kind = "weekly" if kind == "weekly-report" else "carrier-performance"
    target = latest_output_path(build_status(), report_kind)
    if not target or not target.exists():
        raise RuntimeError("报告进程结束后未找到可用Word文件")
    return target


def _run_recovered_general_task(original: dict) -> None:
    """Retry an interrupted report/audio task without requiring a browser connection."""
    kind = str(original.get("kind") or "")
    retry_count = int(original.get("retry_count") or 0) + 1
    root_id = str(original.get("recovery_of") or original.get("task_id") or "")
    task = start_general_task_run(
        kind,
        str(original.get("title") or "自动恢复任务"),
        str(original.get("scope") or ""),
        str(original.get("script") or ""),
        recovery_of=root_id,
        retry_count=retry_count,
        target_path=str(original.get("target_path") or ""),
    )
    task_id = str(task["task_id"])
    append_general_task_log(
        task_id,
        f"[自动恢复] 服务已恢复，正在重试中断任务（第{retry_count}/{GENERAL_TASK_MAX_AUTO_RETRIES}次）。",
    )
    try:
        if kind == "audio-generation":
            target = Path(str(original.get("target_path") or ""))
            if not target.exists() or target.parent.resolve() != ROOT.resolve():
                raise RuntimeError("原音频任务的报告文件已不存在或不在允许目录")
            heartbeat_general_task_run(
                task_id,
                "生成语音摘要",
                "服务恢复后正在重新生成音频。",
                append_log=True,
            )
            result = synthesize_report_audio(target, force=True)
            if not result.get("ok"):
                raise RuntimeError(str(result.get("error") or "音频生成失败"))
            detail = f"自动恢复成功：{(result.get('audio') or {}).get('name') or target.name}"
        elif kind in {"weekly-report", "carrier-performance"}:
            script_name = str(original.get("script") or "")
            allowed_scripts = {
                "weekly-report": "generate_weekly_report.py",
                "carrier-performance": "generate_carrier_performance_report.py",
            }
            if script_name != allowed_scripts[kind]:
                raise RuntimeError("任务恢复脚本不在允许列表")
            proc = subprocess.Popen(
                [sys.executable, "-u", str(ROOT / script_name)],
                cwd=str(ROOT),
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
            )
            heartbeat_general_task_run(
                task_id,
                "报告生成",
                "服务恢复后已重新启动报告进程。",
                worker_pid=proc.pid,
                append_log=True,
            )
            if proc.stdout:
                for line in proc.stdout:
                    value = line.rstrip()
                    if value:
                        append_general_task_log(task_id, value)
                        phase, progress = _task_phase_from_payload({"type": "log", "text": value}, "报告生成")
                        heartbeat_general_task_run(
                            task_id,
                            phase,
                            progress,
                            worker_pid=proc.pid,
                            append_log=False,
                        )
            proc.wait()
            if proc.returncode:
                raise RuntimeError(f"报告生成进程返回{proc.returncode}")
            target = _latest_report_for_recovered_task(kind)
            heartbeat_general_task_run(
                task_id,
                "生成语音摘要",
                "Word已恢复生成，正在重新生成音频。",
                append_log=True,
            )
            audio = synthesize_report_audio(target, force=True)
            if not audio.get("ok"):
                raise RuntimeError(str(audio.get("error") or "音频生成失败"))
            detail = f"自动恢复成功：{target.name}及音频均已生成"
        else:
            raise RuntimeError(f"尚未支持自动恢复的任务类型：{kind or '-'}")
        append_general_task_log(task_id, detail)
        finish_general_task_run(task_id, True, detail)
    except Exception as exc:
        detail = f"第{retry_count}次自动恢复失败：{exc}"
        append_general_task_log(task_id, detail)
        finish_general_task_run(task_id, False, detail)


def schedule_interrupted_general_task_retries(interrupted: list[dict]) -> list[str]:
    """Queue bounded retries for every safely replayable general task."""
    scheduled: list[str] = []
    supported = {"weekly-report", "carrier-performance", "audio-generation"}
    seen: set[tuple[str, str]] = set()
    for task in interrupted:
        kind = str(task.get("kind") or "")
        retry_count = int(task.get("retry_count") or 0)
        # The index is newest-first. If several restarts interrupted the same
        # logical operation, retry only the newest record instead of launching
        # duplicate reports/audio jobs after recovery.
        recovery_key = (kind, str(task.get("target_path") or task.get("scope") or ""))
        if recovery_key in seen:
            _mark_general_task_recovery_disposition(
                str(task.get("task_id") or ""), "superseded_by_newer_interruption"
            )
            continue
        seen.add(recovery_key)
        if kind not in supported:
            append_general_task_log(str(task.get("task_id") or ""), "[自动恢复] 任务类型不支持安全重放，未自动重试。")
            _mark_general_task_recovery_disposition(
                str(task.get("task_id") or ""), "unsupported"
            )
            continue
        if retry_count >= GENERAL_TASK_MAX_AUTO_RETRIES:
            append_general_task_log(
                str(task.get("task_id") or ""),
                f"[自动恢复] 已达最大{GENERAL_TASK_MAX_AUTO_RETRIES}次，停止自动重试。",
            )
            _mark_general_task_recovery_disposition(
                str(task.get("task_id") or ""), "retry_limit_reached"
            )
            continue
        timer = threading.Timer(
            GENERAL_TASK_RETRY_DELAY_SECONDS,
            _run_recovered_general_task,
            args=(dict(task),),
        )
        timer.name = f"task-auto-retry-{task.get('task_run_id') or 'unknown'}"
        timer.daemon = True
        timer.start()
        _mark_general_task_recovery_disposition(
            str(task.get("task_id") or ""), "scheduled"
        )
        scheduled.append(str(task.get("task_id") or ""))
    return scheduled


def reconcile_misclassified_general_tasks() -> list[dict]:
    """Correct legacy report tasks that completed the DOCX but failed required audio generation."""
    corrected: list[dict] = []
    with TASK_RUNS_LOCK:
        tasks = _task_read_local_index()
        for task in tasks:
            if task.get("run_status") != "completed" or task.get("kind") not in {"weekly-report", "carrier-performance"}:
                continue
            relative_log = str(task.get("log_path") or "")
            if not relative_log:
                continue
            log_path = ROOT / relative_log
            try:
                content = log_path.read_text(encoding="utf-8", errors="replace")
            except OSError:
                continue
            failure_index = max(
                content.rfind("❌ 语音摘要生成失败"),
                content.rfind("Audio generation failed"),
            )
            success_index = content.rfind("语音摘要生成完成")
            if failure_index < 0 or success_index > failure_index:
                continue
            failure_line = content[failure_index:].splitlines()[0].strip()
            detail = failure_line or "语音摘要生成失败。"
            task.update(
                {
                    "run_status": "failed",
                    "status_detail": detail,
                    "phase": "失败",
                    "progress_detail": detail,
                    "worker_pid": 0,
                }
            )
            raw_id = str(task.get("task_run_id") or str(task.get("task_id") or "").removeprefix("task:"))
            _task_atomic_json(TASK_RUNS_DIR / (raw_id + ".json"), task)
            corrected.append(dict(task))
        if corrected:
            _task_atomic_json(TASK_RUNS_INDEX_PATH, {"tasks": tasks[:500]})
    return corrected


def _normalize_crawl_task(run: dict) -> dict:
    crawl_id = str(run.get("crawl_run_id") or "")
    stream = run.get("stream_log") if isinstance(run.get("stream_log"), dict) else {}
    task_kind = str(run.get("task_kind") or "crawl")
    kind_labels = {
        "strategic-news": "新闻爬虫",
        "executive-intelligence-refresh": "四库刷新",
        "crawl": "爬虫",
    }
    operational_summary = run.get("operational_summary") if isinstance(run.get("operational_summary"), dict) else {}
    model_analysis = operational_summary.get("model_analysis") if isinstance(operational_summary.get("model_analysis"), dict) else {}
    pages_publish = operational_summary.get("pages_publish") if isinstance(operational_summary.get("pages_publish"), dict) else {}
    return {
        "task_id": "crawl:" + crawl_id,
        "task_run_id": crawl_id,
        "kind": task_kind,
        "kind_label": kind_labels.get(task_kind, "后台任务"),
        "title": str(run.get("trigger") or "爬虫任务"),
        "scope": str(run.get("scope") or "未记录范围"),
        "run_status": str(run.get("run_status") or "completed"),
        "started_at_hkt": str(run.get("started_at_hkt") or ""),
        "completed_at_hkt": str(run.get("completed_at_hkt") or ""),
        "duration_ms": int(run.get("duration_ms") or 0),
        "lines": int(stream.get("lines") or 0),
        "bytes": int(stream.get("bytes") or 0),
        "status_detail": str(run.get("status_detail") or ""),
        "interrupted": bool(run.get("interrupted")),
        "backend_pid": int(run.get("backend_pid") or 0),
        "worker_pid": int(run.get("worker_pid") or 0),
        "phase": str(run.get("phase") or ""),
        "progress_detail": str(run.get("progress_detail") or ""),
        "heartbeat_at_hkt": str(run.get("heartbeat_at_hkt") or ""),
        "analysis_model": str(model_analysis.get("model") or ""),
        "analysis_fallback_used": bool(model_analysis.get("fallback_used")),
        "analysis_fallback_reason": str(model_analysis.get("fallback_reason") or ""),
        "evidence_hash": str(model_analysis.get("evidence_hash") or ""),
        "pages_publish_ok": bool(pages_publish.get("ok")),
        "pages_publish_status": str(pages_publish.get("status") or ""),
        "pages_public_url": str(pages_publish.get("public_url") or ""),
        "pages_site_version": str(pages_publish.get("site_version") or ""),
        "pages_publish_error": str(pages_publish.get("error") or ""),
        "source": "crawl-archive",
    }


def _annotate_strategic_task_retries(tasks: list[dict]) -> None:
    """Number repeated attempts for the same durable strategic scan slot."""
    attempts: dict[tuple[str, str, str], int] = {}
    ordered = sorted(
        tasks,
        key=lambda item: str(
            item.get("started_at_hkt") or item.get("completed_at_hkt") or ""
        ),
    )
    for task in ordered:
        if str(task.get("kind") or "") != "strategic-news":
            task["retry_index"] = 0
            continue
        key = (
            str(task.get("kind") or ""),
            str(task.get("title") or ""),
            str(task.get("scope") or ""),
        )
        retry_index = attempts.get(key, 0)
        task["retry_index"] = retry_index
        attempts[key] = retry_index + 1


def _project_monitor_handlers() -> dict[str, dict]:
    """Merge Feishu-card and dashboard checkbox actions by latest handled time."""
    handlers: dict[str, dict] = {}
    try:
        action_state = json.loads(PROJECT_MONITOR_ACTIONS_PATH.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        action_state = {}
    handled_messages = action_state.get("handled_messages") if isinstance(action_state, dict) else {}
    candidates = list(handled_messages.values()) if isinstance(handled_messages, dict) else []
    try:
        for line in PROJECT_MONITOR_WEB_ACTIONS_PATH.read_text(encoding="utf-8").splitlines():
            item = json.loads(line)
            if isinstance(item, dict):
                candidates.append(item)
    except (OSError, ValueError, TypeError):
        pass
    for handled in candidates:
        if not isinstance(handled, dict):
            continue
        incident_id = str(handled.get("incident_id") or "")
        if not incident_id:
            continue
        previous = handlers.get(incident_id, {})
        handled_at = str(handled.get("handled_at_hkt") or handled.get("completed_at_hkt") or "")
        previous_at = str(previous.get("handled_at_hkt") or previous.get("completed_at_hkt") or "")
        if handled_at >= previous_at:
            handlers[incident_id] = handled
    return handlers


def _handler_public_fields(handled: dict) -> dict[str, str]:
    open_id = str(handled.get("operator_id") or "")
    user = AUTH.public_user_by_feishu_open_id(open_id) or {}
    return {
        "handler_id": str(user.get("id") or ""),
        "handler_open_id": open_id,
        "handler_name": str(handled.get("operator_name") or user.get("name") or ""),
        "handler_avatar_url": str(user.get("avatarUrl") or ""),
        "handled_at_hkt": str(handled.get("handled_at_hkt") or handled.get("completed_at_hkt") or ""),
    }


def attach_news_review_actors(snapshot: dict) -> dict:
    """Attach the latest authenticated human reviewer to each reviewed row."""
    reviewers: dict[int, dict[str, str]] = {}
    for event in AUTH.operation_audit(limit=1000):
        if event.get("action") != "news_review.update" or event.get("result") != "success":
            continue
        details = event.get("details") if isinstance(event.get("details"), dict) else {}
        decision_rows = details.get("decision_rows") if isinstance(details.get("decision_rows"), list) else []
        reviewer = {
            "id": str(event.get("actor_id") or ""),
            "name": str(event.get("actor_name") or "未知用户"),
            "avatarUrl": str(event.get("actor_avatar_url") or ""),
            "reviewedAt": str(event.get("at") or ""),
        }
        for raw_row_number in decision_rows:
            try:
                row_number = int(raw_row_number)
            except (TypeError, ValueError):
                continue
            reviewers.setdefault(row_number, reviewer)
    for row in snapshot.get("rows") or []:
        if not isinstance(row, dict):
            continue
        try:
            reviewer = reviewers.get(int(row.get("rowNumber") or 0))
        except (TypeError, ValueError):
            reviewer = None
        if reviewer:
            row["reviewer"] = reviewer
    return snapshot


def _task_incident_metadata() -> dict[str, dict]:
    try:
        monitor_state = json.loads(PROJECT_MONITOR_STATE_PATH.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        monitor_state = {}
    incidents = monitor_state.get("incidents") if isinstance(monitor_state, dict) else {}
    incidents = incidents if isinstance(incidents, dict) else {}

    handlers = _project_monitor_handlers()

    severity_labels = {"P1": "紧急", "P2": "高", "P3": "中"}
    metadata: dict[str, dict] = {}
    for incident in incidents.values():
        if not isinstance(incident, dict):
            continue
        condition_key = str(incident.get("condition_key") or "")
        if not condition_key:
            continue
        incident_id = str(incident.get("incident_id") or "")
        handled = handlers.get(incident_id, {})
        severity = str(incident.get("severity") or "")
        metadata[condition_key] = {
            "incident_id": incident_id,
            "incident_status": str(incident.get("status") or ""),
            "severity": severity,
            "severity_label": severity_labels.get(severity, ""),
            **_handler_public_fields(handled),
        }
    return metadata


def _annotate_task_incidents(tasks: list[dict]) -> None:
    metadata = _task_incident_metadata()
    for task in tasks:
        task_id = str(task.get("task_id") or task.get("task_run_id") or "")
        run_id = str(task.get("task_run_id") or task_id.removeprefix("crawl:"))
        prefixes = ("crawl-task-failed", "crawl-task-stuck") if task_id.startswith("crawl:") else ("general-task-failed", "general-task-stuck")
        keys = [f"{prefix}:{run_id if prefix.startswith('crawl-') else task_id}" for prefix in prefixes]
        incident = next((metadata[key] for key in keys if key in metadata), None)
        if incident:
            task.update(incident)


def load_project_incident_index(limit: int = 100) -> list[dict]:
    """Return the real project-monitor incident ledger for the alarm screen."""
    try:
        monitor_state = json.loads(PROJECT_MONITOR_STATE_PATH.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        monitor_state = {}
    incidents = monitor_state.get("incidents") if isinstance(monitor_state, dict) else {}
    incidents = incidents if isinstance(incidents, dict) else {}

    handlers = _project_monitor_handlers()

    severity_labels = {"P1": "紧急", "P2": "高", "P3": "中"}
    records: list[dict] = []
    for incident in incidents.values():
        if not isinstance(incident, dict):
            continue
        incident_id = str(incident.get("incident_id") or "")
        if not incident_id:
            continue
        diagnosis = incident.get("diagnosis") if isinstance(incident.get("diagnosis"), dict) else {}
        severity = str(diagnosis.get("severity") or incident.get("severity") or "P2").upper()
        if severity not in severity_labels:
            severity = "P2"
        handled = handlers.get(incident_id, {})
        suggestions = diagnosis.get("solutions") or diagnosis.get("suggestions") or incident.get("suggestions") or []
        if not isinstance(suggestions, list):
            suggestions = [str(suggestions)] if suggestions else []
        evidence = incident.get("evidence") if isinstance(incident.get("evidence"), list) else []
        status = str(incident.get("status") or "open")
        records.append({
            "task_id": f"incident:{incident_id}",
            "incident_id": incident_id,
            "incident_status": status,
            "kind": str(incident.get("component") or "project-monitor"),
            "kind_label": str(incident.get("task_name") or "项目故障"),
            "title": str(incident.get("task_name") or incident.get("summary") or "项目故障"),
            "scope": str(incident.get("component") or incident.get("condition_key") or "项目监控"),
            "run_status": "failed" if status == "open" else "completed",
            "severity": severity,
            "severity_label": severity_labels[severity],
            **_handler_public_fields(handled),
            "summary": str(incident.get("summary") or ""),
            "error": str(diagnosis.get("fault_cause") or incident.get("error") or ""),
            "impact": str(diagnosis.get("fault_impact") or incident.get("impact") or ""),
            "suggestions": [str(item) for item in suggestions if str(item).strip()],
            "evidence": [str(item) for item in evidence if str(item).strip()],
            "phase": "已处理" if handled else ("待处理" if status == "open" else "已恢复"),
            "occurred_at_hkt": str(incident.get("occurred_at_hkt") or incident.get("first_seen_at_hkt") or ""),
            "started_at_hkt": str(incident.get("first_seen_at_hkt") or incident.get("occurred_at_hkt") or ""),
            "heartbeat_at_hkt": str(incident.get("last_seen_at_hkt") or ""),
            "completed_at_hkt": str(incident.get("resolved_at_hkt") or ""),
            "source": "project-monitor",
        })
    records.sort(
        key=lambda item: str(item.get("occurred_at_hkt") or item.get("started_at_hkt") or ""),
        reverse=True,
    )
    return records[: max(1, min(500, int(limit or 100)))]


def count_project_incidents() -> int:
    """Count every valid incident in the ledger independently of the page limit."""
    try:
        monitor_state = json.loads(PROJECT_MONITOR_STATE_PATH.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        return 0
    incidents = monitor_state.get("incidents") if isinstance(monitor_state, dict) else {}
    if not isinstance(incidents, dict):
        return 0
    return sum(
        1
        for incident in incidents.values()
        if isinstance(incident, dict) and str(incident.get("incident_id") or "")
    )


def load_unified_task_index(limit: int = 50) -> list[dict]:
    tasks = [_task_public_record(item) for item in _task_read_local_index()]
    tasks.extend(
        _normalize_crawl_task(item)
        for item in load_crawl_run_index()
        if isinstance(item, dict) and item.get("crawl_run_id")
    )
    _annotate_strategic_task_retries(tasks)
    _annotate_task_incidents(tasks)
    tasks.sort(
        key=lambda item: str(item.get("started_at_hkt") or item.get("completed_at_hkt") or ""),
        reverse=True,
    )
    return tasks[:limit]


def load_unified_task_log(task_id: str) -> dict:
    task_id = str(task_id or "").strip()
    if task_id.startswith("crawl:"):
        crawl_id = task_id.removeprefix("crawl:")
        result = load_crawl_run_log(crawl_id)
        if result.get("ok"):
            run = result.get("run") if isinstance(result.get("run"), dict) else {}
            result["task"] = _normalize_crawl_task(run)
            indexed = next(
                (
                    item
                    for item in load_unified_task_index(limit=500)
                    if str(item.get("task_id") or "") == task_id
                ),
                {},
            )
            result["task"]["retry_index"] = int(
                indexed.get("retry_index") or 0
            )
        return result
    if not task_id.startswith("task:"):
        return {"ok": False, "error": "无效的任务编号。"}
    record = next(
        (item for item in _task_read_local_index() if str(item.get("task_id") or "") == task_id),
        None,
    )
    if not record:
        return {"ok": False, "error": "未找到该任务记录。"}
    task = _task_public_record(record)
    relative_path = str(record.get("log_path") or "")
    log_path = ROOT / relative_path if relative_path else None
    try:
        content = log_path.read_text(encoding="utf-8", errors="replace") if log_path and log_path.exists() else ""
    except OSError as exc:
        return {"ok": False, "error": "任务日志读取失败：" + str(exc)}
    run = {
        "run_status": task.get("run_status"),
        "started_at_hkt": task.get("started_at_hkt"),
        "completed_at_hkt": task.get("completed_at_hkt"),
        "duration_ms": task.get("duration_ms"),
    }
    return {
        "ok": True,
        "task": task,
        "run": run,
        "content": content,
        "lines": int(task.get("lines") or 0),
        "bytes": int(task.get("bytes") or 0),
    }


def _task_phase_from_payload(payload: dict, current: str) -> tuple[str, str]:
    event_type = str(payload.get("type") or "")
    if event_type == "agent_trace":
        trace = payload.get("trace") if isinstance(payload.get("trace"), dict) else {}
        return str(trace.get("node") or "Agent 审核"), str(trace.get("message") or "Agent 正在处理审核节点。")
    if event_type == "done":
        return "任务收尾", str(payload.get("message") or "执行步骤已结束，正在持久化最终状态。")
    text = str(payload.get("text") or payload.get("message") or "").strip()
    if not text:
        return current or "执行中", "任务仍在执行，等待下一条业务进度。"
    if "语音" in text or "TTS" in text:
        phase = "生成语音摘要"
    elif "飞书" in text or "同步" in text:
        phase = "飞书同步"
    elif any(token in text for token in ("搜索验证", "事实抽取", "质量审计", "冲突仲裁", "主体校验", "Agent")):
        phase = "Agent 审核"
    elif "补爬" in text:
        phase = "缺口补爬"
    elif any(token in text for token in ("crawl row", "抓取", "状态码", "URL")):
        phase = "网页抓取"
    elif any(token in text for token in ("报告", "模板", "周报", "业绩摘要")):
        phase = "报告生成"
    else:
        phase = current or "执行中"
    return phase, text[:360]


def observe_task_progress(handler: BaseHTTPRequestHandler, payload: dict) -> None:
    if not isinstance(payload, dict) or not getattr(handler, "_task_monitor_kind", ""):
        return
    phase, detail = _task_phase_from_payload(payload, str(getattr(handler, "_task_monitor_phase", "") or ""))
    handler._task_monitor_phase = phase
    handler._task_monitor_detail = detail
    if getattr(handler, "_task_monitor_kind", "") == "crawl":
        CRAWL_PIPELINE_STATE.update({"phase": phase, "detail": detail})


def _task_monitor_loop(handler: BaseHTTPRequestHandler, owner: threading.Thread) -> None:
    stop_event = handler._task_monitor_stop
    while not stop_event.wait(TASK_HEARTBEAT_INTERVAL_SECONDS):
        task_id = str(getattr(handler, "_task_monitor_id", "") or "")
        kind = str(getattr(handler, "_task_monitor_kind", "") or "")
        phase = str(getattr(handler, "_task_monitor_phase", "") or "执行中")
        detail = str(getattr(handler, "_task_monitor_detail", "") or "任务仍在执行，等待下一条业务进度。")
        worker_pid = int(getattr(handler, "_task_worker_pid", 0) or 0)
        if not owner.is_alive():
            reason = f"任务执行线程意外结束；最后阶段：{phase}；最后进度：{detail}"
            if kind == "crawl":
                mark_crawl_run_interrupted(task_id, reason)
            elif kind == "general":
                append_general_task_log(task_id, "[任务中断] " + reason)
                finish_general_task_run(task_id, False, reason)
            return
        if kind == "crawl":
            heartbeat_crawl_run(task_id, phase, detail, worker_pid=worker_pid, append_log=True)
        elif kind == "general":
            heartbeat_general_task_run(task_id, phase, detail, worker_pid=worker_pid, append_log=True)


def start_task_lifecycle_monitor(
    handler: BaseHTTPRequestHandler,
    kind: str,
    task_id: str,
    phase: str,
) -> None:
    handler._task_monitor_kind = kind
    handler._task_monitor_id = task_id
    handler._task_monitor_phase = phase
    handler._task_monitor_detail = "后台任务已启动，持续监控中。"
    handler._task_worker_pid = 0
    handler._task_monitor_stop = threading.Event()
    if kind == "crawl":
        heartbeat_crawl_run(task_id, phase, handler._task_monitor_detail, append_log=False)
    else:
        heartbeat_general_task_run(task_id, phase, handler._task_monitor_detail, append_log=False)
    monitor = threading.Thread(
        target=_task_monitor_loop,
        args=(handler, threading.current_thread()),
        name=f"task-monitor-{task_id}",
        daemon=True,
    )
    handler._task_monitor_thread = monitor
    monitor.start()


def stop_task_lifecycle_monitor(handler: BaseHTTPRequestHandler) -> None:
    stop_event = getattr(handler, "_task_monitor_stop", None)
    if stop_event:
        stop_event.set()


_ORIGINAL_WRITE_SSE = write_sse
_ORIGINAL_STREAM_REPORT_GENERATION = stream_report_generation


def write_sse(handler: BaseHTTPRequestHandler, payload: dict) -> None:
    observe_task_progress(handler, payload)
    task_id = str(getattr(handler, "_general_task_run_id", "") or "")
    if task_id and isinstance(payload, dict):
        event_type = str(payload.get("type") or "")
        if event_type == "log" and payload.get("text"):
            append_general_task_log(task_id, payload.get("text"))
        elif event_type == "done":
            ok = bool(payload.get("ok", True))
            detail = str(payload.get("error") or payload.get("message") or "")
            append_general_task_log(task_id, "任务完成。" if ok else "任务失败：" + (detail or "未提供原因"))
            finish_general_task_run(task_id, ok, detail)
            handler._general_task_finished = True
    _ORIGINAL_WRITE_SSE(handler, payload)


def stream_report_generation(handler: BaseHTTPRequestHandler, script_name: str, report_kind: str) -> None:
    if report_kind == "weekly":
        kind = "weekly-report"
        title = "生成周报"
        scope = "战略部每周周报"
    else:
        kind = "carrier-performance"
        title = "生成业绩摘要"
        scope = "运营商业绩摘要"
    task = start_general_task_run(kind, title, scope, script_name)
    task_id = str(task["task_id"])
    handler._general_task_run_id = task_id
    handler._general_task_finished = False
    start_task_lifecycle_monitor(handler, "general", task_id, "报告生成")
    try:
        result = _ORIGINAL_STREAM_REPORT_GENERATION(handler, script_name, report_kind)
        if not handler._general_task_finished:
            append_general_task_log(task_id, "任务执行结束。")
            finish_general_task_run(task_id, True, "")
            handler._general_task_finished = True
        return result
    except Exception as exc:
        append_general_task_log(task_id, "任务异常：" + str(exc))
        finish_general_task_run(task_id, False, str(exc))
        handler._general_task_finished = True
        raise
    finally:
        stop_task_lifecycle_monitor(handler)
        handler._general_task_run_id = ""


def start_audio_generation_task(target: Path, force: bool = True) -> tuple[dict, bool]:
    target_key = str(target.resolve())
    with TASK_RUNS_LOCK:
        existing = next(
            (
                task
                for task in _task_read_local_index()
                if task.get("kind") == "audio-generation"
                and task.get("run_status") == "running"
                and str(task.get("target_path") or "") == target_key
            ),
            None,
        )
    if existing:
        return _task_public_record(existing), False

    task = start_general_task_run(
        "audio-generation",
        "生成音频摘要",
        target.name,
        "tts_service.py",
    )
    task_id = str(task["task_id"])
    raw_id = str(task["task_run_id"])
    with TASK_RUNS_LOCK:
        tasks = _task_read_local_index()
        for record in tasks:
            if str(record.get("task_id") or "") == task_id:
                record["target_path"] = target_key
                task = dict(record)
                break
        _task_atomic_json(TASK_RUNS_INDEX_PATH, {"tasks": tasks[:500]})
        _task_atomic_json(TASK_RUNS_DIR / (raw_id + ".json"), task)

    def worker() -> None:
        code = (
            "import sys, json\n"
            "from pathlib import Path\n"
            "from tts_service import synthesize_report_audio\n"
            "try:\n"
            "    res = synthesize_report_audio(Path(sys.argv[1]), force=sys.argv[2] == 'True')\n"
            "    print(json.dumps({'ok': True, 'result': res}))\n"
            "except Exception as e:\n"
            "    print(json.dumps({'ok': False, 'error': str(e)}))\n"
        )
        proc_audio: subprocess.Popen[str] | None = None
        try:
            append_general_task_log(task_id, f"开始为报告生成音频摘要：{target.name}")
            proc_audio = subprocess.Popen(
                [sys.executable, "-c", code, str(target), str(bool(force))],
                cwd=str(ROOT),
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            heartbeat_general_task_run(
                task_id,
                "生成语音摘要",
                "公司内网 TTS 正在生成并统一处理音频。",
                worker_pid=proc_audio.pid,
                append_log=True,
            )
            while proc_audio.poll() is None:
                try:
                    proc_audio.wait(timeout=TASK_HEARTBEAT_INTERVAL_SECONDS)
                except subprocess.TimeoutExpired:
                    heartbeat_general_task_run(
                        task_id,
                        "生成语音摘要",
                        "公司内网 TTS 仍在处理，任务持续跟踪中。",
                        worker_pid=proc_audio.pid,
                        append_log=False,
                    )
            stdout, stderr = proc_audio.communicate()
            try:
                payload = json.loads(stdout or "{}")
            except json.JSONDecodeError as exc:
                raise RuntimeError(f"音频服务返回无法解析：{stderr.strip() or stdout.strip() or exc}") from exc
            if proc_audio.returncode != 0 or not payload.get("ok"):
                raise RuntimeError(str(payload.get("error") or stderr.strip() or "音频生成失败"))
            result = payload.get("result") if isinstance(payload.get("result"), dict) else {}
            if not result.get("ok"):
                raise RuntimeError(str(result.get("error") or "音频生成失败"))
            audio = result.get("audio") if isinstance(result.get("audio"), dict) else {}
            backend = str(result.get("backend") or "unknown")
            audio_name = str(audio.get("name") or "音频摘要")
            detail = f"音频摘要已生成：{audio_name}（{backend}）"
            append_general_task_log(task_id, detail)
            finish_general_task_run(task_id, True, detail)
        except Exception as exc:
            detail = str(exc) or "音频生成失败"
            append_general_task_log(task_id, "音频生成失败：" + detail)
            finish_general_task_run(task_id, False, detail)

    threading.Thread(
        target=worker,
        name="audio-task-" + raw_id,
        daemon=True,
    ).start()
    return _task_public_record(task), True


class AppHandler(BaseHTTPRequestHandler):
    server_version = "WeeklyReportUI/1.0"

    @staticmethod
    def download_disposition(path: Path) -> str:
        encoded_name = quote(path.name, safe="")
        fallback_name = f"weekly-report{path.suffix.lower() or '.docx'}"
        return f"attachment; filename=\"{fallback_name}\"; filename*=UTF-8''{encoded_name}"

    def do_HEAD(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path in {"/", "/company-data", "/company-data.html", "/executive-dashboard-demo", "/executive-dashboard-demo.html", "/static/index.html"}:
            if not AUTH.authorize_page(self, parsed.path):
                return
        if parsed.path.startswith(("/static/", "/outputs/", "/audio/", "/generated-charts/", "/references/", "/references-raw/")):
            if not AUTH.authorize_resource(self, parsed.path):
                return
        if parsed.path == "/":
            self.serve_head(STATIC_DIR / "index.html")
            return
        if parsed.path in {"/company-data", "/company-data.html"}:
            self.serve_head(STATIC_DIR / "company-data.html")
            return
        if parsed.path.startswith("/static/"):
            self.serve_head(STATIC_DIR / parsed.path.removeprefix("/static/"))
            return
        if parsed.path.startswith("/outputs/"):
            name = Path(unquote(parsed.path.removeprefix("/outputs/"))).name
            target = ROOT / name
            if is_report_path(target):
                self.serve_head(target, download=True)
                return
        if parsed.path.startswith("/audio/"):
            name = Path(unquote(parsed.path.removeprefix("/audio/"))).name
            target = AUDIO_DIR / name
            if target.exists() and target.suffix.lower() in {".wav", ".mp3"}:
                self.serve_head(target)
                return
        if parsed.path.startswith("/generated-charts/"):
            target = generated_chart_path(unquote(parsed.path.removeprefix("/generated-charts/")))
            if target and target.exists():
                self.serve_head(target)
                return
        if parsed.path.startswith("/references/"):
            target = reference_path(unquote(parsed.path.removeprefix("/references/")))
            if target and target.exists():
                self.serve_reference_head(target)
                return
        if parsed.path.startswith("/references-raw/"):
            target = reference_path(unquote(parsed.path.removeprefix("/references-raw/")))
            if target and target.exists():
                self.serve_head(target)
                return
        if parsed.path.startswith("/archives/"):
            target = ROOT / unquote(parsed.path.lstrip("/"))
            if is_report_path(target):
                self.serve_head(target, download=True)
                return
        self.send_response(404)
        self.end_headers()

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path
        if AUTH.handle(self, "GET", parsed):
            return
        if path in {"/", "/company-data", "/company-data.html", "/executive-dashboard-demo", "/executive-dashboard-demo.html", "/static/index.html"}:
            if not AUTH.authorize_page(self, path):
                return
        if path.startswith(("/static/", "/outputs/", "/audio/", "/generated-charts/", "/references/", "/references-raw/")):
            if not AUTH.authorize_resource(self, path):
                return
        if path.startswith("/api/") and not AUTH.authorize_api(self, path, "GET"):
            return
        if path == "/":
            self.serve_file(STATIC_DIR / "index.html")
            return
        if path in {"/company-data", "/company-data.html"}:
            self.serve_file(STATIC_DIR / "company-data.html")
            return
        if path in {"/executive-dashboard-demo", "/executive-dashboard-demo.html"}:
            self.serve_file(STATIC_DIR / "executive-dashboard-demo.html")
            return
        if path == "/api/status":
            json_response(self, {"ok": True, "status": build_status()})
            return
        if path == "/api/health":
            json_response(self, {"ok": True, "status": build_status()})
            return
        if path == "/api/weekly-report-preview":
            try:
                json_response(self, {"ok": True, "preview": build_weekly_report_generation_preview()})
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, status=503)
            return
        if path == "/api/report-audio":
            path_str = (parse_qs(parsed.query).get("path") or [""])[0]
            target = report_target_from_rel(path_str)
            if not target:
                json_response(self, {"ok": False, "error": "report file not found"}, 404)
                return
            json_response(self, {"ok": True, "audio": audio_info_for_report(target)})
            return
        if path == "/api/chat-starters":
            json_response(self, {"ok": True, "starters": sample_chat_starters()})
            return
        if path == "/api/strategic-briefs":
            try:
                from strategic_briefing import public_snapshot

                json_response(self, {"ok": True, **public_snapshot()})
            except Exception as exc:
                json_response(
                    self,
                    {"ok": False, "error": str(exc), "items": []},
                    status=500,
                )
            return
        if path == "/api/subscriptions":
            if not is_loopback_client(str(self.client_address[0])):
                json_response(self, {"ok": False, "error": "订阅管理后台仅允许本机访问"}, status=403)
                return
            try:
                service = subscription_service()
                summary = service.list_summary()
                status = build_status()
                reports = [
                    {
                        "name": str(item.get("name") or ""),
                        "path": str(item.get("path_str") or ""),
                        "report_type": str(item.get("reportType") or ""),
                        "mtime_text": str(item.get("mtimeText") or ""),
                        "audio": bool((item.get("audio") or {}).get("exists")) if isinstance(item.get("audio"), dict) else False,
                    }
                    for item in (status.get("outputs") or [])
                    if isinstance(item, dict) and item.get("reportType") in {"weekly", "carrier-performance"}
                ]
                card_actions = service.config.get("card_actions") if isinstance(service.config.get("card_actions"), dict) else {}
                json_response(self, {
                    "ok": True,
                    **summary,
                    "targets": service.available_targets(),
                    "frequencies": [
                        {"key": key, "label": FREQUENCY_LABELS[key]}
                        for key in ("twice_daily", "once_daily")
                    ],
                    "frequency_scope": "news",
                    "report_cadence": {
                        "key": "biweekly_on_publish",
                        "label": REPORT_CADENCE_LABEL,
                    },
                    "report_modes": [
                        {"key": key, "label": REPORT_MODE_LABELS[key]}
                        for key in ("pdf", "pdf_audio", "audio")
                    ],
                    "news_categories": [
                        {"key": key, "label": label}
                        for key, label in NEWS_CATEGORY_LABELS.items()
                    ],
                    "reports": reports,
                    "test_target": {
                        "callback_open_id": str(card_actions.get("primary_handler_open_id") or ""),
                        "delivery_open_id": str(card_actions.get("primary_handler_open_id") or ""),
                        "name": str(card_actions.get("primary_handler_expected_name") or "系统管理员"),
                    },
                })
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/subscriptions/avatar":
            if not is_loopback_client(str(self.client_address[0])):
                json_response(self, {"ok": False, "error": "订阅管理后台仅允许本机访问"}, status=403)
                return
            try:
                open_id = (parse_qs(parsed.query).get("openId") or [""])[0]
                source_url = subscription_service().avatar_source_url(open_id)
                request = urllib.request.Request(source_url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(request, timeout=15) as response:
                    body = response.read(2_000_000)
                    content_type = str(response.headers.get_content_type() or "image/png")
                if not content_type.startswith("image/") or not body:
                    raise ValueError("飞书头像返回格式无效")
                self.send_response(200)
                self.send_header("Content-Type", content_type)
                self.send_header("Content-Length", str(len(body)))
                self.send_header("Cache-Control", "private, max-age=3600")
                self.send_header("X-Content-Type-Options", "nosniff")
                self.end_headers()
                self.wfile.write(body)
            except ValueError as exc:
                json_response(self, {"ok": False, "error": str(exc)}, status=404)
            except Exception as exc:
                json_response(self, {"ok": False, "error": f"飞书头像读取失败：{exc}"}, status=502)
            return
        if path == "/api/news-review-sheet":
            try:
                from news_review_sheet import review_sheet_snapshot

                json_response(self, {"ok": True, **attach_news_review_actors(review_sheet_snapshot())})
            except Exception as exc:
                json_response(
                    self,
                    {"ok": False, "error": str(exc), "headers": [], "rows": []},
                    status=503,
                )
            return
        if path == "/api/executive-intelligence":
            try:
                from executive_intelligence import build_executive_intelligence_snapshot

                json_response(
                    self,
                    {"ok": True, **build_executive_intelligence_snapshot()},
                )
            except Exception as exc:
                json_response(
                    self,
                    {"ok": False, "error": str(exc), "domains": [], "relations": []},
                    status=500,
                )
            return
        if path == "/api/scheduler-overview":
            try:
                json_response(self, build_scheduler_overview())
            except Exception as exc:
                json_response(
                    self,
                    {"ok": False, "error": str(exc), "configured_rows": 0, "source_groups": [], "next_runs": []},
                    status=503,
                )
            return
        if path == "/api/executive-company-benchmarks":
            try:
                json_response(self, build_company_benchmarks())
            except Exception as exc:
                json_response(
                    self,
                    {"ok": False, "error": str(exc), "companies": [], "metrics": {}, "values": {}},
                    status=500,
                )
            return
        if path == "/api/company-metrics":
            json_response(
                self,
                {
                    "ok": True,
                    "data": build_company_metrics_payload(),
                    "curation": load_curation_status(),
                },
            )
            return
        if path == "/api/data-curation":
            json_response(self, {"ok": True, "curation": load_curation_status()})
            return
        if path == "/api/agent-trace":
            query = parse_qs(parsed.query)
            try:
                limit = max(1, min(1000, int(query.get("limit", ["300"])[0])))
            except Exception:
                limit = 300
            json_response(
                self,
                {
                    "ok": True,
                    "trace": load_agent_trace(limit=limit),
                    "summary": load_curation_status(),
                },
            )
            return
        if path == "/api/agent-skills":
            json_response(self, {"ok": True, "skills": available_agent_skills()})
            return
        if path == "/api/agent-runs":
            query = parse_qs(parsed.query)
            try:
                limit = max(1, min(100, int(query.get("limit", ["20"])[0])))
            except Exception:
                limit = 20
            json_response(self, {"ok": True, "runs": list_agent_runs(limit=limit)})
            return
        if path == "/api/agent-memory":
            query = parse_qs(parsed.query)
            try:
                limit = max(1, min(100, int(query.get("limit", ["50"])[0])))
            except Exception:
                limit = 50
            json_response(self, {"ok": True, "memories": load_memories(limit=limit)})
            return
        if path == "/api/chat-threads":
            query = parse_qs(parsed.query)
            thread_id = str(query.get("id", [""])[0] or "")
            if thread_id:
                thread = get_chat_thread(thread_id)
                json_response(self, {"ok": bool(thread), "thread": thread}, 200 if thread else 404)
            else:
                json_response(self, {"ok": True, "threads": chat_thread_summaries()})
            return
        if path == "/api/agent-dataset-lineage":
            query = parse_qs(parsed.query)
            raw_ids = query.get("datasetId", []) + query.get("datasetIds", [])
            dataset_ids = {item for raw in raw_ids for item in str(raw).split(",") if item}
            json_response(self, {"ok": True, "lineage": dataset_lineage(dataset_ids or None)})
            return
        if path == "/api/agent-datasets":
            json_response(
                self,
                {
                    "ok": True,
                    "root": "agent_knowledge",
                    "allowedExtensions": sorted(UPLOAD_ALLOWED_SUFFIXES),
                    "datasets": list_knowledge_datasets(),
                },
            )
            return
        if path == "/api/task-runs":
            query = parse_qs(parsed.query)
            try:
                limit = max(1, min(100, int(query.get("limit", ["50"])[0])))
            except Exception:
                limit = 50
            json_response(self, {"ok": True, "tasks": load_unified_task_index(limit)})
            return
        if path == "/api/project-incidents":
            query = parse_qs(parsed.query)
            try:
                limit = max(1, min(500, int(query.get("limit", ["100"])[0])))
            except Exception:
                limit = 100
            incidents = load_project_incident_index(limit)
            json_response(self, {"ok": True, "incidents": incidents, "total": count_project_incidents()})
            return
        if path == "/api/task-run-log":
            query = parse_qs(parsed.query)
            task_id = str(query.get("id", [""])[0] or "")
            result = load_unified_task_log(task_id)
            json_response(self, result, 200 if result.get("ok") else 404)
            return
        if path == "/api/crawl-runs":
            query = parse_qs(parsed.query)
            try:
                limit = max(1, min(500, int(query.get("limit", ["20"])[0])))
            except Exception:
                limit = 20
            task_kind = str(query.get("taskKind", [""])[0] or "").strip()
            runs = load_crawl_run_history(task_kind=task_kind)
            json_response(
                self,
                {
                    "ok": True,
                    "runs": runs[:limit],
                    "total": len(runs),
                    "taskKind": task_kind,
                    "truncated": len(runs) > limit,
                },
            )
            return
        if path == "/api/crawl-run-log":
            query = parse_qs(parsed.query)
            crawl_run_id = str(query.get("id", [""])[0] or "")
            result = load_crawl_run_log(crawl_run_id)
            if result.get("ok"):
                result["newsItems"] = strategic_news_items_for_crawl_run(result.get("run"))
                result.update(strategic_news_process_items_for_crawl_run(result.get("run")))
            json_response(self, result, 200 if result.get("ok") else 404)
            return
        if path == "/api/curation-quality-records":
            query = parse_qs(parsed.query)
            run_id = str(query.get("runId", [""])[0] or "")
            result = load_curation_quality_records(run_id)
            status = 200 if result.get("ok") else (202 if result.get("pending") else 400)
            json_response(self, result, status)
            return

        if path == "/api/ai-config":
            json_response(self, {"ok": True, "config": load_ai_config(include_key=False)})
            return
        if path.startswith("/outputs/"):
            name = Path(unquote(path.removeprefix("/outputs/"))).name
            target = ROOT / name
            if not is_report_path(target):
                json_response(self, {"ok": False, "error": "file not allowed"}, 404)
                return
            self.serve_file(target, download=True)
            return
        if path.startswith("/audio/"):
            name = Path(unquote(path.removeprefix("/audio/"))).name
            target = AUDIO_DIR / name
            if not target.exists() or target.suffix.lower() not in {".wav", ".mp3"} or target.parent != AUDIO_DIR:
                json_response(self, {"ok": False, "error": "audio not found"}, 404)
                return
            self.serve_file(target)
            return
        if path.startswith("/generated-charts/"):
            target = generated_chart_path(unquote(path.removeprefix("/generated-charts/")))
            if not target or not target.exists():
                json_response(self, {"ok": False, "error": "chart not found"}, 404)
                return
            self.serve_file(target)
            return
        if path.startswith("/references/"):
            target = reference_path(unquote(path.removeprefix("/references/")))
            if not target:
                json_response(self, {"ok": False, "error": "reference not allowed"}, 404)
                return
            self.serve_reference(target)
            return
        if path.startswith("/references-raw/"):
            target = reference_path(unquote(path.removeprefix("/references-raw/")))
            if not target:
                json_response(self, {"ok": False, "error": "reference not allowed"}, 404)
                return
            self.serve_file(target)
            return
        if path.startswith("/static/"):
            self.serve_file(STATIC_DIR / path.removeprefix("/static/"))
            return
        json_response(self, {"ok": False, "error": "not found"}, 404)

    def do_POST(self):
        parsed = urlparse(self.path)
        if AUTH.handle(self, "POST", parsed):
            return
        if parsed.path.startswith("/api/") and not AUTH.authorize_api(self, parsed.path, "POST"):
            return
        if parsed.path == "/api/project-incidents/resolve":
            actor = AUTH.current_actor(self)
            if not actor:
                json_response(self, {"ok": False, "error": "登录状态已失效，请重新登录"}, 401)
                return
            incident_id = ""
            try:
                payload = read_request_json(self)
                incident_id = str(payload.get("incidentId") or "").strip()
                result = CardActionHandler(runtime_root=ROOT).mark_incident_handled_from_web(
                    incident_id,
                    str(actor.get("feishuOpenId") or ""),
                    str(actor.get("feishuUnionId") or ""),
                )
                AUTH.record_operation(
                    actor=actor,
                    action="fault.mark_handled",
                    target=incident_id,
                    details={
                        "handler_name": result.get("operator_name"),
                        "feishu_sync": result.get("feishu_sync"),
                        "sheet_row": result.get("sheet_row"),
                    },
                )
                records = load_project_incident_index(500)
                incident = next((item for item in records if item.get("incident_id") == incident_id), None)
                json_response(self, {"ok": True, "result": result, "incident": incident})
            except ValueError as exc:
                AUTH.record_operation(
                    actor=actor,
                    action="fault.mark_handled",
                    target=incident_id,
                    result="failure",
                    details={"error": str(exc)[:240]},
                )
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            except RuntimeError as exc:
                AUTH.record_operation(
                    actor=actor,
                    action="fault.mark_handled",
                    target=incident_id,
                    result="failure",
                    details={"error": str(exc)[:240]},
                )
                json_response(self, {"ok": False, "error": str(exc)}, 409)
            except Exception as exc:
                AUTH.record_operation(
                    actor=actor,
                    action="fault.mark_handled",
                    target=incident_id,
                    result="failure",
                    details={"error": "飞书同步失败"},
                )
                json_response(self, {"ok": False, "error": f"飞书同步失败：{exc}"}, 502)
            return
        if parsed.path == "/api/competitor-insight-stream":
            try:
                payload = read_request_json(self)
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
                return
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()
            try:
                result = generate_competitor_insight(payload, stream_callback=lambda event: write_sse(self, event))
                write_sse(self, {"type": "done", "ok": True, **result})
            except Exception as exc:
                write_sse(self, {"type": "error", "ok": False, "error": str(exc)})
            self.close_connection = True
            return
        if parsed.path == "/api/competitor-insight":
            try:
                result = generate_competitor_insight(read_request_json(self))
                json_response(self, {"ok": True, **result})
            except ValueError as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 503)
            return
        if parsed.path == "/api/subscriptions":
            if not is_loopback_client(str(self.client_address[0])):
                json_response(self, {"ok": False, "error": "订阅管理后台仅允许本机访问"}, 403)
                return
            try:
                payload = read_request_json(self)
                action = str(payload.get("action") or "")
                service = subscription_service()
                if action == "publish":
                    result = service.publish_entry_card(
                        target_id=str(payload.get("targetId") or ""),
                        target_type=str(payload.get("targetType") or "chat"),
                    )
                elif action == "update":
                    services = payload.get("services") if isinstance(payload.get("services"), list) else []
                    result = service.update_subscriber(
                        str(payload.get("openId") or ""),
                        services=services,
                        status=str(payload.get("status") or "active"),
                        frequency=str(payload.get("newsFrequency") or payload.get("frequency") or "once_daily"),
                        report_mode=str(payload.get("reportMode") or "pdf"),
                        news_item_limit=int(payload.get("newsItemLimit") or 10),
                        news_categories=payload.get("newsCategories"),
                    )
                elif action == "updateReportSchedule":
                    result = service.update_report_schedule(
                        days=payload.get("days"),
                        time_hm=payload.get("time"),
                        enabled=payload.get("enabled") is True,
                    )
                elif action == "updateNewsSchedule":
                    result = service.update_news_schedule(
                        enabled=payload.get("enabled") is True,
                    )
                elif action == "refreshDirectory":
                    result = service.refresh_people_directory()
                elif action == "searchPeople":
                    result = {
                        "query": str(payload.get("query") or ""),
                        "people": service.search_people_directory(str(payload.get("query") or "")),
                    }
                elif action == "searchDirectory":
                    query = str(payload.get("query") or "")
                    result = {
                        "query": query,
                        "people": service.search_people_directory(query),
                        "chats": service.search_chat_directory(query),
                    }
                elif action == "addCandidates":
                    ids = payload.get("directoryOpenIds") if isinstance(payload.get("directoryOpenIds"), list) else []
                    result = service.add_directory_candidates(ids)
                elif action == "invite":
                    ids = payload.get("callbackOpenIds") if isinstance(payload.get("callbackOpenIds"), list) else []
                    result = service.invite_users(
                        ids,
                        confirm_invite=payload.get("confirmInvite") is True,
                        invited_by="local_admin",
                    )
                elif action == "inviteTarget":
                    result = service.invite_target(
                        str(payload.get("targetId") or ""),
                        target_type=str(payload.get("targetType") or ""),
                        confirm_invite=payload.get("confirmInvite") is True,
                    )
                elif action == "pushLatest":
                    result = push_latest_subscription_content(
                        service,
                        target_open_id=str(payload.get("targetOpenId") or ""),
                        confirm_bulk=payload.get("confirmBulk") is True,
                    )
                elif action == "push":
                    result = service.push(
                        service=str(payload.get("service") or ""),
                        mode=str(payload.get("mode") or "text"),
                        path=str(payload.get("path") or ""),
                        title=str(payload.get("title") or ""),
                        body=str(payload.get("body") or ""),
                        test_open_id=str(payload.get("testOpenId") or ""),
                        target_open_id=str(payload.get("targetOpenId") or ""),
                        confirm_bulk=payload.get("confirmBulk") is True,
                    )
                else:
                    raise ValueError("未知订阅管理动作")
                json_response(self, {"ok": True, "result": result, "subscriptions": service.list_summary()})
            except ValueError as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 500)
            return
        if parsed.path == "/api/executive-intelligence/regenerate-discovery":
            if not INTELLIGENCE_INSIGHT_REFRESH_LOCK.acquire(timeout=60):
                json_response(self, {"ok": False, "error": "数据解读服务仍在处理上一项任务，请稍后重试。"}, 409)
                return
            try:
                from executive_intelligence_pipeline import regenerate_model_discovery

                payload = read_request_json(self)
                index = int(payload.get("index"))
                source_domain = str(payload.get("from") or "").strip()
                target_domain = str(payload.get("to") or "").strip()
                if not source_domain or not target_domain:
                    raise ValueError("from和to不能为空")
                json_response(self, regenerate_model_discovery(index, source_domain, target_domain))
            except (TypeError, ValueError) as exc:
                json_response(self, {"ok": False, "error": public_intelligence_error_message(exc)}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": public_intelligence_error_message(exc)}, 500)
            finally:
                INTELLIGENCE_INSIGHT_REFRESH_LOCK.release()
            return
        if parsed.path == "/api/executive-intelligence/regenerate-insight":
            if not INTELLIGENCE_INSIGHT_REFRESH_LOCK.acquire(timeout=60):
                json_response(self, {"ok": False, "error": "数据解读服务仍在处理上一项任务，请稍后重试。"}, 409)
                return
            try:
                from executive_intelligence_pipeline import regenerate_model_focus_summary

                payload = read_request_json(self)
                domain_id = str(payload.get("domain") or "").strip()
                focus_id = str(payload.get("focus") or "").strip()
                if not domain_id or not focus_id:
                    raise ValueError("domain和focus不能为空")
                if payload.get("stream"):
                    start_ndjson_response(self)
                    try:
                        result = regenerate_model_focus_summary(
                            domain_id,
                            focus_id,
                            progress=lambda message: write_ndjson_event(
                                self, {"type": "status", "message": message}
                            ),
                        )
                        for chunk in re.findall(r"[^，。；！？]+[，。；！？]?", result.get("analysis") or ""):
                            write_ndjson_event(self, {"type": "delta", "text": chunk})
                        write_ndjson_event(self, {
                            "type": "complete",
                            **{key: value for key, value in result.items() if key != "analysis"},
                        })
                    except Exception as exc:
                        write_ndjson_event(self, {
                            "type": "error",
                            "message": public_intelligence_error_message(exc),
                        })
                    self.close_connection = True
                else:
                    json_response(self, regenerate_model_focus_summary(domain_id, focus_id))
            except ValueError as exc:
                json_response(self, {"ok": False, "error": public_intelligence_error_message(exc)}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": public_intelligence_error_message(exc)}, 500)
            finally:
                INTELLIGENCE_INSIGHT_REFRESH_LOCK.release()
            return
        if parsed.path == "/api/news-review-sheet/update":
            actor = AUTH.current_actor(self)
            changes = []
            try:
                from news_review_sheet import update_review_sheet_cells

                payload = read_request_json(self)
                changes = payload.get("changes")
                if not isinstance(changes, list):
                    raise ValueError("changes 必须是数组")
                result = update_review_sheet_cells(changes)
                decision_rows = sorted({
                    int(item.get("rowNumber") or 0)
                    for item in changes
                    if isinstance(item, dict) and int(item.get("columnIndex", -1)) in {0, 1}
                    and str(item.get("before") or "") != str(item.get("value") or "")
                })
                if not int(result.get("changedCount") or 0):
                    decision_rows = []
                AUTH.record_operation(
                    actor=actor,
                    action="news_review.update",
                    target=str(result.get("sheetId") or "news-review-sheet"),
                    details={
                        "changed_count": int(result.get("changedCount") or 0),
                        "decision_rows": decision_rows,
                        "cells": [
                            {
                                "row": int(item.get("rowNumber") or 0),
                                "column": int(item.get("columnIndex", -1)),
                                "before": str(item.get("before") or "")[:120],
                                "after": str(item.get("value") or "")[:120],
                            }
                            for item in changes[:200]
                            if isinstance(item, dict)
                        ],
                        "feishu_readback": bool(result.get("readbackVerified")),
                    },
                )
                json_response(self, {"ok": True, **attach_news_review_actors(result)})
            except (ValueError, RuntimeError) as exc:
                AUTH.record_operation(
                    actor=actor,
                    action="news_review.update",
                    target="news-review-sheet",
                    result="failure",
                    details={"error": str(exc)[:240]},
                )
                json_response(self, {"ok": False, "error": str(exc)}, 409)
            except Exception as exc:
                AUTH.record_operation(
                    actor=actor,
                    action="news_review.update",
                    target="news-review-sheet",
                    result="failure",
                    details={"error": str(exc)[:240]},
                )
                json_response(self, {"ok": False, "error": str(exc)}, 500)
            return
        if parsed.path == "/api/crawl":
            if not CRAWL_PIPELINE_LOCK.acquire(blocking=False):
                json_response(
                    self,
                    {"ok": False, "error": "已有手动全量爬虫正在运行，请等待其完成。", "active": dict(CRAWL_PIPELINE_STATE)},
                    409,
                )
                return
            CRAWL_PIPELINE_STATE.clear()
            CRAWL_PIPELINE_STATE.update({"status": "running", "startedAt": datetime.now().astimezone().isoformat(timespec="seconds")})
            try:
                json_response(self, run_crawl())
            finally:
                CRAWL_PIPELINE_STATE.clear()
                CRAWL_PIPELINE_LOCK.release()
            return
        if parsed.path == "/api/crawl-stream":
            # crawl-pipeline-lock:v1
            if not CRAWL_PIPELINE_LOCK.acquire(blocking=False):
                json_response(
                    self,
                    {"ok": False, "error": "已有手动全量爬虫正在运行，请等待其完成。", "active": dict(CRAWL_PIPELINE_STATE)},
                    409,
                )
                return
            CRAWL_PIPELINE_STATE.clear()
            CRAWL_PIPELINE_STATE.update({"status": "starting", "startedAt": datetime.now().astimezone().isoformat(timespec="seconds")})
            crawl_run_id = ""
            proc = None
            try:
                self.send_response(200)
                self.send_header("Content-Type", "text/event-stream; charset=utf-8")
                self.send_header("Cache-Control", "no-cache")
                self.end_headers()

                started = time.time()
                started_at_hkt = datetime.now().astimezone().isoformat(timespec="seconds")
                crawl_scope = "全量（第2-34行）"
                started_record = start_crawl_run(trigger="手动全量", scope=crawl_scope)
                crawl_run_id = str(started_record["crawl_run_id"])
                CRAWL_PIPELINE_STATE.update({"status": "running", "crawlRunId": crawl_run_id, "startedAt": started_at_hkt, "scope": crawl_scope})
                self._crawl_stream_log_path = Path(started_record["stream_log_path"])
                self._crawl_stream_mirror_path = ROOT / "latest_crawl_stream.log"
                start_task_lifecycle_monitor(self, "crawl", crawl_run_id, "网页抓取")
                try:
                    self._crawl_stream_mirror_path.write_text("", encoding="utf-8")
                except OSError:
                    pass
                write_sse(
                    self,
                    {
                        "type": "run_start",
                        "crawlRunId": crawl_run_id,
                        "startedAt": started_at_hkt,
                        "trigger": "手动全量",
                        "scope": crawl_scope,
                    },
                )
                crawl_env = os.environ.copy()
                crawl_env.pop("CMHK_ROWS", None)
                crawl_env["CMHK_CRAWL_TRIGGER"] = "手动全量"
                crawl_env["CMHK_CRAWL_SCOPE"] = "全量（第2-34行）"
                proc = subprocess.Popen(
                    [sys.executable, "-u", str(ROOT / "crawl.py")],
                    cwd=str(ROOT),
                    env=crawl_env,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    text=True,
                    bufsize=1,
                )
                self._task_worker_pid = proc.pid
                for line in proc.stdout:
                    write_sse(self, sse_payload_from_process_line(line.strip()))

                proc.wait()
                self._task_worker_pid = 0
                self._task_monitor_phase = "飞书同步"
                self._task_monitor_detail = "网页抓取进程已结束，正在同步结果并执行后续审核。"
                log_sheet_id = ""
                log_sheet_title = ""
                crawl_failed_count = 0
                sync_result = {}
                metrics_refresh = {}
                trace_sync = {}

                # Sync to Feishu after full crawl
                if proc.returncode == 0 and (ROOT / "write_payload.json").exists():
                    sync_proc = subprocess.run(
                        [sys.executable, str(ROOT / "daily_crawl_and_write.py"), "--sync-only"],
                        env=crawl_env,
                        capture_output=True,
                        text=True,
                    )
                    if sync_proc.returncode != 0:
                        write_sse(self, {"type": "log", "text": f"同步飞书失败: {sync_proc.stderr[-500:]}"})
                        proc.returncode = sync_proc.returncode
                    else:
                        sync_result = json_object_from_output(sync_proc.stdout)
                        log_sheet_id = str(sync_result.get("log_sheet_id") or "")
                        log_sheet_title = str(sync_result.get("log_sheet_title") or "")
                        write_sse(
                            self,
                            {
                                "type": "log",
                                "text": "✅ 飞书表格同步成功！"
                                + (f" 日志页：{log_sheet_title}" if log_sheet_title else ""),
                            },
                        )

                    # Update supplementary JSON configs with newly extracted data
                    update_proc = subprocess.run([sys.executable, str(ROOT / "update_sources_from_crawl.py")], capture_output=True, text=True)
                    if update_proc.returncode != 0:
                        write_sse(self, {"type": "log", "text": f"⚠️ 业绩补充桥接更新异常: {update_proc.stderr[-200:]}"})
                    else:
                        if update_proc.stdout.strip():
                            write_sse(self, {"type": "log", "text": f"ℹ️ 业绩补充配置同步：{update_proc.stdout.strip()}"})

                    performance_sync = run_carrier_performance_sync()
                    if performance_sync["ok"]:
                        payload = json.dumps(
                            {"type": "log", "text": "✅ 运营商业绩摘要补充页已同步并通过五类字段校验。"},
                            ensure_ascii=False,
                        )
                    else:
                        payload = json.dumps(
                            {
                                "type": "log",
                                "text": "运营商业绩摘要补充页同步失败: "
                                + (performance_sync["stderr"] or performance_sync["stdout"])[-500:],
                            },
                            ensure_ascii=False,
                        )
                        proc.returncode = proc.returncode or performance_sync["returnCode"]
                    write_sse(self, json.loads(payload))

                    payload = json.dumps(
                        {
                            "type": "log",
                            "text": "开始多 Agent 数据整理：来源分类、事实抽取、主体校验、质量审计、冲突仲裁和缺口补爬...",
                        },
                        ensure_ascii=False,
                    )
                    write_sse(self, json.loads(payload))
                    metrics_refresh = stream_company_metrics_refresh(self)
                    if metrics_refresh["ok"]:
                        summary = metrics_refresh["summary"]
                        payload = json.dumps(
                            {
                                "type": "log",
                                "text": (
                                    "✅ 公司指标页已更新："
                                    f"{summary.get('companies', 0)} 家公司、"
                                    f"{summary.get('metrics', 0)} 类指标、"
                                    f"{summary.get('records', 0)} 条通过校验的记录。"
                                ),
                            },
                            ensure_ascii=False,
                        )
                    else:
                        payload = json.dumps(
                            {
                                "type": "log",
                                "text": "❌ 公司指标页 AI 整理失败: "
                                + (metrics_refresh["stderr"] or metrics_refresh["stdout"])[-500:],
                            },
                            ensure_ascii=False,
                        )
                        proc.returncode = proc.returncode or metrics_refresh["returnCode"]
                    write_sse(self, json.loads(payload))

                    if metrics_refresh["ok"] and log_sheet_id:
                        latest_curation = load_curation_status()
                        agent_run_id = str(latest_curation.get("run_id") or "")
                        write_sse(
                            self,
                            {
                                "type": "agent_trace",
                                "trace": {
                                    "ts": datetime.now().astimezone().isoformat(timespec="seconds"),
                                    "run_id": agent_run_id,
                                    "node": "飞书审计日志",
                                    "phase": "tool_call",
                                    "event_type": "tool_call",
                                    "message": f"将 Agent 处理流程和结果写入飞书日志页 {log_sheet_title or log_sheet_id}。",
                                    "tool": "daily_crawl_and_write.py --append-agent-trace",
                                    "input": {
                                        "sheetId": log_sheet_id,
                                        "sheetTitle": log_sheet_title,
                                        "runId": agent_run_id,
                                    },
                                },
                            },
                        )
                        trace_sync = append_agent_trace_to_feishu_log(log_sheet_id, agent_run_id)
                        trace_result = trace_sync.get("result") or {}
                        write_sse(
                            self,
                            {
                                "type": "agent_trace",
                                "trace": {
                                    "ts": datetime.now().astimezone().isoformat(timespec="seconds"),
                                    "run_id": agent_run_id,
                                    "node": "飞书审计日志",
                                    "phase": "tool_result",
                                    "event_type": "tool_result",
                                    "message": (
                                        f"Agent 流程已写入飞书，共 {trace_result.get('trace_rows', 0)} 条并完成回读校验。"
                                        if trace_sync["ok"]
                                        else "Agent 流程写入飞书失败。"
                                    ),
                                    "tool": "daily_crawl_and_write.py --append-agent-trace",
                                    "result": {
                                        "ok": trace_sync["ok"],
                                        "sheetId": log_sheet_id,
                                        "sheetTitle": log_sheet_title,
                                        "range": trace_result.get("range", ""),
                                        "traceRows": trace_result.get("trace_rows", 0),
                                        "error": (trace_sync["stderr"] or trace_sync["stdout"])[-500:]
                                        if not trace_sync["ok"]
                                        else "",
                                    },
                                },
                            },
                        )
                        if not trace_sync["ok"]:
                            write_sse(
                                self,
                                {
                                    "type": "log",
                                    "text": (
                                        "⚠️ 爬取、主表同步和 Agent 整理均已完成；"
                                        "仅飞书审计日志追加失败，可稍后重试，不影响本轮数据结果。"
                                    ),
                                },
                            )
                    elif metrics_refresh["ok"]:
                        write_sse(
                            self,
                            {
                                "type": "log",
                                "text": "⚠️ Agent 已完成，但未取得本次飞书日志页 ID，未能追加 Agent 审计区块。",
                            },
                        )

                try:
                    run_log_path = ROOT / "run_log.json"
                    if run_log_path.exists():
                        with run_log_path.open("r", encoding="utf-8") as f:
                            run_log_data = json.load(f)
                        success_items = []
                        failure_items = []
                        for item in run_log_data:
                            url = item.get("url", "")
                            status = int(item.get("http_status") or 0)
                            used_fallback = str(
                                item.get("evidence_fallback_used") or ""
                            ).lower() in {"1", "true", "yes"}
                            if 200 <= status < 400 and not used_fallback:
                                success_items.append({"url": url, "reason": "OK"})
                            else:
                                reason = (
                                    item.get("fallback_reason")
                                    if used_fallback
                                    else item.get("error")
                                    or item.get("skip_reason")
                                    or f"HTTP {status}"
                                )
                                failure_items.append({"url": url, "reason": reason})
                        crawl_failed_count = len(failure_items)
                        summary_payload = json.dumps({
                            "type": "crawl_summary",
                            "success": success_items,
                            "failed": failure_items,
                            "total": len(run_log_data)
                        }, ensure_ascii=False)
                        write_sse(self, json.loads(summary_payload))
                except Exception as e:
                    pass

                duration_ms = round((time.time() - started) * 1000)
                runtime_sync_script = ROOT / "sync_scheduler_runtime.sh"
                if proc.returncode == 0 and runtime_sync_script.exists():
                    try:
                        runtime_sync = subprocess.run(
                            [str(runtime_sync_script)],
                            cwd=str(ROOT),
                            capture_output=True,
                            text=True,
                            timeout=120,
                        )
                    except Exception as exc:
                        write_sse(
                            self,
                            {
                                "type": "log",
                                "text": (
                                    "后台调度运行副本同步异常，已保留本轮结果并继续登记："
                                    f"{type(exc).__name__}: {exc}"
                                ),
                            },
                        )
                    else:
                        if runtime_sync.returncode != 0:
                            write_sse(
                                self,
                                {
                                    "type": "log",
                                    "text": f"后台调度运行副本同步失败，已继续登记：{runtime_sync.stderr[-500:]}",
                                },
                            )
                crawl_run_record = register_crawl_run(
                    crawl_return_code=proc.returncode,
                    duration_ms=duration_ms,
                    sync_result=sync_result,
                    metrics_refresh=metrics_refresh,
                    trace_sync=trace_sync,
                    trigger="手动全量",
                    scope=crawl_scope,
                    crawl_run_id=crawl_run_id,
                    started_at_hkt=started_at_hkt,
                    stream_log_path=self._crawl_stream_log_path,
                )
                write_sse(
                    self,
                    {
                        "type": "log",
                        "text": (
                            "爬虫运行日志索引已保存："
                            f"{crawl_run_record.get('crawl_run_id')}；"
                            f"飞书日志页：{(crawl_run_record.get('feishu') or {}).get('log_sheet_title') or '未写入'}。"
                        ),
                    },
                )
                write_sse(self, {
                    "type": "done",
                    "ok": proc.returncode == 0,
                    "completedWithWarnings": proc.returncode == 0 and crawl_failed_count > 0,
                    "failedUrlCount": crawl_failed_count,
                    "durationMs": duration_ms,
                    "status": build_status(),
                    "crawlRunRegistry": crawl_run_record,
                })
                return
            except Exception as exc:
                if proc is not None and proc.poll() is None:
                    proc.terminate()
                    try:
                        proc.wait(timeout=5)
                    except subprocess.TimeoutExpired:
                        proc.kill()
                detail = f"爬虫流水线异常中断：{type(exc).__name__}: {exc}"
                if crawl_run_id:
                    mark_crawl_run_interrupted(crawl_run_id, detail)
                try:
                    write_sse(self, {"type": "log", "text": "[任务中断] " + detail})
                    write_sse(self, {"type": "done", "ok": False, "interrupted": True, "message": detail})
                except (BrokenPipeError, ConnectionResetError, OSError):
                    pass
                return
            finally:
                stop_task_lifecycle_monitor(self)
                CRAWL_PIPELINE_STATE.clear()
                CRAWL_PIPELINE_LOCK.release()
        if parsed.path == "/api/generate-stream":
            stream_report_generation(self, "generate_weekly_report.py", "weekly")
            return

        if parsed.path == "/api/generate-carrier-performance-stream":
            stream_report_generation(
                self,
                "generate_carrier_performance_report.py",
                "carrier-performance",
            )
            return

        if parsed.path == "/api/generate":
            json_response(self, run_report_generation())
            return

        if parsed.path == "/api/generate-carrier-performance":
            json_response(self, run_carrier_performance_generation())
            return

        if parsed.path == "/api/agent-datasets/upload":
            try:
                result = write_uploaded_knowledge_dataset(read_request_json(self))
                json_response(self, {"ok": True, **result, "datasets": list_knowledge_datasets()})
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return

        if parsed.path == "/api/audio/generate":
            try:
                body = read_request_json(self)
                target = report_target_from_rel(str(body.get("path") or ""))
                if not target:
                    raise ValueError("文件不存在或不允许生成音频")
                task, created = start_audio_generation_task(target, bool(body.get("force", False)))
                json_response(
                    self,
                    {
                        "ok": True,
                        "queued": created,
                        "alreadyRunning": not created,
                        "task": task,
                    },
                    202,
                )
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/report-file":
            try:
                json_response(self, {"ok": True, "status": update_report_file(read_request_json(self))})
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/delete-files":
            try:
                body = read_request_json(self)
                result = delete_report_files(body.get("paths", []))
                json_response(self, {"ok": True, **result})
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/ai-config":
            try:
                json_response(self, {"ok": True, "config": save_ai_config(read_request_json(self)), "status": build_status()})
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/chat-image-analyze":
            try:
                json_response(self, {"ok": True, **analyze_chat_image(read_request_json(self))})
            except urllib.error.HTTPError as exc:
                detail = exc.read().decode("utf-8", errors="ignore")[:600]
                json_response(self, {"ok": False, "error": f"视觉模型返回 HTTP {exc.code}：{detail}"}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/chat-audio-transcribe":
            try:
                content_length = int(self.headers.get("Content-Length") or 0)
                if content_length > (CHAT_AUDIO_MAX_BYTES * 2):
                    raise ValueError("单次语音不能超过 20 MB")
                json_response(self, {"ok": True, **transcribe_chat_audio(read_request_json(self))})
            except urllib.error.HTTPError as exc:
                detail = exc.read().decode("utf-8", errors="ignore")[:600]
                json_response(self, {"ok": False, "error": f"语音模型返回 HTTP {exc.code}：{detail}"}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/ai-models":
            try:
                payload = read_request_json(self)
                base_url = str(payload.get("base_url") or "").strip().rstrip("/")
                if not re.match(r"^https?://", base_url, flags=re.I):
                    raise ValueError("Base URL 必须以 http:// 或 https:// 开头")
                if not is_internal_ai_base_url(base_url):
                    raise ValueError("只能访问公司内网模型服务")
                saved = load_ai_config(include_key=True)
                api_key = str(payload.get("api_key") or saved.get("api_key") or "").strip()
                if not api_key:
                    raise ValueError("请输入 API Key，或先保存一个有效 Key")
                request = urllib.request.Request(
                    f"{base_url}/models",
                    headers={"Authorization": f"Bearer {api_key}", "Accept": "application/json"},
                    method="GET",
                )
                with urllib.request.urlopen(request, timeout=30) as response:
                    model_payload = json.loads(response.read().decode("utf-8"))
                models = sorted(
                    {
                        str(item.get("id") or "").strip()
                        for item in (model_payload.get("data") or [])
                        if isinstance(item, dict) and str(item.get("id") or "").strip()
                    },
                    key=str.lower,
                )
                if not models:
                    raise ValueError("服务已连接，但没有返回可用模型")
                json_response(self, {"ok": True, "models": models, "count": len(models)})
            except urllib.error.HTTPError as exc:
                detail = exc.read().decode("utf-8", errors="ignore")[:600]
                json_response(self, {"ok": False, "error": f"模型服务返回 HTTP {exc.code}：{detail}"}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/ai-test":
            started = time.monotonic()
            try:
                payload = read_request_json(self)
                config = load_ai_config(include_key=True)
                provider = str(payload.get("provider") or config.get("provider") or "deepseek").lower()
                base_url = str(payload.get("base_url") or config.get("base_url") or "").strip().rstrip("/")
                model = str(payload.get("model") or config.get("model") or "").strip()
                api_key = str(payload.get("api_key") or config.get("api_key") or "").strip()
                if not base_url or not model or not api_key:
                    raise ValueError("Base URL、模型和 API Key 均不能为空")
                if not re.match(r"^https?://", base_url, flags=re.I):
                    raise ValueError("Base URL 必须以 http:// 或 https:// 开头")
                if not is_internal_ai_base_url(base_url):
                    raise ValueError("只能访问公司内网模型服务")
                if provider == "openai":
                    url = f"{base_url}/responses"
                    body = {"model": model, "input": "Reply OK", "max_output_tokens": 16}
                else:
                    url = f"{base_url}/chat/completions"
                    body = {
                        "model": model,
                        "messages": [{"role": "user", "content": "只回复OK"}],
                        "max_tokens": 32,
                        "stream": False,
                    }
                body.update(config.get("extra_parameters") or {})
                request = urllib.request.Request(
                    url,
                    data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    method="POST",
                )
                wait_for_internal_ai_slot("ai-settings-test")
                with urllib.request.urlopen(request, timeout=45) as response:
                    response.read()
                result = {
                    "ok": True,
                    "provider": provider,
                    "model": model,
                    "latency_ms": round((time.monotonic() - started) * 1000),
                }
                json_response(self, {"ok": True, "result": result, "status": build_status()})
            except urllib.error.HTTPError as exc:
                detail = exc.read().decode("utf-8", errors="ignore")[:600]
                json_response(self, {"ok": False, "error": f"模型服务返回 HTTP {exc.code}：{detail}"}, 400)
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/rag-token-estimate":
            payload = read_request_json(self)
            text = str(payload.get("text") or "")
            model = str(payload.get("model") or "")
            json_response(
                self,
                {
                    "ok": True,
                    "tokens": estimate_tokens(text, model=model or None),
                    "chars": len(text),
                    "model": model or None,
                    "counter": "tiktoken_or_heuristic",
                },
            )
            return
        if parsed.path == "/api/agent-memory/delete":
            payload = read_request_json(self)
            memory_id = str(payload.get("id") or "")
            json_response(self, {"ok": delete_memory(memory_id), "id": memory_id})
            return
        if parsed.path == "/api/chat-approval":
            payload = read_request_json(self)
            request_id = str(payload.get("requestId") or "")
            action_id = str(payload.get("actionId") or "")
            decision = "allow" if payload.get("decision") == "allow" else "deny"
            resolved = resolve_chat_approval(request_id, action_id, decision)
            json_response(
                self,
                {"ok": resolved, "requestId": request_id, "actionId": action_id, "decision": decision},
                200 if resolved else 404,
            )
            return
        if parsed.path == "/api/chat-threads":
            try:
                payload = read_request_json(self)
                action = str(payload.get("action") or "save")
                if action == "delete":
                    thread_id = str(payload.get("id") or "")
                    json_response(self, {"ok": delete_chat_thread(thread_id), "threads": chat_thread_summaries()})
                elif action == "pin":
                    thread_id = str(payload.get("id") or "")
                    pinned = bool(payload.get("pinned"))
                    thread = set_chat_thread_pinned(thread_id, pinned)
                    json_response(
                        self,
                        {"ok": bool(thread), "thread": thread, "threads": chat_thread_summaries()},
                        200 if thread else 404,
                    )
                else:
                    thread = upsert_chat_thread(payload)
                    json_response(self, {"ok": True, "thread": thread, "threads": chat_thread_summaries()})
            except Exception as exc:
                json_response(self, {"ok": False, "error": str(exc)}, 400)
            return
        if parsed.path == "/api/chat":
            json_response(self, {"ok": False, "error": "deprecated API, use stream"}, 404)
            return
        if parsed.path == "/api/chat-stream":
            payload = read_request_json(self)
            message = str(payload.get("message") or "")
            request_id = re.sub(r"[^A-Za-z0-9_.:-]", "", str(payload.get("requestId") or ""))[:160]
            if not request_id:
                request_id = f"chat-{uuid.uuid4().hex}"
            web_search_enabled = bool(payload.get("webSearchEnabled"))
            thinking_enabled = bool(payload.get("thinkingEnabled"))
            selected_skill_ids = payload.get("selectedSkillIds")
            if not isinstance(selected_skill_ids, list):
                selected_skill_ids = []
            selected_dataset_ids = payload.get("selectedDatasetIds")
            if not isinstance(selected_dataset_ids, list):
                selected_dataset_ids = []
            approved_action_ids = payload.get("approvedActionIds")
            if not isinstance(approved_action_ids, list):
                approved_action_ids = []
            conversation_history = payload.get("conversationHistory")
            if not isinstance(conversation_history, list):
                conversation_history = []
            emit_context_events = bool(payload.get("emitContextEvents", True))
            loaded_skill_ids = payload.get("loadedSkillIds")
            if not isinstance(loaded_skill_ids, list):
                loaded_skill_ids = []
            active_thread_id = re.sub(
                r"[^A-Za-z0-9_.:-]", "", str(payload.get("threadId") or "")
            )[:160]
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            for event in stream_agent_with_approvals(
                message,
                request_id=request_id,
                force_web_search=web_search_enabled,
                selected_skill_ids=[str(item) for item in selected_skill_ids],
                selected_dataset_ids=[str(item) for item in selected_dataset_ids],
                thinking_enabled=thinking_enabled,
                approved_action_ids=[str(item) for item in approved_action_ids],
                conversation_history=conversation_history,
                emit_context_events=emit_context_events,
                loaded_skill_ids=[str(item) for item in loaded_skill_ids],
                runtime_context=request_runtime_context(self),
                active_thread_id=active_thread_id,
            ):
                body = json.dumps(event, ensure_ascii=False)
                self.wfile.write(f"data: {body}\n\n".encode("utf-8"))
                self.wfile.flush()
                if event.get("type") == "done":
                    self.close_connection = True
            return
        json_response(self, {"ok": False, "error": "not found"}, 404)

    def log_message(self, fmt: str, *args) -> None:
        print(f"[web] {self.address_string()} - {fmt % args}")

    def serve_file(self, path: Path, download: bool = False) -> None:
        if not path.exists() or not path.is_file():
            json_response(self, {"ok": False, "error": "file not found"}, 404)
            return
        body = path.read_bytes()
        content_type = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
        if content_type.startswith("text/") or path.suffix.lower() in {".md", ".tsv", ".json"}:
            content_type = f"{content_type}; charset=utf-8"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        if path.suffix.lower() == ".html":
            self.send_header("Cache-Control", "no-store")
        elif path.suffix.lower() in {".css", ".js"}:
            self.send_header("Cache-Control", "no-cache, must-revalidate")
        if download:
            self.send_header("Content-Disposition", self.download_disposition(path))
        self.end_headers()
        self.wfile.write(body)

    def serve_reference(self, path: Path) -> None:
        if not path.exists() or not path.is_file():
            json_response(self, {"ok": False, "error": "file not found"}, 404)
            return
        suffix = path.suffix.lower()
        if suffix in {".md", ".tsv", ".json", ".txt", ".docx", ".pdf"}:
            raw = read_display_text(path)
            title = path.name
            raw_ref = quote(path.relative_to(ROOT).as_posix(), safe="/")
            body = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{escape(title)}</title>
  <style>
    body {{ margin: 0; padding: 24px; background: #f8fafc; color: #172033; font: 14px/1.7 -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; }}
    .bar {{ position: sticky; top: 0; margin: -24px -24px 18px; padding: 14px 24px; background: rgba(248, 250, 252, 0.96); border-bottom: 1px solid #d8e3ee; backdrop-filter: blur(8px); }}
    h1 {{ margin: 0 0 4px; font-size: 18px; }}
    a {{ color: #0067b1; font-weight: 700; text-decoration: none; }}
    pre {{ margin: 0; padding: 18px; overflow: auto; white-space: pre-wrap; word-break: break-word; background: #fff; border: 1px solid #d8e3ee; border-radius: 8px; box-shadow: 0 1px 4px rgba(15, 29, 46, 0.06); font: 13px/1.75 ui-monospace, SFMono-Regular, Menlo, Consolas, "PingFang SC", "Microsoft YaHei", monospace; }}
  </style>
</head>
<body>
  <div class="bar">
    <h1>{escape(title)}</h1>
    <a href="/references-raw/{raw_ref}" target="_blank" rel="noopener noreferrer">打开原始文件</a>
  </div>
  <pre>{escape(raw)}</pre>
</body>
</html>""".encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        self.serve_file(path)

    def serve_reference_head(self, path: Path) -> None:
        if not path.exists() or not path.is_file():
            self.send_response(404)
            self.end_headers()
            return
        if path.suffix.lower() in {".md", ".tsv", ".json", ".txt", ".docx", ".pdf"}:
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.end_headers()
            return
        self.serve_head(path)

    def serve_head(self, path: Path, download: bool = False) -> None:
        if not path.exists() or not path.is_file():
            self.send_response(404)
            self.end_headers()
            return
        content_type = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
        if content_type.startswith("text/") or path.suffix.lower() in {".md", ".tsv", ".json"}:
            content_type = f"{content_type}; charset=utf-8"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(path.stat().st_size))
        if path.suffix.lower() == ".html":
            self.send_header("Cache-Control", "no-store")
        elif path.suffix.lower() in {".css", ".js"}:
            self.send_header("Cache-Control", "no-cache, must-revalidate")
        if download:
            self.send_header("Content-Disposition", self.download_disposition(path))
        self.end_headers()


def main() -> None:
    interrupted = reconcile_interrupted_crawl_runs()
    if interrupted:
        print(f"Reconciled {len(interrupted)} interrupted crawl run(s)", flush=True)
    interrupted_tasks = reconcile_interrupted_general_tasks()
    if interrupted_tasks:
        print(f"Reconciled {len(interrupted_tasks)} interrupted report task(s)", flush=True)
    recovery_candidates = pending_interrupted_general_task_retries()
    if recovery_candidates:
        scheduled_retries = schedule_interrupted_general_task_retries(recovery_candidates)
        print(f"Scheduled {len(scheduled_retries)} interrupted task retry/retries", flush=True)
    corrected_tasks = reconcile_misclassified_general_tasks()
    if corrected_tasks:
        print(f"Corrected {len(corrected_tasks)} misclassified report task(s)", flush=True)
    start_scheduler_with_backend()
    port = int(os.environ.get("PORT", "8765"))
    host = os.environ.get("HOST", "0.0.0.0")
    server = ThreadingHTTPServer((host, port), AppHandler)
    print(f"Weekly report UI: http://{host}:{port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
